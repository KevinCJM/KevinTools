import tempfile
import unittest
from pathlib import Path

from docx import Document

from src.template_parser import TemplateParser


class TemplateParserAbstractEnTests(unittest.TestCase):
    def test_parse_abstract_en_roles(self) -> None:
        document = Document()
        document.add_paragraph("摘要")
        document.add_paragraph("这是中文摘要内容")
        document.add_paragraph("")
        document.add_paragraph("Abstract")
        document.add_paragraph("This is the English abstract")

        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "abstracts.docx"
            document.save(str(path))
            parser = TemplateParser()
            result = parser.parse(str(path))

        self.assertIn("abstract_title", result.roles)
        self.assertIn("abstract_body", result.roles)
        self.assertIn("abstract_en_title", result.roles)
        self.assertIn("abstract_en_body", result.roles)


if __name__ == "__main__":
    unittest.main()
