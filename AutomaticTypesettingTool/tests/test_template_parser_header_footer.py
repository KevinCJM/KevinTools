import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.shared import Pt

from src.template_parser import TemplateParser


class TemplateParserHeaderFooterTests(unittest.TestCase):
    def test_parse_header_footer_styles(self) -> None:
        document = Document()
        section = document.sections[0]

        header = section.header
        header.paragraphs[0].text = "Header"
        header_run = header.paragraphs[0].runs[0]
        header_run.font.name = "Times New Roman"
        header_run.font.size = Pt(9)

        footer = section.footer
        footer.paragraphs[0].text = "Footer"
        footer_run = footer.paragraphs[0].runs[0]
        footer_run.font.name = "Arial"
        footer_run.font.size = Pt(8)

        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "header_footer.docx"
            document.save(str(path))
            parser = TemplateParser()
            result = parser.parse(str(path))

        header_footer = result.meta.get("header_footer", {})
        sections = header_footer.get("sections", [])
        self.assertTrue(sections)
        headers = sections[0].get("headers", [])
        footers = sections[0].get("footers", [])
        self.assertTrue(headers)
        self.assertTrue(footers)

        header_style = headers[0].get("style", {})
        footer_style = footers[0].get("style", {})
        self.assertEqual(header_style.get("font_name"), "Times New Roman")
        self.assertEqual(footer_style.get("font_name"), "Arial")
        self.assertEqual(header_style.get("font_size_name"), "小五")


if __name__ == "__main__":
    unittest.main()
