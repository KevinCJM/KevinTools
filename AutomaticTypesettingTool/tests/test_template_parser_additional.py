import tempfile
import types
import unittest
from datetime import datetime
from pathlib import Path
from unittest import mock
from zipfile import ZipFile

from docx import Document
from lxml import etree
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt

from src import config
from src.style_reader import FontSpec, ResolvedStyle, StyleDefinition
from src.style_rule import StyleRule
from src.template_parser import (
    ParseLogState,
    SampleStats,
    TemplateParser,
    _add_sample,
    _apply_fallbacks,
    _build_style_rule,
    _choose_best_candidate,
    _collect_paragraph_samples,
    _collect_body_candidates_by_stack,
    _describe_candidate,
    _detect_heading_levels,
    _parse_heading_level_from_name,
    _ensure_readable_file,
    _ensure_required_roles,
    _extract_paragraph_alignment,
    _extract_paragraph_line_spacing,
    _extract_paragraph_outline_level,
    _extract_run_bold,
    _extract_run_font_size,
    _extract_run_fonts,
    _length_to_pt,
    _load_role_map,
    _match_role,
    _match_special_role_by_style_name,
    _parse_int,
    _parse_theme_map,
    _read_docx_parts,
    _read_paragraph_line_rule,
    _resolve_line_spacing,
    _resolve_role_conflicts,
    _sample_mode,
    _tag_warnings_for_style,
    _validate_strict,
    _warn,
    _write_log,
    _W_NS,
)


FIXTURES_DIR = Path(__file__).resolve().parents[1] / "tests" / "fixtures"


def _fixture(name: str) -> Path:
    return FIXTURES_DIR / name


def _make_ppr(line_rule: str | None = None, line: int | None = None, outline: int | None = None):
    p_pr = etree.Element(f"{{{_W_NS}}}pPr")
    if line_rule is not None or line is not None:
        spacing = etree.SubElement(p_pr, f"{{{_W_NS}}}spacing")
        if line_rule is not None:
            spacing.set(f"{{{_W_NS}}}lineRule", line_rule)
        if line is not None:
            spacing.set(f"{{{_W_NS}}}line", str(line))
    if outline is not None:
        outline_elem = etree.SubElement(p_pr, f"{{{_W_NS}}}outlineLvl")
        outline_elem.set(f"{{{_W_NS}}}val", str(outline))
    return types.SimpleNamespace(pPr=p_pr)


def _make_run(
    text: str,
    font_name: str | None = None,
    size: object | None = None,
    bold: bool | None = None,
    ascii_name: str | None = None,
    hansi_name: str | None = None,
    east_asia_name: str | None = None,
):
    r_pr = None
    if ascii_name or hansi_name or east_asia_name:
        r_pr = etree.Element(f"{{{_W_NS}}}rPr")
        r_fonts = etree.SubElement(r_pr, f"{{{_W_NS}}}rFonts")
        if ascii_name:
            r_fonts.set(f"{{{_W_NS}}}ascii", ascii_name)
        if hansi_name:
            r_fonts.set(f"{{{_W_NS}}}hAnsi", hansi_name)
        if east_asia_name:
            r_fonts.set(f"{{{_W_NS}}}eastAsia", east_asia_name)
    element = types.SimpleNamespace(rPr=r_pr)
    font = types.SimpleNamespace(name=font_name, size=size, bold=bold)
    return types.SimpleNamespace(_element=element, font=font, bold=bold, text=text)


def _make_paragraph(
    style_id: str | None,
    style_name: str | None = None,
    style_alignment: object | None = None,
    alignment: object | None = None,
    line_rule: object | None = None,
    line_value: object | None = None,
    space_before: object | None = None,
    space_after: object | None = None,
    ppr: object | None = None,
    runs: list | None = None,
    text: str | None = None,
):
    if style_id is None:
        style = None
    else:
        if style_name is None:
            style_name = style_id
        paragraph_format = (
            types.SimpleNamespace(alignment=style_alignment)
            if style_alignment is not None
            else None
        )
        style = types.SimpleNamespace(
            style_id=style_id,
            name=style_name,
            paragraph_format=paragraph_format,
        )
    paragraph_format = types.SimpleNamespace(
        line_spacing_rule=line_rule,
        line_spacing=line_value,
        space_before=space_before,
        space_after=space_after,
    )
    return types.SimpleNamespace(
        style=style,
        alignment=alignment,
        paragraph_format=paragraph_format,
        _element=ppr or types.SimpleNamespace(pPr=None),
        runs=runs or [],
        text=text or "",
    )


def _make_resolved(
    style_id: str,
    name: str,
    font_name: str | None = None,
    font_size_pt: float | None = None,
    bold: bool | None = None,
    alignment: str | None = None,
    space_before_pt: float | None = None,
    space_after_pt: float | None = None,
    line_rule: str | None = None,
    line_twips: int | None = None,
):
    fonts = FontSpec(ascii=font_name)
    return ResolvedStyle(
        style_id=style_id,
        name=name,
        fonts=fonts,
        font_name=font_name,
        font_size_pt=font_size_pt,
        bold=bold,
        alignment=alignment,
        space_before_pt=space_before_pt,
        space_after_pt=space_after_pt,
        line_rule=line_rule,
        line_twips=line_twips,
        outline_level=None,
    )


def _make_style_def(style_id: str, name: str | None, outline_level: int | None) -> StyleDefinition:
    fonts = FontSpec()
    return StyleDefinition(
        style_id=style_id,
        name=name,
        based_on=None,
        fonts=fonts,
        font_size_pt=None,
        bold=None,
        alignment=None,
        space_before_pt=None,
        space_after_pt=None,
        line_rule=None,
        line_twips=None,
        outline_level=outline_level,
    )


class TemplateParserAdditionalTests(unittest.TestCase):
    def setUp(self) -> None:
        config.ensure_base_dirs()
        for log_file in config.LOG_DIR.glob("template_parser_*.log"):
            log_file.unlink()

    def test_ensure_readable_file_directory(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            with self.assertRaises(IsADirectoryError):
                _ensure_readable_file(Path(tmpdir))

    def test_read_docx_parts_missing_styles(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "missing_styles.docx"
            with ZipFile(path, "w") as archive:
                archive.writestr("word/document.xml", "<w:document/>")
            with self.assertRaises(ValueError):
                _read_docx_parts(path)

    def test_parse_theme_map_no_scheme(self) -> None:
        theme_xml = b"""<?xml version="1.0" encoding="UTF-8"?>
<a:theme xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" />
"""
        self.assertEqual(_parse_theme_map(None), {})
        self.assertEqual(_parse_theme_map(theme_xml), {})

    def test_load_role_map_validation_errors(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            role_path = Path(tmpdir)
            with self.assertRaises(ValueError):
                _load_role_map(None, role_path, False, {"Heading1"}, {"heading 1"}, None)

        with tempfile.TemporaryDirectory() as tmpdir:
            missing_path = Path(tmpdir) / "missing.json"
            with self.assertRaises(FileNotFoundError):
                _load_role_map(None, missing_path, True, set(), set(), None)

        with self.assertRaises(ValueError):
            _load_role_map(["bad"], None, False, set(), set(), None)

        with self.assertRaises(ValueError):
            _load_role_map({1: "body_L1"}, None, False, set(), set(), None)

        with self.assertRaises(ValueError):
            _load_role_map({"Normal": "bad"}, None, False, set(), {"normal"}, None)

    def test_load_role_map_id_and_warning(self) -> None:
        log_state = ParseLogState(template_path=_fixture("TPL_BASIC.docx"), start_time=datetime.now())
        name_map, id_map = _load_role_map(
            {"id:Heading1": "body_L1", "Heading1": "body_L1"},
            None,
            False,
            {"Heading1"},
            {"heading 1"},
            log_state,
        )
        self.assertEqual(id_map.get("heading1"), "body_L1")
        self.assertTrue(log_state.warnings)
        self.assertNotIn("heading1", name_map)

    def test_load_role_map_case_insensitive(self) -> None:
        name_map, id_map = _load_role_map(
            {"HeAdInG 1": "Body", "id:Heading1": "title_L1"},
            None,
            False,
            {"Heading1"},
            {"heading 1"},
            None,
        )
        self.assertEqual(name_map.get("heading 1"), "body_L1")
        self.assertEqual(id_map.get("heading1"), "title_L1")

    def test_load_role_map_accepts_multilevel_and_special_roles(self) -> None:
        name_map, id_map = _load_role_map(
            {
                "Heading 2": "title_L2",
                "Body Style": "body_L3",
                "摘要": "abstract_title",
                "参考文献": "reference_body",
                "Figure Caption": "figure_caption",
            },
            None,
            False,
            {"Heading2", "BodyStyle"},
            {"heading 2", "body style", "摘要", "参考文献", "figure caption"},
            None,
        )
        self.assertEqual(name_map.get("heading 2"), "title_L2")
        self.assertEqual(name_map.get("body style"), "body_L3")
        self.assertEqual(name_map.get("摘要"), "abstract_title")
        self.assertEqual(name_map.get("参考文献"), "reference_body")
        self.assertEqual(name_map.get("figure caption"), "figure_caption")
        self.assertEqual(id_map, {})

    def test_load_role_map_invalid_multilevel_role(self) -> None:
        with self.assertRaises(ValueError):
            _load_role_map({"Heading 0": "title_L0"}, None, False, set(), set(), None)

    def test_detect_heading_levels_outline_normalization(self) -> None:
        styles = {
            "H1": _make_style_def("H1", "Heading 1", outline_level=0),
            "H2": _make_style_def("H2", "Heading 2", outline_level=1),
            "Normal": _make_style_def("Normal", "Normal", outline_level=None),
        }
        levels, overflow, min_outline_level = _detect_heading_levels(
            styles,
            samples={},
            name_map={},
            id_map={},
            max_heading_level=6,
        )
        self.assertEqual(levels, [1, 2])
        self.assertEqual(overflow, [])
        self.assertEqual(min_outline_level, 0)

    def test_detect_heading_levels_explicit_overflow_and_name(self) -> None:
        styles = {
            "H1": _make_style_def("H1", "Heading 1", outline_level=None),
            "H3": _make_style_def("H3", "Heading 3", outline_level=None),
        }
        samples = {"H2": SampleStats(count=1, style_name="Heading 2")}
        levels, overflow, min_outline_level = _detect_heading_levels(
            styles,
            samples=samples,
            name_map={"heading 4": "title_L4"},
            id_map={"h7": "title_L7"},
            max_heading_level=3,
        )
        self.assertEqual(levels, [1, 2, 3])
        self.assertEqual(overflow, [4, 7])
        self.assertIsNone(min_outline_level)

    def test_collect_body_candidates_outline_normalization(self) -> None:
        paragraphs = [
            _make_paragraph(
                "Title1",
                style_name="Title One",
                ppr=_make_ppr(outline=0),
                text="Title One",
            ),
            _make_paragraph(
                "Body1",
                style_name="Body One",
                text="Body One",
            ),
            _make_paragraph(
                "Title2",
                style_name="Title Two",
                ppr=_make_ppr(outline=1),
                text="Title Two",
            ),
            _make_paragraph(
                "Body2",
                style_name="Body Two",
                text="Body Two",
            ),
        ]
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            temp_path = Path(tmp.name)
        try:
            Document().save(temp_path)
            with mock.patch("src.template_parser._iter_paragraphs", return_value=paragraphs):
                candidates = _collect_body_candidates_by_stack(
                    temp_path,
                    name_map={},
                    id_map={},
                    max_heading_level=None,
                    outline_level_max=3,
                    outline_level_min=0,
                )
        finally:
            temp_path.unlink()
        self.assertIn("body_L1", candidates)
        self.assertIn("body_L2", candidates)

    def test_collect_body_candidates_explicit_body_priority(self) -> None:
        paragraphs = [
            _make_paragraph(
                "Heading1",
                style_name="Heading 1",
                ppr=_make_ppr(outline=0),
                text="Body Content",
            )
        ]
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            temp_path = Path(tmp.name)
        try:
            Document().save(temp_path)
            with mock.patch("src.template_parser._iter_paragraphs", return_value=paragraphs):
                candidates = _collect_body_candidates_by_stack(
                    temp_path,
                    name_map={"heading 1": "body_L2"},
                    id_map={},
                    max_heading_level=None,
                    outline_level_max=3,
                    outline_level_min=0,
                )
        finally:
            temp_path.unlink()
        self.assertIn("body_L2", candidates)
        self.assertNotIn("body_L1", candidates)
        self.assertEqual(candidates["body_L2"]["Heading1"]["source"], "explicit")

    def test_collect_body_candidates_reference_mode_across_blank(self) -> None:
        paragraphs = [
            _make_paragraph(
                "RefTitle",
                style_name="ReferenceTitle",
                text="参考文献",
            ),
            _make_paragraph(
                "Normal",
                style_name="Normal",
                text="[1] 引用",
            ),
            _make_paragraph(
                "Normal",
                style_name="Normal",
                text="",
            ),
            _make_paragraph(
                "Normal",
                style_name="Normal",
                text="[2] 引用",
            ),
            _make_paragraph(
                "Heading1",
                style_name="Heading 1",
                text="第三章",
            ),
            _make_paragraph(
                "Normal",
                style_name="Normal",
                text="正文",
            ),
        ]
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            temp_path = Path(tmp.name)
        try:
            Document().save(temp_path)
            with mock.patch("src.template_parser._iter_paragraphs", return_value=paragraphs):
                candidates = _collect_body_candidates_by_stack(
                    temp_path,
                    name_map={},
                    id_map={},
                    max_heading_level=6,
                    outline_level_max=None,
                    outline_level_min=None,
                )
        finally:
            temp_path.unlink()
        self.assertIn("reference_body", candidates)
        self.assertEqual(candidates["reference_body"]["Normal"]["count"], 2)

    def test_collect_body_candidates_abstract_blank_ends(self) -> None:
        paragraphs = [
            _make_paragraph(
                "AbstractTitle",
                style_name="AbstractTitle",
                text="摘要",
            ),
            _make_paragraph(
                "Normal",
                style_name="Normal",
                text="摘要内容",
            ),
            _make_paragraph(
                "Normal",
                style_name="Normal",
                text="",
            ),
            _make_paragraph(
                "Normal",
                style_name="Normal",
                text="摘要后正文",
            ),
        ]
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            temp_path = Path(tmp.name)
        try:
            Document().save(temp_path)
            with mock.patch("src.template_parser._iter_paragraphs", return_value=paragraphs):
                candidates = _collect_body_candidates_by_stack(
                    temp_path,
                    name_map={},
                    id_map={},
                    max_heading_level=6,
                    outline_level_max=None,
                    outline_level_min=None,
                )
        finally:
            temp_path.unlink()
        self.assertIn("abstract_body", candidates)
        self.assertEqual(candidates["abstract_body"]["Normal"]["count"], 1)
    def test_parse_heading_level_from_name_rules(self) -> None:
        self.assertEqual(_parse_heading_level_from_name("Heading 2"), 2)
        self.assertEqual(_parse_heading_level_from_name("Title3"), 3)
        self.assertEqual(_parse_heading_level_from_name("Chapter 4"), 4)
        self.assertEqual(_parse_heading_level_from_name("Chapter"), 1)
        self.assertEqual(_parse_heading_level_from_name("章节标题"), 1)
        self.assertEqual(_parse_heading_level_from_name("标题二"), 2)
        self.assertEqual(_parse_heading_level_from_name("二级标题"), 2)
        self.assertEqual(_parse_heading_level_from_name("一级标题"), 1)

    def test_match_role_paths(self) -> None:
        resolved = types.SimpleNamespace(style_id="Heading1", name="Heading 1", outline_level=None)
        stats = SampleStats(count=1, first_index=0, outline_min=0)
        self.assertEqual(
            _match_role(resolved, {}, {"heading1": "body_L1"}, 0, 0, 6, stats),
            ("body_L1", "explicit"),
        )
        self.assertEqual(
            _match_role(resolved, {"heading 1": "title_L1"}, {}, 0, 0, 6, stats),
            ("title_L1", "explicit"),
        )
        self.assertEqual(
            _match_role(types.SimpleNamespace(style_id=None, name=None, outline_level=None), {}, {}, 0, 0, 6, stats),
            ("title_L1", "outline"),
        )
        self.assertEqual(
            _match_role(types.SimpleNamespace(style_id=None, name="正文样式", outline_level=None), {}, {}, None, None, 6, None),
            ("body_L1", "keyword"),
        )
        self.assertIsNone(
            _match_role(types.SimpleNamespace(style_id=None, name=None, outline_level=None), {}, {}, None, None, 6, None)
        )
        self.assertEqual(
            _match_role(types.SimpleNamespace(style_id=None, name="Heading 2", outline_level=None), {}, {}, None, None, 6, stats),
            ("title_L2", "keyword"),
        )
        self.assertIsNone(
            _match_role(types.SimpleNamespace(style_id=None, name="Heading 8", outline_level=None), {}, {}, None, None, 6, stats)
        )
        text_stats = SampleStats(count=1, text_samples=["摘要"])
        self.assertEqual(
            _match_role(types.SimpleNamespace(style_id=None, name="Custom", outline_level=None), {}, {}, None, None, 6, text_stats),
            ("abstract_title", "text"),
        )

    def test_match_role_outline_level_max(self) -> None:
        resolved = types.SimpleNamespace(style_id=None, name=None, outline_level=1)
        self.assertIsNone(_match_role(resolved, {}, {}, 0, 1, 6, None))
        self.assertEqual(_match_role(resolved, {}, {}, 1, 1, 6, None), ("title_L1", "outline"))

    def test_match_role_keywords(self) -> None:
        self.assertEqual(
            _match_role(types.SimpleNamespace(style_id=None, name="Title 1", outline_level=None), {}, {}, None, None, 6, None),
            ("title_L1", "keyword"),
        )
        self.assertEqual(
            _match_role(types.SimpleNamespace(style_id=None, name="Normal", outline_level=None), {}, {}, None, None, 6, None),
            ("body_L1", "keyword"),
        )
        self.assertEqual(_match_special_role_by_style_name("摘要标题"), "abstract_title")
        self.assertEqual(_match_special_role_by_style_name("References Title"), "reference_title")
        self.assertEqual(_match_special_role_by_style_name("Figure Caption"), "figure_caption")
        self.assertEqual(_match_special_role_by_style_name("Table Caption"), "table_caption")
        self.assertIsNone(_match_special_role_by_style_name("Caption"))
        self.assertIsNone(_match_special_role_by_style_name("图表目录"))

    def test_choose_best_candidate_order(self) -> None:
        resolved_a = _make_resolved(
            "A",
            "A",
            font_name="Arial",
            font_size_pt=12.0,
            bold=True,
            alignment="LEFT",
            space_before_pt=0.0,
            space_after_pt=0.0,
            line_rule="exact",
            line_twips=240,
        )
        resolved_b = _make_resolved("B", "B")
        stats_a = SampleStats(count=1, first_index=2)
        stats_b = SampleStats(count=5, first_index=1)
        candidates = [
            {"resolved": resolved_b, "source": "keyword", "stats": stats_b, "order": 2},
            {"resolved": resolved_a, "source": "explicit", "stats": stats_a, "order": 1},
        ]
        selected = _choose_best_candidate("body_L1", candidates)
        self.assertIsNotNone(selected)
        self.assertEqual(selected[0].font_name, "Arial")
        self.assertIsNone(_choose_best_candidate("body_L1", []))

    def test_choose_best_candidate_prefers_count(self) -> None:
        resolved_a = _make_resolved("A", "A")
        resolved_b = _make_resolved(
            "B",
            "B",
            font_name="Complete",
            font_size_pt=12.0,
            bold=True,
            alignment="LEFT",
            space_before_pt=0.0,
            space_after_pt=0.0,
            line_rule="exact",
            line_twips=240,
        )
        stats_a = SampleStats(count=5, first_index=2)
        stats_b = SampleStats(count=1, first_index=1)
        candidates = [
            {"resolved": resolved_a, "source": "keyword", "stats": stats_a, "order": 2},
            {"resolved": resolved_b, "source": "keyword", "stats": stats_b, "order": 1},
        ]
        selected = _choose_best_candidate("body_L1", candidates)
        self.assertIsNotNone(selected)
        self.assertEqual(selected[0].font_name, resolved_a.font_name)

    def test_resolve_role_conflicts_prefers_next_candidate(self) -> None:
        resolved_a = _make_resolved("A", "A")
        resolved_b = _make_resolved("B", "B")
        resolved_c = _make_resolved("C", "C")
        stats_a = SampleStats(count=5, first_index=1)
        stats_b = SampleStats(count=1, first_index=2)
        stats_c = SampleStats(count=2, first_index=3)
        role_candidates = {
            "title_L1": [
                {"resolved": resolved_a, "source": "explicit", "stats": stats_a, "order": 1},
                {"resolved": resolved_b, "source": "keyword", "stats": stats_b, "order": 2},
            ],
            "title_L2": [
                {"resolved": resolved_a, "source": "keyword", "stats": stats_a, "order": 1},
                {"resolved": resolved_c, "source": "keyword", "stats": stats_c, "order": 2},
            ],
        }
        log_state = ParseLogState(template_path=_fixture("TPL_BASIC.docx"), start_time=datetime.now())
        selected = _resolve_role_conflicts(role_candidates, log_state=log_state)
        self.assertEqual(selected["title_L1"]["resolved"].style_id, "A")
        self.assertEqual(selected["title_L2"]["resolved"].style_id, "C")
        self.assertFalse(any(warning.rule == "shared_style" for warning in log_state.warnings))

    def test_resolve_role_conflicts_allows_shared_style(self) -> None:
        resolved_a = _make_resolved("A", "A")
        stats = SampleStats(count=1, first_index=1)
        role_candidates = {
            "title_L1": [{"resolved": resolved_a, "source": "explicit", "stats": stats, "order": 1}],
            "title_L2": [{"resolved": resolved_a, "source": "keyword", "stats": stats, "order": 2}],
        }
        log_state = ParseLogState(template_path=_fixture("TPL_BASIC.docx"), start_time=datetime.now())
        selected = _resolve_role_conflicts(role_candidates, log_state=log_state)
        self.assertEqual(selected["title_L1"]["resolved"].style_id, "A")
        self.assertEqual(selected["title_L2"]["resolved"].style_id, "A")
        self.assertTrue(any(warning.rule == "shared_style" for warning in log_state.warnings))

    def test_build_style_rule_uses_samples(self) -> None:
        resolved = _make_resolved("S1", "S1")
        stats = SampleStats(count=1, first_index=0)
        stats.weights["line_spacing"] = {("SINGLE", 1.0, "MULTIPLE"): (1.0, 1)}
        stats.weights["font_name"] = {"SampleFont": (1.0, 1)}
        rule = _build_style_rule("body_L1", resolved, stats)
        self.assertEqual(rule.line_spacing_rule, "SINGLE")
        self.assertEqual(rule.font_name, "SampleFont")

    def test_build_style_rule_prefers_east_asia_font(self) -> None:
        resolved = _make_resolved("S1", "S1", font_name="AsciiFont")
        stats = SampleStats(count=1, first_index=0)
        stats.weights["font_name_eastAsia"] = {"EastFont": (1.0, 1)}
        rule = _build_style_rule("body_L1", resolved, stats)
        self.assertEqual(rule.font_name_eastAsia, "EastFont")
        self.assertEqual(rule.font_name_ascii, "AsciiFont")
        self.assertEqual(rule.font_name, "EastFont")

    def test_build_style_rule_keeps_resolved_fields(self) -> None:
        resolved = _make_resolved(
            "S1",
            "S1",
            font_name="ResolvedFont",
            font_size_pt=12.0,
            bold=True,
            alignment="RIGHT",
            space_before_pt=6.0,
            space_after_pt=4.0,
            line_rule="exact",
            line_twips=240,
        )
        stats = SampleStats(count=3, first_index=0)
        stats.weights["alignment"] = {"LEFT": (3.0, 1)}
        stats.weights["font_size_pt"] = {10.0: (3.0, 1)}
        stats.weights["bold"] = {False: (3.0, 1)}
        stats.weights["line_spacing"] = {("MULTIPLE", 2.0, "MULTIPLE"): (3.0, 1)}
        rule = _build_style_rule("body_L1", resolved, stats)
        self.assertEqual(rule.alignment, "RIGHT")
        self.assertEqual(rule.font_size_pt, 12.0)
        self.assertTrue(rule.bold)
        self.assertEqual(rule.line_spacing_rule, "EXACTLY")
        self.assertEqual(rule.line_spacing_value, 12.0)
        self.assertEqual(rule.line_spacing_unit, "PT")

    def test_resolve_line_spacing(self) -> None:
        resolved = types.SimpleNamespace(line_rule="exact", line_twips=240)
        self.assertEqual(_resolve_line_spacing(resolved), ("EXACTLY", 12.0, "PT"))
        resolved = types.SimpleNamespace(line_rule="atLeast", line_twips=400)
        self.assertEqual(_resolve_line_spacing(resolved), ("AT_LEAST", 20.0, "PT"))
        resolved = types.SimpleNamespace(line_rule="auto", line_twips=360)
        self.assertEqual(_resolve_line_spacing(resolved), ("MULTIPLE", 1.5, "MULTIPLE"))
        resolved = types.SimpleNamespace(line_rule=None, line_twips=240)
        self.assertEqual(_resolve_line_spacing(resolved), ("MULTIPLE", 1.0, "MULTIPLE"))

    def test_collect_paragraph_samples_import_error(self) -> None:
        with mock.patch("builtins.__import__") as mocked_import:
            def _raise(name, *args, **kwargs):
                if name == "docx":
                    raise ImportError("missing")
                return __import__(name, *args, **kwargs)

            mocked_import.side_effect = _raise
            with self.assertRaises(ImportError):
                _collect_paragraph_samples(_fixture("TPL_BASIC.docx"), None)

    def test_collect_paragraph_samples_custom(self) -> None:
        paragraphs = [
            _make_paragraph(None),
            _make_paragraph(""),
            _make_paragraph(
                "Style1",
                alignment=None,
                runs=[
                    _make_run(" "),
                    _make_run("text", font_name="RunFont"),
                ],
            ),
        ]
        with mock.patch("src.template_parser._iter_paragraphs", return_value=paragraphs):
            stats = _collect_paragraph_samples(_fixture("TPL_BASIC.docx"), None)
        self.assertIn("Style1", stats)

    def test_collect_paragraph_samples_records_style_name(self) -> None:
        paragraphs = [
            _make_paragraph(
                "Style1",
                style_name="Style One",
                runs=[_make_run("text", font_name="RunFont")],
            )
        ]
        with mock.patch("src.template_parser._iter_paragraphs", return_value=paragraphs):
            stats = _collect_paragraph_samples(_fixture("TPL_BASIC.docx"), None)
        self.assertEqual(stats["Style1"].style_name, "Style One")

    def test_collect_paragraph_samples_groups_by_style_name(self) -> None:
        paragraphs = [
            _make_paragraph(
                "Style1",
                style_name="SameStyle",
                runs=[_make_run("text", font_name="FontA")],
            ),
            _make_paragraph(
                "Style2",
                style_name="SameStyle",
                runs=[_make_run("text", font_name="FontB")],
            ),
        ]
        with mock.patch("src.template_parser._iter_paragraphs", return_value=paragraphs):
            stats = _collect_paragraph_samples(_fixture("TPL_BASIC.docx"), None)
        self.assertIn("Style1", stats)
        self.assertIn("Style2", stats)
        self.assertIs(stats["Style1"], stats["Style2"])
        self.assertEqual(stats["Style1"].count, 2)

    def test_collect_paragraph_samples_tables(self) -> None:
        doc = Document()
        doc.add_paragraph("Heading outside table", style="Heading 1")
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.paragraphs[0].text = "Table paragraph"
        nested = cell.add_table(rows=1, cols=1)
        nested.cell(0, 0).paragraphs[0].text = "Nested table paragraph"
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            temp_path = Path(tmp.name)
        try:
            doc.save(temp_path)
            stats = _collect_paragraph_samples(temp_path, None)
        finally:
            temp_path.unlink()
        normal_key = next((key for key in stats if key.lower() == "normal"), None)
        self.assertIsNotNone(normal_key)
        self.assertGreaterEqual(stats[normal_key].count, 2)

    def test_add_sample_and_sample_mode(self) -> None:
        stats = SampleStats()
        index = _add_sample(stats, "field", "A", 1.0, 0)
        index = _add_sample(stats, "field", "A", 1.0, index)
        self.assertEqual(index, 1)
        self.assertEqual(_sample_mode(stats, "field"), "A")
        self.assertIsNone(_sample_mode(stats, "missing"))

    def test_extract_paragraph_alignment(self) -> None:
        self.assertIsNone(_extract_paragraph_alignment(_make_paragraph("S"), WD_ALIGN_PARAGRAPH))
        self.assertEqual(
            _extract_paragraph_alignment(
                _make_paragraph("S", alignment=WD_ALIGN_PARAGRAPH.LEFT),
                WD_ALIGN_PARAGRAPH,
            ),
            "LEFT",
        )
        self.assertEqual(
            _extract_paragraph_alignment(
                _make_paragraph("S", alignment=WD_ALIGN_PARAGRAPH.CENTER),
                WD_ALIGN_PARAGRAPH,
            ),
            "CENTER",
        )
        self.assertEqual(
            _extract_paragraph_alignment(
                _make_paragraph("S", alignment=WD_ALIGN_PARAGRAPH.RIGHT),
                WD_ALIGN_PARAGRAPH,
            ),
            "RIGHT",
        )
        self.assertEqual(
            _extract_paragraph_alignment(
                _make_paragraph("S", alignment=WD_ALIGN_PARAGRAPH.JUSTIFY),
                WD_ALIGN_PARAGRAPH,
            ),
            "JUSTIFY",
        )
        self.assertEqual(
            _extract_paragraph_alignment(
                _make_paragraph("S", alignment=WD_ALIGN_PARAGRAPH.DISTRIBUTE),
                WD_ALIGN_PARAGRAPH,
            ),
            "JUSTIFY",
        )
        self.assertEqual(
            _extract_paragraph_alignment(
                _make_paragraph("S", alignment=None, style_alignment=WD_ALIGN_PARAGRAPH.CENTER),
                WD_ALIGN_PARAGRAPH,
            ),
            "CENTER",
        )
        self.assertEqual(
            _extract_paragraph_alignment(
                _make_paragraph("S", alignment="UNKNOWN"),
                WD_ALIGN_PARAGRAPH,
            ),
            "JUSTIFY",
        )

    def test_extract_paragraph_line_spacing_variants(self) -> None:
        paragraph = _make_paragraph("S", line_rule=WD_LINE_SPACING.SINGLE)
        self.assertEqual(
            _extract_paragraph_line_spacing(paragraph, WD_LINE_SPACING, None, "S", 1),
            ("SINGLE", 1.0, "MULTIPLE"),
        )
        paragraph = _make_paragraph("S", line_rule=WD_LINE_SPACING.ONE_POINT_FIVE)
        self.assertEqual(
            _extract_paragraph_line_spacing(paragraph, WD_LINE_SPACING, None, "S", 1),
            ("ONE_POINT_FIVE", 1.5, "MULTIPLE"),
        )
        paragraph = _make_paragraph("S", line_rule=WD_LINE_SPACING.DOUBLE)
        self.assertEqual(
            _extract_paragraph_line_spacing(paragraph, WD_LINE_SPACING, None, "S", 1),
            ("DOUBLE", 2.0, "MULTIPLE"),
        )
        paragraph = _make_paragraph("S", line_rule=WD_LINE_SPACING.MULTIPLE, line_value=1.2)
        self.assertEqual(
            _extract_paragraph_line_spacing(paragraph, WD_LINE_SPACING, None, "S", 1),
            ("MULTIPLE", 1.2, "MULTIPLE"),
        )
        paragraph = _make_paragraph("S", line_rule=WD_LINE_SPACING.EXACTLY, line_value=Pt(18))
        self.assertEqual(
            _extract_paragraph_line_spacing(paragraph, WD_LINE_SPACING, None, "S", 1),
            ("EXACTLY", 18.0, "PT"),
        )
        paragraph = _make_paragraph("S", line_rule=WD_LINE_SPACING.AT_LEAST, line_value=Pt(20))
        self.assertEqual(
            _extract_paragraph_line_spacing(paragraph, WD_LINE_SPACING, None, "S", 1),
            ("AT_LEAST", 20.0, "PT"),
        )
        paragraph = _make_paragraph("S", line_rule=None, line_value=1.1)
        self.assertEqual(
            _extract_paragraph_line_spacing(paragraph, WD_LINE_SPACING, None, "S", 1),
            ("MULTIPLE", 1.1, "MULTIPLE"),
        )
        paragraph = _make_paragraph(
            "S",
            line_rule=None,
            line_value=Pt(16),
            ppr=_make_ppr(line_rule="exact", line=360),
        )
        self.assertEqual(
            _extract_paragraph_line_spacing(paragraph, WD_LINE_SPACING, None, "S", 1),
            ("EXACTLY", 18.0, "PT"),
        )
        paragraph = _make_paragraph(
            "S",
            line_rule=None,
            line_value=Pt(16),
            ppr=_make_ppr(line_rule="auto", line=480),
        )
        self.assertEqual(
            _extract_paragraph_line_spacing(paragraph, WD_LINE_SPACING, None, "S", 1),
            ("MULTIPLE", 2.0, "MULTIPLE"),
        )
        log_state = ParseLogState(template_path=_fixture("TPL_BASIC.docx"), start_time=datetime.now())
        paragraph = _make_paragraph("S", line_rule=None, line_value=Pt(16))
        spacing = _extract_paragraph_line_spacing(paragraph, WD_LINE_SPACING, log_state, "S", 1)
        self.assertEqual(spacing[0], "EXACTLY")
        self.assertTrue(log_state.warnings)

    def test_length_to_pt_and_parse_int(self) -> None:
        class DummyLength:
            pt = 12.5

        self.assertEqual(_length_to_pt(DummyLength()), 12.5)
        self.assertIsNone(_length_to_pt({}))
        self.assertIsNone(_parse_int("bad"))

    def test_read_paragraph_line_rule(self) -> None:
        paragraph = types.SimpleNamespace(_element=types.SimpleNamespace(pPr=None))
        self.assertEqual(_read_paragraph_line_rule(paragraph), (None, None))
        paragraph = types.SimpleNamespace(_element=_make_ppr())
        self.assertEqual(_read_paragraph_line_rule(paragraph), (None, None))
        paragraph = types.SimpleNamespace(_element=_make_ppr(line_rule="exact", line=240))
        self.assertEqual(_read_paragraph_line_rule(paragraph), ("exact", 240))

    def test_extract_run_helpers(self) -> None:
        run = _make_run("text", font_name="Fallback")
        self.assertEqual(_extract_run_fonts(run)[0], "Fallback")
        run = _make_run(
            "text",
            font_name="Fallback",
            ascii_name="Ascii",
            hansi_name="HAnsi",
            east_asia_name="East",
        )
        font_name, ascii_name, hansi_name, east_asia_name = _extract_run_fonts(run)
        self.assertEqual(font_name, "East")
        self.assertEqual(ascii_name, "Ascii")
        self.assertEqual(hansi_name, "HAnsi")
        self.assertEqual(east_asia_name, "East")

        run = _make_run("text", size=Pt(11))
        self.assertEqual(_extract_run_font_size(run), 11.0)
        run = _make_run("text")
        self.assertIsNone(_extract_run_font_size(run))

        run = _make_run("text", bold=True)
        self.assertTrue(_extract_run_bold(run))
        run = _make_run("text", bold=None)
        run.font.bold = False
        self.assertFalse(_extract_run_bold(run))

    def test_extract_paragraph_outline_level(self) -> None:
        paragraph = types.SimpleNamespace(_element=_make_ppr(outline=1))
        self.assertEqual(_extract_paragraph_outline_level(paragraph), 1)
        paragraph = types.SimpleNamespace(_element=types.SimpleNamespace(pPr=None))
        self.assertIsNone(_extract_paragraph_outline_level(paragraph))

    def test_apply_fallbacks_and_strict_validation(self) -> None:
        log_state = ParseLogState(template_path=_fixture("TPL_BASIC.docx"), start_time=datetime.now())
        rules = _apply_fallbacks(
            {},
            allow_fallback=False,
            strict=False,
            log_state=log_state,
            required_on_presence_map=None,
            global_body_rule=None,
        )
        self.assertIn("body_L1", rules)
        self.assertTrue(log_state.warnings)
        rules = _apply_fallbacks(
            {},
            allow_fallback=True,
            strict=False,
            log_state=None,
            required_on_presence_map=None,
            global_body_rule=None,
        )
        self.assertEqual(rules["body_L1"].font_name, "宋体")
        with self.assertRaises(ValueError):
            _validate_strict({"body_L1": StyleRule(role="body_L1")})
        with self.assertRaises(ValueError):
            _ensure_required_roles({"body_L1": StyleRule(role="body_L1")})
        _ensure_required_roles(
            {
                "body_L1": StyleRule(role="body_L1"),
                "title_L1": StyleRule(role="title_L1"),
            }
        )

    def test_apply_fallbacks_default_rules(self) -> None:
        rules = _apply_fallbacks(
            {},
            allow_fallback=True,
            strict=False,
            log_state=None,
            required_on_presence_map=None,
            global_body_rule=None,
        )
        chapter = rules["title_L1"]
        body = rules["body_L1"]
        self.assertEqual(chapter.font_name, "宋体")
        self.assertEqual(chapter.font_size_pt, 16.0)
        self.assertTrue(chapter.bold)
        self.assertEqual(chapter.alignment, "CENTER")
        self.assertEqual(chapter.line_spacing_rule, "ONE_POINT_FIVE")
        self.assertEqual(chapter.line_spacing_value, 1.5)
        self.assertEqual(chapter.line_spacing_unit, "MULTIPLE")
        self.assertEqual(chapter.space_before_pt, 12.0)
        self.assertEqual(chapter.space_after_pt, 12.0)
        self.assertEqual(body.font_name, "宋体")
        self.assertEqual(body.font_size_pt, 12.0)
        self.assertFalse(body.bold)
        self.assertEqual(body.alignment, "JUSTIFY")
        self.assertEqual(body.line_spacing_rule, "ONE_POINT_FIVE")
        self.assertEqual(body.line_spacing_value, 1.5)
        self.assertEqual(body.line_spacing_unit, "MULTIPLE")
        self.assertEqual(body.space_before_pt, 0.0)
        self.assertEqual(body.space_after_pt, 0.0)

    def test_apply_fallbacks_preserves_existing_fields(self) -> None:
        rules = {
            "body_L1": StyleRule(
                role="body_L1",
                font_name="自定义字体",
                font_size_pt=10.0,
            )
        }
        rules = _apply_fallbacks(
            rules,
            allow_fallback=True,
            strict=False,
            log_state=None,
            required_on_presence_map=None,
            global_body_rule=None,
        )
        body = rules["body_L1"]
        self.assertEqual(body.font_name, "自定义字体")
        self.assertEqual(body.font_size_pt, 10.0)
        self.assertEqual(body.bold, False)

    def test_apply_fallbacks_conditional_required(self) -> None:
        body_rule = StyleRule(
            role="body_L1",
            font_name="BodyFont",
            font_size_pt=12.0,
            bold=False,
            alignment="LEFT",
            line_spacing_rule="SINGLE",
            line_spacing_value=1.0,
            line_spacing_unit="MULTIPLE",
            space_before_pt=0.0,
            space_after_pt=0.0,
        )
        rules = {
            "abstract_title": StyleRule(role="abstract_title"),
            "body_L1": body_rule,
        }
        rules = _apply_fallbacks(
            rules,
            allow_fallback=True,
            strict=False,
            log_state=None,
            required_on_presence_map=None,
            global_body_rule=None,
        )
        self.assertIn("abstract_body", rules)
        self.assertEqual(rules["abstract_body"].font_name, "BodyFont")

    def test_apply_fallbacks_conditional_missing_no_fallback(self) -> None:
        rules = {"abstract_title": StyleRule(role="abstract_title")}
        rules = _apply_fallbacks(
            rules,
            allow_fallback=False,
            strict=False,
            log_state=None,
            required_on_presence_map=None,
            global_body_rule=None,
        )
        self.assertIn("abstract_body", rules)
        self.assertIsNone(rules["abstract_body"].font_name)

    def test_apply_fallbacks_body_level_uses_global(self) -> None:
        global_rule = StyleRule(
            role="body_L1",
            font_name="GlobalBody",
            font_size_pt=10.0,
            bold=False,
            alignment="LEFT",
            line_spacing_rule="SINGLE",
            line_spacing_value=1.0,
            line_spacing_unit="MULTIPLE",
            space_before_pt=0.0,
            space_after_pt=0.0,
        )
        rules = _apply_fallbacks(
            {},
            allow_fallback=True,
            strict=False,
            log_state=None,
            required_roles=["body_L2"],
            required_on_presence_map=None,
            global_body_rule=global_rule,
        )
        self.assertEqual(rules["body_L2"].font_name, "GlobalBody")

    def test_validate_strict_checks_optional_roles(self) -> None:
        with self.assertRaises(ValueError):
            _validate_strict({"abstract_body": StyleRule(role="abstract_body")})
    def test_log_helpers(self) -> None:
        log_state = ParseLogState(template_path=_fixture("TPL_BASIC.docx"), start_time=datetime.now())
        log_state.elapsed_sec = 0.123
        log_state.style_count = 3
        log_state.detected_heading_levels = [1, 2]
        log_state.detected_heading_levels_overflow = [7]
        log_state.role_sources["body_L1"] = "source=keyword style_name=正文 style_id=Style1"
        _warn(
            log_state,
            rule="missing_fields",
            reason="missing",
            style_id="Style1",
            paragraph_index=2,
        )
        _tag_warnings_for_style(log_state, "body_L1", "Style1")
        _write_log(log_state)
        logs = sorted(config.LOG_DIR.glob("template_parser_*.log"))
        self.assertTrue(logs)
        content = logs[-1].read_text(encoding="utf-8")
        self.assertIn("template_path:", content)
        self.assertIn("elapsed_sec:", content)
        self.assertIn("styles_count: 3", content)
        self.assertIn("detected_heading_levels: [1, 2]", content)
        self.assertIn("detected_heading_levels_overflow: [7]", content)
        self.assertIn("role_source[body_L1]: source=keyword", content)
        self.assertIn("warnings_count", content)
        self.assertIn("role=body_L1", content)
        self.assertIn("paragraph_index=2", content)

    def test_fallback_warning_includes_chain(self) -> None:
        log_state = ParseLogState(template_path=_fixture("TPL_BASIC.docx"), start_time=datetime.now())
        rules = _apply_fallbacks(
            {},
            allow_fallback=True,
            strict=False,
            log_state=log_state,
            required_roles=["title_L2"],
            required_on_presence_map=None,
            global_body_rule=None,
        )
        self.assertIn("title_L2", rules)
        reasons = [warning.reason for warning in log_state.warnings]
        self.assertTrue(any("via title_L1" in reason for reason in reasons))

    def test_describe_candidate(self) -> None:
        resolved = _make_resolved("Style1", "Name")
        desc = _describe_candidate({"resolved": resolved, "source": "keyword"})
        self.assertIn("style_id=Style1", desc)

    def test_parse_strict_error_logged(self) -> None:
        parser = TemplateParser()
        with self.assertRaises(ValueError):
            parser.parse(str(_fixture("TPL_INVALID.docx")))
        logs = sorted(config.LOG_DIR.glob("template_parser_*.log"))
        self.assertTrue(logs)
        content = logs[-1].read_text(encoding="utf-8")
        self.assertIn("error:", content)


if __name__ == "__main__":
    unittest.main()
