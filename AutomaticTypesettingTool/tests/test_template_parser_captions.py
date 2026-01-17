import base64
import tempfile
import unittest
from pathlib import Path

from src.template_parser import TemplateParser


_PNG_BASE64 = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR4nGNgYAAAAAMA"
    "ASsJTYQAAAAASUVORK5CYII="
)


class TemplateParserCaptionTests(unittest.TestCase):
    def _build_doc(self, root: Path, builder) -> Path:
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover - dependency required
            raise unittest.SkipTest("python-docx not installed") from exc
        path = root / "captions.docx"
        doc = Document()
        builder(doc, root)
        doc.save(str(path))
        return path

    def _write_png(self, path: Path) -> None:
        path.write_bytes(base64.b64decode(_PNG_BASE64))

    def test_table_caption_near_table(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = self._build_doc(
                Path(tmpdir),
                lambda doc, _: self._build_table_caption(doc),
            )
            parser = TemplateParser()
            result = parser.parse(str(path))
            self.assertIn("table_caption", result.roles)

    def test_figure_caption_near_image(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = self._build_doc(
                Path(tmpdir),
                lambda doc, root: self._build_figure_caption(doc, root),
            )
            parser = TemplateParser()
            result = parser.parse(str(path))
            self.assertIn("figure_caption", result.roles)

    def test_figure_caption_english(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = self._build_doc(
                Path(tmpdir),
                lambda doc, root: self._build_figure_caption_english(doc, root),
            )
            parser = TemplateParser()
            result = parser.parse(str(path))
            self.assertIn("figure_caption", result.roles)

    def test_table_caption_english(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = self._build_doc(
                Path(tmpdir),
                lambda doc, _: self._build_table_caption_english(doc),
            )
            parser = TemplateParser()
            result = parser.parse(str(path))
            self.assertIn("table_caption", result.roles)

    def test_caption_far_from_object(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = self._build_doc(
                Path(tmpdir),
                lambda doc, root: self._build_figure_caption_far(doc, root),
            )
            parser = TemplateParser()
            result = parser.parse(str(path))
            self.assertNotIn("figure_caption", result.roles)

    def test_caption_directory_ignored(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = self._build_doc(
                Path(tmpdir),
                lambda doc, _: self._build_caption_directory(doc),
            )
            parser = TemplateParser()
            result = parser.parse(str(path))
            self.assertNotIn("figure_caption", result.roles)
            self.assertNotIn("table_caption", result.roles)

    def test_figure_body_detected(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = self._build_doc(
                Path(tmpdir),
                lambda doc, root: self._build_figure_body(doc, root),
            )
            parser = TemplateParser()
            result = parser.parse(str(path))
            self.assertIn("figure_body", result.roles)

    def test_table_body_detected(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = self._build_doc(
                Path(tmpdir),
                lambda doc, _: self._build_table_body(doc),
            )
            parser = TemplateParser()
            result = parser.parse(str(path))
            self.assertIn("table_body", result.roles)

    def _build_table_caption(self, doc) -> None:
        doc.add_paragraph("表 1 表格标题")
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "cell"

    def _build_figure_caption(self, doc, root: Path) -> None:
        from docx.shared import Inches

        image_path = root / "tiny.png"
        self._write_png(image_path)
        doc.add_paragraph().add_run().add_picture(str(image_path), width=Inches(0.1))
        doc.add_paragraph("图 1 图片标题")

    def _build_figure_caption_english(self, doc, root: Path) -> None:
        from docx.shared import Inches

        image_path = root / "tiny.png"
        self._write_png(image_path)
        doc.add_paragraph().add_run().add_picture(str(image_path), width=Inches(0.1))
        doc.add_paragraph("Figure 1 Sample image")

    def _build_table_caption_english(self, doc) -> None:
        doc.add_paragraph("Table 1 Sample table")
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "cell"

    def _build_figure_caption_far(self, doc, root: Path) -> None:
        from docx.shared import Inches

        image_path = root / "tiny.png"
        self._write_png(image_path)
        doc.add_paragraph().add_run().add_picture(str(image_path), width=Inches(0.1))
        doc.add_paragraph("gap1")
        doc.add_paragraph("gap2")
        doc.add_paragraph("gap3")
        doc.add_paragraph("图 1 图片标题")

    def _build_caption_directory(self, doc) -> None:
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "cell"
        doc.add_paragraph("图表目录")

    def _build_figure_body(self, doc, root: Path) -> None:
        from docx.shared import Inches

        image_path = root / "tiny.png"
        self._write_png(image_path)
        doc.add_paragraph().add_run().add_picture(str(image_path), width=Inches(0.1))

    def _build_table_body(self, doc) -> None:
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "table body"


if __name__ == "__main__":
    unittest.main()
