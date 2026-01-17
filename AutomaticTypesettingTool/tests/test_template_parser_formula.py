import tempfile
import unittest
from pathlib import Path

from lxml import etree

from src.template_parser import TemplateParser


class TemplateParserFormulaTests(unittest.TestCase):
    def _add_math(self, paragraph) -> None:
        math = etree.Element(
            "{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath"
        )
        paragraph._p.append(math)

    def _add_ole_object(self, paragraph, prog_id: str) -> None:
        w_object = etree.Element(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object"
        )
        ole = etree.SubElement(
            w_object,
            "{urn:schemas-microsoft-com:office:office}OLEObject",
        )
        ole.set("ProgID", prog_id)
        paragraph._p.append(w_object)

    def _build_doc(self, root: Path, builder) -> Path:
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover - dependency required
            raise unittest.SkipTest("python-docx not installed") from exc
        path = root / "formula.docx"
        doc = Document()
        builder(doc)
        doc.save(str(path))
        return path

    def test_formula_block_and_number(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = self._build_doc(
                Path(tmpdir),
                lambda doc: self._build_block_formula(doc),
            )
            parser = TemplateParser()
            result = parser.parse(str(path))
            roles = result.roles
            self.assertIn("formula_block", roles)
            self.assertIn("formula_number", roles)

    def test_formula_inline(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = self._build_doc(
                Path(tmpdir),
                lambda doc: self._build_inline_formula(doc),
            )
            parser = TemplateParser()
            result = parser.parse(str(path))
            roles = result.roles
            self.assertIn("formula_inline", roles)

    def test_formula_ole_object(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = self._build_doc(
                Path(tmpdir),
                lambda doc: self._build_ole_formula(doc),
            )
            parser = TemplateParser()
            result = parser.parse(str(path))
            roles = result.roles
            self.assertIn("formula_block", roles)

    def test_non_formula_ole_ignored(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = self._build_doc(
                Path(tmpdir),
                lambda doc: self._build_ole_non_formula(doc),
            )
            parser = TemplateParser()
            result = parser.parse(str(path))
            roles = result.roles
            self.assertNotIn("formula_block", roles)
            self.assertNotIn("formula_inline", roles)
            self.assertNotIn("formula_number", roles)

    def test_formula_in_table(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = self._build_doc(
                Path(tmpdir),
                lambda doc: self._build_table_formula(doc),
            )
            parser = TemplateParser()
            result = parser.parse(str(path))
            roles = result.roles
            self.assertIn("formula_block", roles)

    def _build_block_formula(self, doc) -> None:
        paragraph = doc.add_paragraph()
        self._add_math(paragraph)
        paragraph.add_run("(1)")

    def _build_inline_formula(self, doc) -> None:
        paragraph = doc.add_paragraph("根据公式")
        self._add_math(paragraph)
        paragraph.add_run("可得")

    def _build_ole_formula(self, doc) -> None:
        paragraph = doc.add_paragraph()
        self._add_ole_object(paragraph, "Equation.3")

    def _build_ole_non_formula(self, doc) -> None:
        paragraph = doc.add_paragraph()
        self._add_ole_object(paragraph, "Excel.Sheet.12")

    def _build_table_formula(self, doc) -> None:
        table = doc.add_table(rows=1, cols=1)
        paragraph = table.cell(0, 0).paragraphs[0]
        self._add_math(paragraph)


if __name__ == "__main__":
    unittest.main()
