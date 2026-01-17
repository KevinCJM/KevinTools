import tempfile
import unittest
from pathlib import Path

from docx import Document

from src.template_parser import TemplateParser


class TemplateParserTitleSpacingTests(unittest.TestCase):
    def test_parse_title_spacing(self) -> None:
        document = Document()
        document.add_paragraph("")
        document.add_paragraph("摘要")
        document.add_paragraph("")
        document.add_paragraph("")
        document.add_paragraph("这是摘要正文")
        document.add_paragraph("")
        document.add_paragraph("目录")
        document.add_paragraph("")
        document.add_paragraph("第一章")

        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "spacing.docx"
            document.save(str(path))
            parser = TemplateParser()
            result = parser.parse(str(path))

        spacing = result.meta.get("title_spacing", {})
        abstract_spacing = spacing.get("abstract_title", {})
        toc_spacing = spacing.get("toc_title", {})
        self.assertEqual(abstract_spacing.get("before"), 1)
        self.assertEqual(abstract_spacing.get("after"), 2)
        self.assertEqual(toc_spacing.get("before"), 1)
        self.assertEqual(toc_spacing.get("after"), 1)


if __name__ == "__main__":
    unittest.main()
