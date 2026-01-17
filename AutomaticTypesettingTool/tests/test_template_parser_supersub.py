import tempfile
import unittest
from pathlib import Path

from src.template_parser import TemplateParser


class TemplateParserSuperSubTests(unittest.TestCase):
    def _build_doc(self, root: Path, builder) -> Path:
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover - dependency required
            raise unittest.SkipTest("python-docx not installed") from exc
        path = root / "supsub.docx"
        doc = Document()
        builder(doc)
        doc.save(str(path))
        return path

    def test_superscript_detected(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = self._build_doc(
                Path(tmpdir),
                lambda doc: self._build_superscript(doc),
            )
            parser = TemplateParser()
            result = parser.parse(str(path))
            self.assertIn("superscript", result.roles)
            role = result.roles["superscript"]
            self.assertEqual(role.font_name, "Times New Roman")
            self.assertEqual(role.font_size_pt, 9.0)
            self.assertTrue(role.bold)

    def test_subscript_detected(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = self._build_doc(
                Path(tmpdir),
                lambda doc: self._build_subscript(doc),
            )
            parser = TemplateParser()
            result = parser.parse(str(path))
            self.assertIn("subscript", result.roles)
            role = result.roles["subscript"]
            self.assertEqual(role.font_name, "Arial")
            self.assertEqual(role.font_size_pt, 8.0)
            self.assertFalse(role.bold)

    def test_subscript_inherits_paragraph_style(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = self._build_doc(
                Path(tmpdir),
                lambda doc: self._build_subscript_inherit(doc),
            )
            parser = TemplateParser()
            result = parser.parse(str(path))
            self.assertIn("subscript", result.roles)
            role = result.roles["subscript"]
            self.assertEqual(role.font_name, "Courier New")
            self.assertEqual(role.font_size_pt, 13.0)

    def test_supersub_with_body_role(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = self._build_doc(
                Path(tmpdir),
                lambda doc: self._build_supersub_with_body(doc),
            )
            parser = TemplateParser()
            result = parser.parse(str(path))
            self.assertIn("superscript", result.roles)
            self.assertIn("subscript", result.roles)
            self.assertIn("body_L1", result.roles)

    def _build_superscript(self, doc) -> None:
        from docx.shared import Pt

        paragraph = doc.add_paragraph("X")
        run = paragraph.add_run("2")
        run.font.superscript = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)
        run.bold = True

    def _build_subscript(self, doc) -> None:
        from docx.shared import Pt

        paragraph = doc.add_paragraph("H")
        run = paragraph.add_run("2")
        run.font.subscript = True
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.bold = False

    def _build_subscript_inherit(self, doc) -> None:
        from docx.enum.style import WD_STYLE_TYPE
        from docx.shared import Pt

        style = doc.styles.add_style("SubStyle", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Courier New"
        style.font.size = Pt(13)
        paragraph = doc.add_paragraph("H")
        paragraph.style = style
        run = paragraph.add_run("2")
        run.font.subscript = True

    def _build_supersub_with_body(self, doc) -> None:
        from docx.shared import Pt

        doc.add_paragraph("Body text for baseline style.")
        paragraph = doc.add_paragraph("X")
        run = paragraph.add_run("2")
        run.font.superscript = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)
        paragraph = doc.add_paragraph("H")
        run = paragraph.add_run("2")
        run.font.subscript = True
        run.font.name = "Arial"
        run.font.size = Pt(8)


if __name__ == "__main__":
    unittest.main()
