import tempfile
import unittest
from pathlib import Path

from src.section_rules import BodyRangeRule, SectionPosition, SectionRule
from src.template_parser import TemplateParser


class TemplateParserCoverSectionTests(unittest.TestCase):
    def _build_doc(self, root: Path) -> Path:
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt
        except ImportError as exc:  # pragma: no cover - dependency required
            raise unittest.SkipTest("python-docx not installed") from exc
        path = root / "cover.docx"
        doc = Document()
        title_para = doc.add_paragraph("封面标题")
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if title_para.runs:
            title_para.runs[0].font.size = Pt(22)
        doc.add_paragraph("专业：计算机")
        doc.add_paragraph("摘要")
        doc.save(str(path))
        return path

    def test_cover_content_keywords_detected(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = self._build_doc(Path(tmpdir))
            rule = SectionRule(
                key="cover",
                display_name="首页",
                title_keywords=(),
                content_keywords=("专业",),
                position=SectionPosition.FIRST_PAGE,
                body_range=BodyRangeRule.UNTIL_BLANK,
            )
            parser = TemplateParser(section_rules=[rule])
            result = parser.parse(str(path))
            self.assertIn("cover_info", result.roles)
            self.assertIn("cover_title", result.roles)
            self.assertNotIn("document_title", result.roles)


if __name__ == "__main__":
    unittest.main()
