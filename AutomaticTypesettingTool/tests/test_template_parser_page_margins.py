import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Inches

from src.template_parser import TemplateParser


class TemplateParserPageMarginsTests(unittest.TestCase):
    def test_parse_page_margins_sections(self) -> None:
        document = Document()
        section0 = document.sections[0]
        section0.top_margin = Inches(1)
        section0.left_margin = Inches(1)
        document.add_paragraph("原创声明")

        section1 = document.add_section(WD_SECTION.NEW_PAGE)
        section1.top_margin = Inches(1.5)
        section1.left_margin = Inches(1.2)
        document.add_paragraph("第一章 绪论")

        section2 = document.add_section(WD_SECTION.NEW_PAGE)
        section2.top_margin = Inches(2)
        section2.left_margin = Inches(1.4)
        document.add_paragraph("参考文献")

        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "margins.docx"
            document.save(str(path))
            parser = TemplateParser()
            result = parser.parse(str(path))

        page_margins = result.meta.get("page_margins", {})
        sections = page_margins.get("sections", [])
        self.assertEqual(len(sections), 3)
        logical_parts = {section.get("logical_part") for section in sections}
        self.assertIn("statement", logical_parts)
        self.assertIn("main", logical_parts)
        self.assertIn("back", logical_parts)
        summary = page_margins.get("summary", {})
        self.assertEqual(summary.get("sections"), 3)
        self.assertEqual(summary.get("distinct_margins"), 3)
        first_margins = sections[0].get("margins", {})
        self.assertAlmostEqual(first_margins.get("top", 0), 72.0, places=1)


if __name__ == "__main__":
    unittest.main()
