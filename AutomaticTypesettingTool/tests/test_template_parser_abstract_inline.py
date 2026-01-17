import tempfile
import unittest
from pathlib import Path

from docx import Document

from src.template_parser import TemplateParser


class TemplateParserAbstractInlineTests(unittest.TestCase):
    def test_parse_inline_abstract(self) -> None:
        document = Document()
        document.add_paragraph("摘要：这是中文摘要内容")
        document.add_paragraph("关键词：测试，摘要")
        document.add_paragraph("")
        document.add_paragraph("Abstract: This is the English abstract")
        document.add_paragraph("Key words: test")

        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "abstract_inline.docx"
            document.save(str(path))
            parser = TemplateParser()
            result = parser.parse(str(path))

        roles = result.roles
        self.assertIn("abstract_title", roles)
        self.assertIn("abstract_body", roles)
        self.assertIn("keyword_line", roles)
        self.assertIn("abstract_en_title", roles)
        self.assertIn("abstract_en_body", roles)


if __name__ == "__main__":
    unittest.main()
