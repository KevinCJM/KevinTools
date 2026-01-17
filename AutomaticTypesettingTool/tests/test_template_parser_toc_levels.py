import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from src.template_parser import TemplateParser


class TemplateParserTocLevelsTests(unittest.TestCase):
    def test_parse_toc_levels(self) -> None:
        document = Document()
        styles = document.styles
        existing_names = {style.name for style in styles}
        for name in ("TOC 1", "TOC 2", "TOC 3"):
            if name not in existing_names:
                styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
                existing_names.add(name)

        document.add_paragraph("目录")
        document.add_paragraph("第一章", style="TOC 1")
        document.add_paragraph("1.1", style="TOC 2")
        document.add_paragraph("1.1.1", style="TOC 3")

        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "toc.docx"
            document.save(str(path))
            parser = TemplateParser()
            result = parser.parse(str(path))

        self.assertIn("toc_body_L1", result.roles)
        self.assertIn("toc_body_L2", result.roles)
        self.assertIn("toc_body_L3", result.roles)
        toc_levels = result.meta.get("toc_levels", {})
        self.assertIn(1, toc_levels.get("levels", []))
        self.assertIn(2, toc_levels.get("levels", []))
        self.assertIn(3, toc_levels.get("levels", []))
        styles_map = toc_levels.get("styles", {})
        self.assertIn("TOC 1", styles_map.get("1", []))
        self.assertIn("TOC 2", styles_map.get("2", []))
        self.assertIn("TOC 3", styles_map.get("3", []))


if __name__ == "__main__":
    unittest.main()
