import json
import tempfile
import unittest
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches
from lxml import etree

from src import config
from src.template_parser import TemplateParser, _W_NS, ParseResult


FIXTURES_DIR = Path(__file__).resolve().parents[1] / "tests" / "fixtures"


class RegressionNewFeaturesTests(unittest.TestCase):
    def test_regression_new_features(self) -> None:
        expected_path = FIXTURES_DIR / "regression_expected.json"
        expected = json.loads(expected_path.read_text(encoding="utf-8"))

        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            toc_path = _create_toc_template(tmpdir_path)
            margin_table_path = _create_margin_table_template(tmpdir_path)
            footnote_path = _create_footnote_template(tmpdir_path)

            parser = TemplateParser()
            toc_result = parser.parse(str(toc_path))
            margin_result = parser.parse(str(margin_table_path))
            footnote_result = parser.parse(str(footnote_path))

        actual = {
            "toc_template": _summarize_toc(toc_result),
            "margin_table_template": _summarize_margin_table(margin_result),
            "footnote_template": _summarize_footnotes(footnote_result),
        }
        actual = _normalize_summary(actual)
        expected = _normalize_summary(expected)
        _write_actual_summary(actual)
        self.assertEqual(actual, expected)


def _create_toc_template(tmpdir: Path) -> Path:
    document = Document()
    styles = document.styles
    existing_names = {style.name for style in styles}
    for name in ("TOC 1", "TOC 2", "TOC 3"):
        if name not in existing_names:
            styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            existing_names.add(name)
    document.add_paragraph("目录")
    document.add_paragraph("第一章", style="TOC 1")
    document.add_paragraph("1.1", style="TOC 2")
    document.add_paragraph("1.1.1", style="TOC 3")
    path = tmpdir / "regression_toc.docx"
    document.save(str(path))
    return path


def _create_margin_table_template(tmpdir: Path) -> Path:
    document = Document()
    section0 = document.sections[0]
    section0.top_margin = Inches(1)
    section0.left_margin = Inches(1)
    document.add_paragraph("原创声明")

    section1 = document.add_section(WD_SECTION.NEW_PAGE)
    section1.top_margin = Inches(1.5)
    section1.left_margin = Inches(1.2)
    document.add_paragraph("第一章 绪论")

    section2 = document.add_section(WD_SECTION.NEW_PAGE)
    section2.top_margin = Inches(2)
    section2.left_margin = Inches(1.4)
    document.add_paragraph("参考文献")

    for _ in range(5):
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "x"

    path = tmpdir / "regression_margins_table.docx"
    document.save(str(path))
    xml_bytes = _update_table_borders(path)
    _rewrite_document_xml(path, xml_bytes)
    return path


def _create_footnote_template(tmpdir: Path) -> Path:
    document = Document()
    document.add_paragraph("正文")
    path = tmpdir / "regression_footnotes.docx"
    document.save(str(path))
    _inject_footnote_parts(path)
    return path


def _summarize_toc(result: ParseResult) -> dict[str, object]:
    roles = []
    for role in ("toc_body_L1", "toc_body_L2", "toc_body_L3", "toc_title"):
        if role in result.roles:
            roles.append(role)
    toc_levels = result.meta.get("toc_levels", {}).get("levels", [])
    return {
        "roles": sorted(roles),
        "toc_levels": toc_levels,
    }


def _summarize_margin_table(result: ParseResult) -> dict[str, object]:
    page_summary = result.meta.get("page_margins", {}).get("summary", {})
    table_summary = result.meta.get("table_borders", {}).get("summary", {})
    logical_parts = {
        key: page_summary.get(key)
        for key in ("statement", "main", "back")
        if page_summary.get(key) is not None
    }
    expected_keys = ("grid", "outer_only", "inner_only", "none", "mixed")
    table_summary = {key: table_summary.get(key) for key in expected_keys}
    return {
        "page_margins": {
            "sections": page_summary.get("sections"),
            "distinct_margins": page_summary.get("distinct_margins"),
            "logical_parts": logical_parts,
        },
        "table_borders": {
            "summary": table_summary,
        },
    }


def _summarize_footnotes(result: ParseResult) -> dict[str, object]:
    roles = []
    for role in ("footnote_text", "footnote_reference"):
        if role in result.roles:
            roles.append(role)
    numbering = result.meta.get("footnote_numbering", {})
    return {
        "roles": sorted(roles),
        "footnote_numbering": {
            "format": numbering.get("format"),
            "start": numbering.get("start"),
            "restart": numbering.get("restart"),
        },
    }


def _write_actual_summary(actual: dict[str, object]) -> None:
    config.ensure_base_dirs()
    output_path = config.OUTPUT_DIR / "regression_actual.json"
    output_path.write_text(json.dumps(actual, ensure_ascii=False, indent=2), encoding="utf-8")


def _normalize_summary(summary: dict[str, object]) -> dict[str, object]:
    normalized = json.loads(json.dumps(summary))
    for template in normalized.values():
        if isinstance(template, dict) and "roles" in template:
            roles = template.get("roles")
            if isinstance(roles, list):
                template["roles"] = sorted(roles)
    return normalized


def _update_table_borders(path: Path) -> bytes:
    with ZipFile(path) as archive:
        document_bytes = archive.read("word/document.xml")
    root = etree.fromstring(document_bytes)
    tables = root.findall(".//w:tbl", namespaces={"w": _W_NS})
    patterns = [
        {"top", "bottom", "left", "right", "insideH", "insideV"},
        {"top", "bottom", "left", "right"},
        {"insideH", "insideV"},
        set(),
        {"top", "left"},
    ]
    for table, borders in zip(tables, patterns, strict=False):
        _apply_tbl_borders(table, borders)
    return etree.tostring(root, xml_declaration=True, encoding="utf-8")


def _apply_tbl_borders(table: etree._Element, borders: set[str]) -> None:
    tbl_pr = table.find(f"{{{_W_NS}}}tblPr")
    if tbl_pr is None:
        tbl_pr = etree.SubElement(table, f"{{{_W_NS}}}tblPr")
    tbl_borders = tbl_pr.find(f"{{{_W_NS}}}tblBorders")
    if tbl_borders is None:
        tbl_borders = etree.SubElement(tbl_pr, f"{{{_W_NS}}}tblBorders")
    for key in ("top", "bottom", "left", "right", "insideH", "insideV"):
        child = tbl_borders.find(f"{{{_W_NS}}}{key}")
        if child is None:
            child = etree.SubElement(tbl_borders, f"{{{_W_NS}}}{key}")
        value = "single" if key in borders else "nil"
        child.set(f"{{{_W_NS}}}val", value)


def _rewrite_document_xml(path: Path, xml_bytes: bytes) -> None:
    with ZipFile(path) as archive:
        data = {name: archive.read(name) for name in archive.namelist()}
    data["word/document.xml"] = xml_bytes
    with ZipFile(path, "w") as archive:
        for name, content in data.items():
            archive.writestr(name, content)


def _inject_footnote_parts(path: Path) -> None:
    footnotes_xml = (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f'<w:footnotes xmlns:w="{_W_NS}">'
        f'<w:footnote w:id="1">'
        f'<w:p>'
        f'<w:pPr><w:pStyle w:val="FootnoteText"/></w:pPr>'
        f'<w:r>'
        f'<w:rPr>'
        f'<w:rStyle w:val="FootnoteReference"/>'
        f'<w:rFonts w:ascii="Times New Roman"/>'
        f'<w:sz w:val="20"/>'
        f'<w:b/>'
        f'</w:rPr>'
        f'<w:footnoteRef/>'
        f'</w:r>'
        f'<w:r>'
        f'<w:rPr>'
        f'<w:rFonts w:eastAsia="宋体"/>'
        f'<w:sz w:val="24"/>'
        f'</w:rPr>'
        f'<w:t>示例脚注</w:t>'
        f'</w:r>'
        f'</w:p>'
        f'</w:footnote>'
        f'</w:footnotes>'
    )
    settings_xml = (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f'<w:settings xmlns:w="{_W_NS}">'
        f'<w:footnotePr>'
        f'<w:numFmt w:val="decimal"/>'
        f'<w:numStart w:val="1"/>'
        f'<w:numRestart w:val="eachPage"/>'
        f'</w:footnotePr>'
        f'</w:settings>'
    )
    with ZipFile(path) as archive:
        data = {name: archive.read(name) for name in archive.namelist()}
    data["word/footnotes.xml"] = footnotes_xml.encode("utf-8")
    data["word/settings.xml"] = settings_xml.encode("utf-8")
    with ZipFile(path, "w") as archive:
        for name, content in data.items():
            archive.writestr(name, content)


if __name__ == "__main__":
    unittest.main()
