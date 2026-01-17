from __future__ import annotations

import json
import tempfile
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from time import perf_counter
from typing import Any, Iterable
import re
from zipfile import BadZipFile, ZipFile

from lxml import etree

from . import config
from .style_reader import ResolvedStyle, StyleDefaults, parse_styles_xml, resolve_style
from .section_rules import (
    BodyRangeRule,
    DEFAULT_SECTION_RULES,
    SectionPosition,
    SectionRule,
    serialize_section_rules,
)
from . import template_types
from .style_rule import REQUIRED_FIELDS, StyleRule, serialize_style_rules

_TITLE_KEYWORDS = ("\u6807\u9898", "\u7ae0", "heading 1", "title")
_BODY_KEYWORDS = ("\u6b63\u6587", "normal", "body", "text")
_SOURCE_PRIORITY = {
    "explicit": 0,
    "stack": 1,
    "outline": 2,
    "keyword": 3,
    "text": 4,
    "global": 5,
    "fallback": 6,
}
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_PPR_NS = {"w": _W_NS}
_BASE_TITLE_ROLE = "title_L1"
_BASE_BODY_ROLE = "body_L1"
_LEGACY_TITLE_ROLE = "chapter_title"
_LEGACY_BODY_ROLE = "body"
_GLOBAL_BODY_CANDIDATE_ROLE = "_global_body_candidate"
_TITLE_BODY_PATTERN = re.compile(r"^(title|body)_l([1-9]\d*)$", re.IGNORECASE)
_SPECIAL_ROLES = {
    "document_title",
    "document_title_en",
    "cover_title",
    "cover_info",
    "abstract_title",
    "abstract_body",
    "abstract_en_title",
    "abstract_en_body",
    "reference_title",
    "reference_body",
    "keyword_line",
    "toc_title",
    "toc_body",
    "figure_caption",
    "figure_body",
    "table_caption",
    "table_body",
    "figure_note",
    "table_note",
    "footnote_text",
    "footnote_reference",
}
_ROLE_ALIASES = {
    _LEGACY_TITLE_ROLE: _BASE_TITLE_ROLE,
    _LEGACY_BODY_ROLE: _BASE_BODY_ROLE,
}
_ROLE_GROUP_PATTERNS: list[tuple[re.Pattern[str], str]] = [
    (re.compile(r"^title_L[1-9]\d*$", re.IGNORECASE), "title"),
    (re.compile(r"^body_L[1-9]\d*$", re.IGNORECASE), "body"),
    (re.compile(r"^document_title$", re.IGNORECASE), "special_title"),
    (re.compile(r"^document_title_en$", re.IGNORECASE), "special_title"),
    (re.compile(r"^cover_title$", re.IGNORECASE), "cover"),
    (re.compile(r"^cover_info$", re.IGNORECASE), "cover"),
    (re.compile(r"^abstract_title$", re.IGNORECASE), "special_title"),
    (re.compile(r"^abstract_body$", re.IGNORECASE), "special_body"),
    (re.compile(r"^abstract_en_title$", re.IGNORECASE), "special_title"),
    (re.compile(r"^abstract_en_body$", re.IGNORECASE), "special_body"),
    (re.compile(r"^reference_title$", re.IGNORECASE), "special_title"),
    (re.compile(r"^reference_body$", re.IGNORECASE), "special_body"),
    (re.compile(r"^keyword_line$", re.IGNORECASE), "special_body"),
    (re.compile(r"^toc_title$", re.IGNORECASE), "toc"),
    (re.compile(r"^toc_body$", re.IGNORECASE), "toc"),
    (re.compile(r"^toc_body_L[1-9]\d*$", re.IGNORECASE), "toc"),
    (re.compile(r"^cover_", re.IGNORECASE), "cover"),
    (re.compile(r"^figure_caption$", re.IGNORECASE), "caption"),
    (re.compile(r"^figure_body$", re.IGNORECASE), "special_body"),
    (re.compile(r"^table_caption$", re.IGNORECASE), "caption"),
    (re.compile(r"^table_body$", re.IGNORECASE), "special_body"),
    (re.compile(r"^figure_note$", re.IGNORECASE), "note"),
    (re.compile(r"^table_note$", re.IGNORECASE), "note"),
    (re.compile(r"^footnote_text$", re.IGNORECASE), "note"),
    (re.compile(r"^footnote_reference$", re.IGNORECASE), "note"),
]
_HEADING_NAME_LEVEL_PATTERNS = [
    re.compile(r"\bheading\s*([1-9]\d*)\b", re.IGNORECASE),
    re.compile(r"\btitle\s*([1-9]\d*)\b", re.IGNORECASE),
    re.compile(r"\bchapter\s*([1-9]\d*)\b", re.IGNORECASE),
    re.compile(r"标题\s*([1-9]\d*)"),
    re.compile(r"标题\s*([一二三四五六七八九十])"),
    re.compile(r"([一二三四五六七八九十])\s*级标题"),
]
_CHINESE_NUMERAL_MAP = {
    "一": 1,
    "二": 2,
    "三": 3,
    "四": 4,
    "五": 5,
    "六": 6,
    "七": 7,
    "八": 8,
    "九": 9,
    "十": 10,
}
_ABSTRACT_STYLE_KEYWORDS = ("\u6458\u8981", "abstract")
_REFERENCE_STYLE_KEYWORDS = ("\u53c2\u8003\u6587\u732e", "references", "bibliography")
_CAPTION_DIRECTORY_TERMS = (
    "\u56fe\u8868\u76ee\u5f55",
    "\u56fe\u8868\u6e05\u5355",
    "\u56fe\u8868\u5217\u8868",
    "\u56fe\u8868\u7d22\u5f15",
    "\u56fe\u8868\u8bf4\u660e",
)
_TEXT_ABSTRACT_PATTERN = re.compile(r"^\s*(摘要|abstract)\s*$", re.IGNORECASE)
_TEXT_ABSTRACT_EN_PATTERN = re.compile(r"^\s*abstract\s*$", re.IGNORECASE)
_TEXT_ABSTRACT_INLINE_PATTERN = re.compile(r"^\s*(摘要|abstract)\s*[:：]\s*\S+", re.IGNORECASE)
_TEXT_REFERENCE_PATTERN = re.compile(r"^\s*(参考文献|references)\s*$", re.IGNORECASE)
_TEXT_KEYWORD_PATTERN = re.compile(r"^\s*(关键词|key\s*words?)\s*[:：]?", re.IGNORECASE)
_TEXT_TOC_PATTERN = re.compile(r"^\s*目\s*录\s*$|^\s*contents\s*$", re.IGNORECASE)
_TEXT_FIGURE_PATTERN = re.compile(r"^\s*(图\s*\d+|figure\s*\d+|fig\.\s*\d+)\b", re.IGNORECASE)
_TEXT_TABLE_PATTERN = re.compile(r"^\s*(表\s*\d+|table\s*\d+|tab\.\s*\d+)\b", re.IGNORECASE)
_TEXT_CHAPTER_PATTERN = re.compile(r"^\s*第[一二三四五六七八九十百千0-9]+章\b")
_TEXT_NUMBER_HEADING_PATTERN = re.compile(r"^\s*(\d+(?:\.\d+)*)\s+\S+")
_TEXT_CN_L1_PATTERN = re.compile(r"^\s*[一二三四五六七八九十]+[、.]\s*")
_TEXT_CN_L2_PATTERN = re.compile(r"^\s*（[一二三四五六七八九十]+）\s*")
_TEXT_ENGLISH_TITLE_PATTERN = re.compile(r"^[A-Za-z][A-Za-z0-9\s,:;()\-]+$")
_TEXT_COVER_MARKER_PATTERN = re.compile(r"^\s*(扉页|首页)\s*$")
_TOC_LEVEL_PATTERNS = [
    re.compile(r"\btoc\s*([1-9]\d*)\b", re.IGNORECASE),
    re.compile(r"\btoc([1-9]\d*)\b", re.IGNORECASE),
    re.compile(r"\bcontents?\s*([1-9]\d*)\b", re.IGNORECASE),
    re.compile(r"\u76ee\u5f55\s*([1-9]\d*)"),
    re.compile(r"\u76ee\u5f55([1-9]\d*)"),
]
_NOTE_TEXT_PATTERN = re.compile(r"^\s*(注|note|来源|source)[:：]?", re.IGNORECASE)
_NOTE_STYLE_KEYWORDS = ("\u6ce8", "note", "captionnote")
_FORMULA_NUMBER_PATTERN = re.compile(r"[\(（]\s*\d+(?:\.\d+)*\s*[\)）]")
_M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
_O_NS = "urn:schemas-microsoft-com:office:office"
_PAGE_MARGIN_KEYS = ("top", "bottom", "left", "right", "header", "footer", "gutter")
_DEFAULT_STATEMENT_KEYWORDS = (
    "\u58f0\u660e",
    "\u6388\u6743",
    "\u539f\u521b",
    "\u72ec\u521b\u6027",
    "\u5b66\u4f4d\u8bba\u6587\u539f\u521b\u6027",
    "\u7248\u6743",
)
_DEFAULT_BACK_KEYWORDS = (
    "\u53c2\u8003\u6587\u732e",
    "\u81f4\u8c22",
    "\u9e23\u8c22",
    "\u9644\u5f55",
    "\u53c2\u8003\u8d44\u6599",
)
_FOOTNOTE_TEXT_STYLE_IDS = ("footnotetext", "footnote_text")
_FOOTNOTE_REFERENCE_STYLE_IDS = ("footnotereference", "footnote_reference")
_FOOTNOTE_TEXT_STYLE_NAMES = (
    "footnote text",
    "\u811a\u6ce8\u6587\u672c",
    "\u811a\u6ce8\u6b63\u6587",
    "\u811a\u6ce8\u6587\u5b57",
)
_FOOTNOTE_REFERENCE_STYLE_NAMES = (
    "footnote reference",
    "\u811a\u6ce8\u5f15\u7528",
    "\u811a\u6ce8\u6807\u8bb0",
    "\u811a\u6ce8\u5e8f\u53f7",
    "\u811a\u6ce8\u7f16\u53f7",
)
_COVER_SECTION_KEY = "cover"


@dataclass
class WarningEntry:
    rule: str
    reason: str
    role: str | None = None
    style_id: str | None = None
    paragraph_index: int | None = None


@dataclass
class ParseLogState:
    template_path: Path
    start_time: datetime
    warnings: list[WarningEntry] = field(default_factory=list)
    role_sources: dict[str, str] = field(default_factory=dict)
    role_candidates: dict[str, list[str]] = field(default_factory=dict)
    style_count: int = 0
    detected_heading_levels: list[int] = field(default_factory=list)
    detected_heading_levels_overflow: list[int] = field(default_factory=list)
    error: str | None = None
    elapsed_sec: float | None = None


@dataclass
class SampleStats:
    count: int = 0
    first_index: int | None = None
    outline_min: int | None = None
    style_name: str | None = None
    weights: dict[str, dict[object, tuple[float, int]]] = field(default_factory=dict)
    text_samples: list[str] = field(default_factory=list)


@dataclass
class ParseResult:
    roles: dict[str, StyleRule]
    role_links: list[dict[str, object]]
    meta: dict[str, object]


class TemplateParser:
    def __init__(
        self,
        role_map: dict | None = None,
        role_map_path: str | None = None,
        outline_level_max: int = 0,
        max_heading_level: int | None = None,
        required_roles: list[str] | None = None,
        required_on_presence_map: dict[str, str] | None = None,
        strict: bool = False,
        allow_fallback: bool = True,
        template_type: str | None = None,
        section_rules: Iterable[SectionRule] | None = None,
    ) -> None:
        self.role_map = role_map
        self.role_map_path = role_map_path
        self.outline_level_max = outline_level_max
        self.max_heading_level = (
            config.DEFAULT_MAX_HEADING_LEVEL
            if max_heading_level is None
            else max_heading_level
        )
        self.required_roles = (
            list(config.DEFAULT_REQUIRED_ROLES)
            if required_roles is None
            else list(required_roles)
        )
        self.required_on_presence_map = _normalize_required_on_presence_map(required_on_presence_map)
        self.strict = strict
        self.allow_fallback = allow_fallback
        self._template_type_mode = (template_type or "generic").strip() or "generic"
        self.template_type = self._template_type_mode
        if section_rules is None:
            self._section_rules_mode = None
            self.section_rules = tuple(DEFAULT_SECTION_RULES)
        else:
            self._section_rules_mode = tuple(section_rules)
            self.section_rules = self._section_rules_mode
        if self._section_rules_mode is None:
            self._cover_detection_enabled = any(
                rule.key == _COVER_SECTION_KEY for rule in self.section_rules
            )
        else:
            self._cover_detection_enabled = any(
                rule.key == _COVER_SECTION_KEY for rule in self._section_rules_mode
            )
        self._effective_role_map_path: Path | None = None
        self._role_map_required = False
        self._log_state: ParseLogState | None = None
        self._detected_heading_levels: list[int] = []
        self._detected_heading_levels_overflow: list[int] = []
        self._last_log_state: ParseLogState | None = None
        self._last_global_body_rule: StyleRule | None = None
        self._reset_extra_meta()

    def parse(self, template_path: str) -> ParseResult:
        path = Path(template_path)
        started = perf_counter()
        self._detected_heading_levels = []
        self._detected_heading_levels_overflow = []
        self._reset_extra_meta()
        self._log_state = ParseLogState(
            template_path=path,
            start_time=datetime.now(),
        )
        self._apply_template_type(path)
        try:
            _ensure_readable_file(path)
            self._effective_role_map_path = self._resolve_role_map_path(path)
            roles = self._parse_template(path)
        except Exception as exc:
            self._log_state.error = str(exc)
            self._log_state.elapsed_sec = perf_counter() - started
            _write_log(self._log_state)
            self._log_state = None
            raise
        self._log_state.elapsed_sec = perf_counter() - started
        _write_log(self._log_state)
        role_links = _build_role_links(roles)
        meta = self._build_meta()
        self._last_log_state = self._log_state
        self._log_state = None
        return ParseResult(roles=roles, role_links=role_links, meta=meta)

    def parse_roles(self, template_path: str) -> dict[str, StyleRule]:
        result = self.parse(template_path)
        return result.roles

    def export_json(
        self,
        rules: dict[str, StyleRule] | ParseResult,
        output_path: str | None = None,
    ) -> None:
        if isinstance(rules, ParseResult):
            result = rules
            roles = result.roles
        else:
            result = None
            roles = rules
        if self.strict:
            _validate_strict(
                roles,
                required_roles=self.required_roles,
                required_on_presence_map=self.required_on_presence_map,
            )
        if output_path is None:
            config.ensure_base_dirs()
            output = config.DEFAULT_OUTPUT_PATH
        else:
            output = Path(output_path)
            output.parent.mkdir(parents=True, exist_ok=True)
        if result is not None:
            role_links = result.role_links
            meta = result.meta
        else:
            role_links = []
            meta = self._build_meta()
        payload = serialize_style_rules(
            roles,
            allow_fallback=self.allow_fallback,
            strict=self.strict,
            schema_version="2.0",
            role_links=role_links,
            meta=meta,
        )
        with output.open("w", encoding="utf-8") as handle:
            json.dump(payload, handle, ensure_ascii=False, indent=2)

    def _build_meta(self) -> dict[str, object]:
        return {
            "max_heading_level": self.max_heading_level,
            "detected_heading_levels": list(self._detected_heading_levels),
            "detected_heading_levels_overflow": list(self._detected_heading_levels_overflow),
            "required_roles": _normalize_required_roles(self.required_roles),
            "required_on_presence_map": dict(_normalize_required_on_presence_map(self.required_on_presence_map)),
            "allow_fallback": self.allow_fallback,
            "strict": self.strict,
            "template_type": self.template_type,
            "section_rules": serialize_section_rules(self.section_rules),
            "cover_detection_enabled": self._cover_detection_enabled,
            "page_margins": self._page_margins,
            "table_borders": self._table_borders,
            "title_spacing": self._title_spacing,
            "footnote_numbering": self._footnote_numbering,
            "toc_levels": self._toc_levels,
            "header_footer": self._header_footer,
        }

    def _reset_extra_meta(self) -> None:
        self._page_margins = _default_page_margins()
        self._table_borders = _default_table_borders()
        self._title_spacing = {}
        self._footnote_numbering = {}
        self._toc_levels = _default_toc_levels()
        self._header_footer = {}

    def _apply_template_type(self, template_path: Path) -> None:
        if self._template_type_mode.lower() == "auto" and self._section_rules_mode is None:
            resolved = template_types.detect_template_type(template_path)
            self.template_type = resolved.key
            self.section_rules = resolved.section_rules
        else:
            self.template_type = self._template_type_mode
            if self._section_rules_mode is not None:
                self.section_rules = self._section_rules_mode
            elif self.template_type.lower() == "generic":
                self.section_rules = tuple(DEFAULT_SECTION_RULES)
            else:
                resolved = template_types.resolve_template_type(self.template_type)
                self.section_rules = resolved.section_rules
        self._cover_detection_enabled = any(
            rule.key == _COVER_SECTION_KEY for rule in self.section_rules
        )

    def _resolve_role_map_path(self, template_path: Path) -> Path | None:
        if self.role_map is not None:
            self._role_map_required = False
            return None
        if self.role_map_path:
            self._role_map_required = True
            return Path(self.role_map_path)
        self._role_map_required = False
        return template_path.parent / "role_mapping.json"

    def _parse_template(self, template_path: Path) -> dict[str, StyleRule]:
        section_rules_for_detection = tuple(self.section_rules)
        self._page_margins = _parse_page_margins(
            template_path,
            section_rules_for_detection,
            self._log_state,
            enable_cover_detection=self._cover_detection_enabled,
        )
        self._table_borders = _parse_table_borders(
            template_path,
            self._log_state,
        )
        self._footnote_numbering = _parse_footnote_numbering(
            template_path,
            self._log_state,
        )
        styles_bytes, theme_bytes = _read_docx_parts(template_path)
        theme_map = _parse_theme_map(theme_bytes, log_state=self._log_state)
        with tempfile.TemporaryDirectory() as tmpdir:
            styles_path = Path(tmpdir) / "styles.xml"
            styles_path.write_bytes(styles_bytes)
            styles, defaults = parse_styles_xml(styles_path)
        self._header_footer = _parse_header_footer(
            template_path,
            styles=styles,
            defaults=defaults,
            theme_map=theme_map,
            page_margins=self._page_margins,
            log_state=self._log_state,
        )
        if self._log_state is not None:
            self._log_state.style_count = len(styles)
        if not styles:
            raise ValueError(f"no paragraph styles found in template: {template_path}")

        style_names = {
            style.name.lower()
            for style in styles.values()
            if style.name and style.name.strip()
        }
        name_map, id_map = _load_role_map(
            role_map=self.role_map,
            role_map_path=self._effective_role_map_path,
            role_map_required=self._role_map_required,
            style_ids=set(styles.keys()),
            style_names=style_names,
            log_state=self._log_state,
        )

        samples = _collect_paragraph_samples(template_path, log_state=self._log_state)
        (
            detected_heading_levels,
            detected_heading_levels_overflow,
            outline_level_min,
        ) = _detect_heading_levels(
            styles,
            samples,
            name_map,
            id_map,
            self.max_heading_level,
        )
        outline_levels, outline_min_doc, outline_max_doc = _collect_outline_levels(
            template_path,
        )
        if outline_min_doc is not None:
            if outline_level_min is None or outline_min_doc < outline_level_min:
                outline_level_min = outline_min_doc
        effective_outline_max = self.outline_level_max
        if effective_outline_max is None or effective_outline_max <= 0:
            effective_outline_max = outline_max_doc
        elif outline_max_doc is not None:
            effective_outline_max = min(effective_outline_max, outline_max_doc)

        if outline_levels and outline_level_min is not None:
            normalized = {
                level - outline_level_min + 1
                for level in outline_levels
                if level >= outline_level_min
            }
            if self.max_heading_level is not None:
                normalized = {level for level in normalized if level <= self.max_heading_level}
            detected_heading_levels = sorted(set(detected_heading_levels).union(normalized))

        self._detected_heading_levels = list(detected_heading_levels)
        self._detected_heading_levels_overflow = list(detected_heading_levels_overflow)
        if self._log_state is not None:
            self._log_state.detected_heading_levels = list(self._detected_heading_levels)
            self._log_state.detected_heading_levels_overflow = list(
                self._detected_heading_levels_overflow
            )
        toc_levels: dict[int, set[str]] = {}
        title_spacing: dict[str, dict[str, int]] = {}
        body_role_candidates = _collect_body_candidates_by_stack(
            template_path,
            name_map=name_map,
            id_map=id_map,
            max_heading_level=self.max_heading_level,
            outline_level_max=effective_outline_max,
            outline_level_min=outline_level_min,
            section_rules=section_rules_for_detection,
            enable_cover_detection=self._cover_detection_enabled,
            toc_levels=toc_levels,
            title_spacing=title_spacing,
        )
        self._toc_levels = _serialize_toc_levels(toc_levels)
        self._title_spacing = title_spacing
        style_order = {style_id: index for index, style_id in enumerate(styles.keys())}

        role_candidates: dict[str, list[tuple[str, Any]]] = {}
        for style_id in styles:
            stats = samples.get(style_id)
            if stats is None or stats.count <= 0:
                continue
            resolved = resolve_style(style_id, styles, defaults, theme_map=theme_map)
            match = _match_role(
                resolved,
                name_map,
                id_map,
                outline_level_max=effective_outline_max,
                outline_level_min=outline_level_min,
                max_heading_level=self.max_heading_level,
                stats=stats,
            )
            if match is None:
                continue
            role, source = match
            role_candidates.setdefault(role, []).append(
                {
                    "resolved": resolved,
                    "source": source,
                    "stats": stats,
                    "order": style_order[style_id],
                }
            )

        global_candidates = body_role_candidates.pop(_GLOBAL_BODY_CANDIDATE_ROLE, None)

        def _run_resolved_style() -> ResolvedStyle:
            fonts = defaults.fonts.apply_theme(theme_map)
            return ResolvedStyle(
                style_id="",
                name=None,
                fonts=fonts,
                font_name=fonts.preferred_name(),
                font_size_pt=defaults.font_size_pt,
                bold=defaults.bold,
                alignment=defaults.alignment,
                space_before_pt=defaults.space_before_pt,
                space_after_pt=defaults.space_after_pt,
                line_rule=defaults.line_rule,
                line_twips=defaults.line_twips,
                outline_level=defaults.outline_level,
            )

        def _append_body_candidate(role: str, style_id: str, info: dict[str, object]) -> None:
            stats = info.get("stats") or samples.get(style_id)
            derived_stats = _derive_role_stats(
                stats,
                count=int(info.get("count", 0)),
                first_index=info.get("first_index"),
            )
            if not style_id or style_id not in styles:
                if derived_stats is None:
                    return
                resolved = _run_resolved_style()
                role_candidates.setdefault(role, []).append(
                    {
                        "resolved": resolved,
                        "source": info.get("source", "run"),
                        "stats": derived_stats,
                        "order": int(info.get("first_index", 0)),
                        "prefer_samples": True,
                    }
                )
                return
            resolved = resolve_style(style_id, styles, defaults, theme_map=theme_map)
            role_candidates.setdefault(role, []).append(
                {
                    "resolved": resolved,
                    "source": info.get("source", "stack"),
                    "stats": derived_stats,
                    "order": int(info.get("first_index", style_order.get(style_id, 0))),
                    "prefer_samples": True,
                }
            )

        global_candidate_entries: list[dict[str, Any]] = []
        if global_candidates:
            for style_id, info in global_candidates.items():
                resolved = resolve_style(style_id, styles, defaults, theme_map=theme_map)
                stats = info.get("stats") or samples.get(style_id)
                derived_stats = _derive_role_stats(
                    stats,
                    count=int(info.get("count", 0)),
                    first_index=info.get("first_index"),
                )
                global_candidate_entries.append(
                    {
                        "resolved": resolved,
                        "source": info.get("source", "global"),
                        "stats": derived_stats,
                        "order": int(info.get("first_index", style_order.get(style_id, 0))),
                        "prefer_samples": True,
                    }
                )

        for role, style_map in body_role_candidates.items():
            for style_id, info in style_map.items():
                _append_body_candidate(role, style_id, info)

        if global_candidates and _BASE_BODY_ROLE not in role_candidates:
            for style_id, info in global_candidates.items():
                _append_body_candidate(_BASE_BODY_ROLE, style_id, info)

        global_body_rule: StyleRule | None = None
        if global_candidate_entries:
            selected = _choose_best_candidate(_BASE_BODY_ROLE, global_candidate_entries)
            if selected is not None:
                global_body_rule = selected[0]
        self._last_global_body_rule = global_body_rule

        footnote_text_stats, footnote_ref_stats, footnote_text_ids, footnote_ref_ids = (
            _collect_footnote_samples(template_path, self._log_state)
        )

        selected_candidates = _resolve_role_conflicts(role_candidates, log_state=self._log_state)
        rules: dict[str, StyleRule] = {}
        for role in sorted(selected_candidates):
            candidate = selected_candidates[role]
            rule = _build_style_rule(
                role,
                candidate["resolved"],
                candidate.get("stats"),
                candidate.get("prefer_samples", False),
            )
            rules[role] = rule
            if self._log_state is not None:
                self._log_state.role_sources[role] = _describe_candidate(candidate)
                self._log_state.role_candidates[role] = [
                    _describe_candidate(item) for item in role_candidates.get(role, [])
                ]
                _tag_warnings_for_style(
                    self._log_state,
                    role,
                    candidate["resolved"].style_id,
                )

        footnote_rules = _build_footnote_rules(
            styles,
            defaults,
            theme_map,
            footnote_text_stats,
            footnote_ref_stats,
            footnote_text_ids,
            footnote_ref_ids,
        )
        for role, rule in footnote_rules.items():
            if role not in rules:
                rules[role] = rule

        rules = _apply_fallbacks(
            rules,
            allow_fallback=self.allow_fallback,
            strict=self.strict,
            log_state=self._log_state,
            required_roles=self.required_roles,
            required_on_presence_map=self.required_on_presence_map,
            global_body_rule=global_body_rule,
        )
        if self.strict:
            _validate_strict(
                rules,
                required_roles=self.required_roles,
                required_on_presence_map=self.required_on_presence_map,
            )
        return rules


def _build_role_links(rules: dict[str, StyleRule]) -> list[dict[str, object]]:
    links: list[dict[str, object]] = []
    level_links: list[tuple[int, dict[str, object]]] = []
    for role in rules:
        level = _extract_title_level_from_role(role)
        if level is None:
            continue
        body_role = f"body_L{level}"
        if body_role not in rules:
            continue
        level_links.append(
            (
                level,
                {
                    "title_role": role,
                    "body_role": body_role,
                    "level": level,
                },
            )
        )
    for _, link in sorted(level_links, key=lambda item: item[0]):
        links.append(link)

    section_pairs = [
        ("abstract_title", "abstract_body", "abstract"),
        ("abstract_en_title", "abstract_en_body", "abstract_en"),
        ("reference_title", "reference_body", "reference"),
        ("figure_caption", "figure_note", "figure"),
        ("table_caption", "table_note", "table"),
    ]
    for title_role, body_role, section in section_pairs:
        if title_role in rules and body_role in rules:
            links.append(
                {
                    "title_role": title_role,
                    "body_role": body_role,
                    "section": section,
                }
            )
    section_titles: dict[str, str] = {}
    for role in rules:
        match = re.match(r"^section_(.+)_title$", role, flags=re.IGNORECASE)
        if not match:
            continue
        section_key = match.group(1)
        if section_key:
            section_titles.setdefault(section_key, role)
    for section_key, title_role in sorted(section_titles.items()):
        body_role = f"section_{section_key}_body"
        if body_role in rules:
            links.append(
                {
                    "title_role": title_role,
                    "body_role": body_role,
                    "section": section_key,
                }
            )
    return links


def _normalize_role(role: str) -> str:
    role_clean = role.strip()
    if not role_clean:
        return role_clean
    lower = role_clean.lower()
    alias = _ROLE_ALIASES.get(lower)
    if alias:
        return alias
    match = _TITLE_BODY_PATTERN.match(lower)
    if match:
        prefix = "title" if match.group(1).lower() == "title" else "body"
        level = int(match.group(2))
        return f"{prefix}_L{level}"
    if lower in _SPECIAL_ROLES:
        return lower
    return role_clean


def _resolve_role_group(role: str) -> str | None:
    role_name = _normalize_role(role)
    if not role_name:
        return None
    for pattern, group in _ROLE_GROUP_PATTERNS:
        if pattern.match(role_name):
            return group
    return None


def _validate_role_name(role: str) -> str:
    normalized = _normalize_role(role)
    if not normalized:
        raise ValueError("role mapping value must be a non-empty string")
    if _resolve_role_group(normalized) is None:
        allowed = (
            "title_Lk/body_Lk or one of "
            + ", ".join(sorted(_SPECIAL_ROLES))
        )
        raise ValueError(f"role mapping value must be {allowed}, got {role!r}")
    return normalized


def _extract_title_level_from_role(role: str) -> int | None:
    match = re.match(r"^title_L([1-9]\d*)$", role, flags=re.IGNORECASE)
    if not match:
        return None
    try:
        return int(match.group(1))
    except ValueError:
        return None


def _parse_heading_level_from_name(name: str) -> int | None:
    name = name.strip()
    if not name:
        return None
    for pattern in _HEADING_NAME_LEVEL_PATTERNS:
        match = pattern.search(name)
        if not match:
            continue
        token = match.group(match.lastindex or 1)
        if token in _CHINESE_NUMERAL_MAP:
            return _CHINESE_NUMERAL_MAP[token]
        try:
            return int(token)
        except ValueError:
            return None
    if re.search(r"\bchapter\b", name, re.IGNORECASE):
        return 1
    compact = re.sub(r"\s+", "", name)
    if "章节标题" in compact:
        return 1
    return None


def _is_heading_candidate_name(name: str) -> bool:
    return _parse_heading_level_from_name(name) is not None


def _detect_heading_levels(
    styles: dict[str, Any],
    samples: dict[str, SampleStats],
    name_map: dict[str, str],
    id_map: dict[str, str],
    max_heading_level: int | None,
) -> tuple[list[int], list[int], int | None]:
    detected: set[int] = set()
    overflow: set[int] = set()

    def _record(level: int, explicit: bool) -> None:
        if level <= 0:
            return
        if max_heading_level is not None and level > max_heading_level:
            if explicit:
                overflow.add(level)
            return
        detected.add(level)

    for role in list(name_map.values()) + list(id_map.values()):
        level = _extract_title_level_from_role(role)
        if level is not None:
            _record(level, explicit=True)

    names: set[str] = set()
    for style in styles.values():
        if style.name:
            names.add(style.name)
    for stats in samples.values():
        if stats.style_name:
            names.add(stats.style_name)
    for name in names:
        level = _parse_heading_level_from_name(name)
        if level is not None:
            _record(level, explicit=False)

    explicit_title_style_ids: set[str] = set()
    for style_id, style in styles.items():
        style_id_lower = style_id.lower()
        mapped = id_map.get(style_id_lower)
        if mapped and _extract_title_level_from_role(mapped) is not None:
            explicit_title_style_ids.add(style_id)
            continue
        name = (style.name or "").strip().lower()
        if name and name in name_map:
            mapped = name_map[name]
            if _extract_title_level_from_role(mapped) is not None:
                explicit_title_style_ids.add(style_id)

    candidate_outline_levels: list[int] = []
    candidate_styles: list[str] = []
    for style_id, style in styles.items():
        stats = samples.get(style_id)
        name = style.name or (stats.style_name if stats else "") or ""
        is_heading_name = _is_heading_candidate_name(name)
        has_outline = False
        outline_levels: list[int] = []
        if style.outline_level is not None:
            outline_levels.append(style.outline_level)
        if stats and stats.outline_min is not None:
            outline_levels.append(stats.outline_min)
        if outline_levels:
            has_outline = True
        candidate = is_heading_name or has_outline or style_id in explicit_title_style_ids
        if not candidate:
            continue
        candidate_styles.append(style_id)
        if outline_levels:
            candidate_outline_levels.append(min(outline_levels))

    min_outline_level: int | None = None
    if candidate_outline_levels:
        min_outline_level = min(candidate_outline_levels)
        for style_id in candidate_styles:
            style = styles[style_id]
            outline_levels: list[int] = []
            if style.outline_level is not None:
                outline_levels.append(style.outline_level)
            stats = samples.get(style_id)
            if stats and stats.outline_min is not None:
                outline_levels.append(stats.outline_min)
            if not outline_levels:
                continue
            outline_level = min(outline_levels)
            level = outline_level - min_outline_level + 1
            _record(level, explicit=False)

    return sorted(detected), sorted(overflow), min_outline_level


def _normalize_required_roles(required_roles: Iterable[str] | None) -> list[str]:
    if required_roles is None:
        required_roles = config.DEFAULT_REQUIRED_ROLES
    normalized: list[str] = []
    seen: set[str] = set()
    for role in required_roles:
        role_norm = _normalize_role(role)
        if role_norm in seen:
            continue
        normalized.append(role_norm)
        seen.add(role_norm)
    return normalized


def _normalize_required_on_presence_map(
    required_on_presence_map: dict[str, str] | None,
) -> dict[str, str]:
    if required_on_presence_map is None:
        required_on_presence_map = config.DEFAULT_REQUIRED_ON_PRESENCE_MAP
    normalized: dict[str, str] = {}
    for key, value in required_on_presence_map.items():
        if not isinstance(key, str) or not isinstance(value, str):
            raise ValueError("required_on_presence_map keys and values must be strings")
        key_norm = _normalize_role(key)
        value_norm = _normalize_role(value)
        if _resolve_role_group(key_norm) is None or _resolve_role_group(value_norm) is None:
            raise ValueError(
                f"required_on_presence_map must use valid role names, got {key!r} -> {value!r}"
            )
        normalized[key_norm] = value_norm
    return normalized


def _conditional_required_roles(
    rules: dict[str, StyleRule],
    required_on_presence_map: dict[str, str] | None,
) -> list[str]:
    normalized = _normalize_required_on_presence_map(required_on_presence_map)
    targets: list[str] = []
    for trigger_role, target_role in normalized.items():
        if trigger_role in rules and target_role not in targets:
            targets.append(target_role)
    return targets


def _ensure_required_roles(
    rules: dict[str, Any],
    required_roles: Iterable[str] | None = None,
) -> None:
    required = _normalize_required_roles(required_roles)
    missing = [role for role in required if role not in rules]
    if missing:
        missing_text = ", ".join(missing)
        raise ValueError(f"strict mode missing roles: {missing_text}")


def _validate_strict(
    rules: dict[str, StyleRule],
    required_roles: Iterable[str] | None = None,
    required_on_presence_map: dict[str, str] | None = None,
) -> None:
    required = _normalize_required_roles(required_roles)
    conditional_targets = _conditional_required_roles(rules, required_on_presence_map)
    required_all: list[str] = []
    for role in list(required) + list(conditional_targets):
        if role not in required_all:
            required_all.append(role)
    missing_roles = [role for role in required_all if role not in rules]
    if missing_roles:
        missing_text = ", ".join(missing_roles)
        raise ValueError(f"strict mode missing roles: {missing_text}")
    errors: list[str] = []
    for role, rule in rules.items():
        missing = _missing_required_fields(rule)
        if missing:
            missing_text = ", ".join(missing)
            errors.append(f"{role}: {missing_text}")
    if errors:
        raise ValueError("strict mode missing fields: " + "; ".join(errors))


def _ensure_readable_file(path: Path) -> None:
    if not path.exists():
        raise FileNotFoundError(f"template not found: {path}")
    if not path.is_file():
        raise IsADirectoryError(f"template path is not a file: {path}")
    try:
        with path.open("rb"):
            pass
    except PermissionError as exc:
        raise PermissionError(f"template is not readable: {path}") from exc


def _read_docx_parts(template_path: Path) -> tuple[bytes, bytes | None]:
    try:
        with ZipFile(template_path) as archive:
            styles_bytes = archive.read("word/styles.xml")
            try:
                theme_bytes = archive.read("word/theme/theme1.xml")
            except KeyError:
                theme_bytes = None
    except BadZipFile as exc:
        raise ValueError(f"invalid docx file: {template_path}") from exc
    except KeyError as exc:
        raise ValueError(f"missing required part in docx: {exc}") from exc
    return styles_bytes, theme_bytes


def _parse_theme_map(
    theme_bytes: bytes | None,
    log_state: ParseLogState | None = None,
) -> dict[str, str]:
    if not theme_bytes:
        return {}
    try:
        root = etree.fromstring(theme_bytes)
    except Exception as exc:
        if log_state is not None:
            _warn(
                log_state,
                rule="theme_font",
                reason=f"failed to parse theme1.xml ({exc})",
            )
        return {}
    ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
    scheme = root.find(".//a:fontScheme", namespaces=ns)
    if scheme is None:
        return {}
    mapping: dict[str, str] = {}

    def _read_font(prefix: str, elem: etree._Element | None) -> None:
        if elem is None:
            return
        latin = elem.find("a:latin", namespaces=ns)
        if latin is not None:
            typeface = latin.get("typeface")
            if typeface:
                mapping[f"{prefix}Ascii"] = typeface
                mapping[f"{prefix}HAnsi"] = typeface
        ea = elem.find("a:ea", namespaces=ns)
        if ea is not None:
            typeface = ea.get("typeface")
            if typeface:
                mapping[f"{prefix}EastAsia"] = typeface

    _read_font("major", scheme.find("a:majorFont", namespaces=ns))
    _read_font("minor", scheme.find("a:minorFont", namespaces=ns))
    return mapping


def _load_role_map(
    role_map: dict | None,
    role_map_path: Path | None,
    role_map_required: bool,
    style_ids: set[str],
    style_names: set[str],
    log_state: ParseLogState | None,
) -> tuple[dict[str, str], dict[str, str]]:
    if role_map is not None:
        raw = role_map
    else:
        raw = {}
        if role_map_path:
            if role_map_path.exists():
                if not role_map_path.is_file():
                    raise ValueError(f"role mapping path is not a file: {role_map_path}")
                with role_map_path.open("r", encoding="utf-8") as handle:
                    raw = json.load(handle)
            elif role_map_required:
                raise FileNotFoundError(f"role mapping file not found: {role_map_path}")
    if not isinstance(raw, dict):
        raise ValueError("role mapping must be a JSON object")

    name_map: dict[str, str] = {}
    id_map: dict[str, str] = {}
    style_ids_lower = {style_id.lower() for style_id in style_ids}
    for key, value in raw.items():
        if not isinstance(key, str) or not isinstance(value, str):
            raise ValueError("role mapping keys and values must be strings")
        key_lower = key.strip().lower()
        value_lower = value.strip().lower()
        value_lower = _validate_role_name(value_lower)
        if key_lower.startswith("id:"):
            style_id = key_lower[3:].strip()
            if style_id:
                id_map[style_id] = value_lower
        else:
            if key_lower in style_ids_lower and key_lower not in style_names:
                if log_state is not None:
                    _warn(
                        log_state,
                        rule="role_map",
                        reason="ignored key matching style_id without 'id:' prefix",
                        style_id=key,
                    )
                continue
            if key_lower:
                name_map[key_lower] = value_lower
    return name_map, id_map


def _match_role(
    resolved: Any,
    name_map: dict[str, str],
    id_map: dict[str, str],
    outline_level_max: int | None,
    outline_level_min: int | None,
    max_heading_level: int | None,
    stats: SampleStats | None,
) -> tuple[str, str] | None:
    style_id = (resolved.style_id or "").lower()
    if style_id and style_id in id_map:
        return id_map[style_id], "explicit"
    style_name = (resolved.name or "").lower()
    if style_name and style_name in name_map:
        return name_map[style_name], "explicit"

    outline_level = resolved.outline_level
    if stats and stats.outline_min is not None:
        if outline_level is None or stats.outline_min < outline_level:
            outline_level = stats.outline_min
    if outline_level_max is not None and outline_level is not None and outline_level_min is not None:
        if outline_level <= outline_level_max:
            normalized_level = outline_level - outline_level_min + 1
            if normalized_level > 0 and (
                max_heading_level is None or normalized_level <= max_heading_level
            ):
                return f"title_L{normalized_level}", "outline"

    if style_name:
        special = _match_special_role_by_style_name(style_name)
        if special:
            return special, "keyword"
        title_level = _parse_heading_level_from_name(style_name)
        if title_level is not None:
            if max_heading_level is None or title_level <= max_heading_level:
                return f"title_L{title_level}", "keyword"
        if _match_keyword(style_name, _BODY_KEYWORDS):
            return _BASE_BODY_ROLE, "keyword"
    special_text = _match_special_role_by_text(stats)
    if special_text:
        return special_text, "text"
    return None


def _match_keyword(name: str, keywords: Iterable[str]) -> bool:
    return any(keyword in name for keyword in keywords)


def _match_special_role_by_style_name(name: str) -> str | None:
    name_stripped = name.strip()
    if not name_stripped:
        return None
    lower = name_stripped.lower()
    if lower in {"title", "标题"}:
        return "document_title"
    if "toc" in lower or "contents" in lower or "\u76ee\u5f55" in name_stripped:
        level = _parse_toc_level_from_name(name_stripped)
        if level is not None:
            return f"toc_body_L{level}"
        return "toc_body"
    if any(term in name_stripped for term in _CAPTION_DIRECTORY_TERMS):
        return None
    if any(keyword in lower for keyword in _ABSTRACT_STYLE_KEYWORDS):
        if "abstract" in lower and "摘要" not in name_stripped:
            return "abstract_en_title"
        return "abstract_title"
    if any(keyword in lower for keyword in _REFERENCE_STYLE_KEYWORDS):
        return "reference_title"
    if "caption" in lower:
        if "figure" in lower:
            return "figure_caption"
        if "table" in lower:
            return "table_caption"
        return None
    if "图题" in name_stripped:
        return "figure_caption"
    if "表题" in name_stripped:
        return "table_caption"
    return None


def _match_special_role_by_text(stats: SampleStats | None) -> str | None:
    if stats is None:
        return None
    for text in stats.text_samples:
        inline_role = _match_inline_abstract_role(text)
        if inline_role:
            return inline_role
        if _TEXT_ABSTRACT_EN_PATTERN.match(text):
            return "abstract_en_title"
        if _TEXT_ABSTRACT_PATTERN.match(text):
            return "abstract_title"
        if _TEXT_REFERENCE_PATTERN.match(text):
            return "reference_title"
        if _TEXT_TOC_PATTERN.match(text):
            return "toc_title"
        if _TEXT_FIGURE_PATTERN.match(text):
            return "figure_caption"
        if _TEXT_TABLE_PATTERN.match(text):
            return "table_caption"
    return None


def _match_special_role_by_text_value(text: str) -> str | None:
    if not text:
        return None
    inline_role = _match_inline_abstract_role(text)
    if inline_role:
        return inline_role
    if _TEXT_ABSTRACT_EN_PATTERN.match(text):
        return "abstract_en_title"
    if _TEXT_ABSTRACT_PATTERN.match(text):
        return "abstract_title"
    if _TEXT_REFERENCE_PATTERN.match(text):
        return "reference_title"
    if _TEXT_TOC_PATTERN.match(text):
        return "toc_title"
    if _TEXT_FIGURE_PATTERN.match(text):
        return "figure_caption"
    if _TEXT_TABLE_PATTERN.match(text):
        return "table_caption"
    return None


def _match_inline_abstract_role(text: str) -> str | None:
    if not text:
        return None
    match = _TEXT_ABSTRACT_INLINE_PATTERN.match(text)
    if not match:
        return None
    token = match.group(1)
    if token is None:
        return None
    if token.lower() == "abstract" and "摘要" not in text:
        return "abstract_en_title"
    return "abstract_title"


def _is_cover_info_line(text: str) -> bool:
    if not text:
        return False
    compact = re.sub(r"\s+", "", text)
    if not compact:
        return False
    if compact.startswith(
        (
            "专业",
            "年级",
            "姓名",
            "作者",
            "学号",
            "学院",
            "院系",
            "系别",
            "班级",
            "单位",
            "学校",
            "学生",
            "指导教师",
            "指导老师",
            "导师",
            "日期",
            "时间",
        )
    ):
        return True
    if re.fullmatch(r"\d{4}年\d{1,2}月(\d{1,2}日)?", compact):
        return True
    return False


def _is_cover_marker(text: str) -> bool:
    if not text:
        return False
    return _TEXT_COVER_MARKER_PATTERN.match(text) is not None


def _match_title_role_by_text_value(text: str) -> str | None:
    if not text:
        return None
    if _TEXT_CHAPTER_PATTERN.match(text):
        return "title_L1"
    if _TEXT_CN_L1_PATTERN.match(text):
        return "title_L1"
    if _TEXT_CN_L2_PATTERN.match(text):
        return "title_L2"
    match = _TEXT_NUMBER_HEADING_PATTERN.match(text)
    if match:
        segments = match.group(1).split(".")
        level = len(segments)
        if level > 0:
            return f"title_L{level}"
    return None


def _is_keyword_line(text: str) -> bool:
    if not text:
        return False
    return _TEXT_KEYWORD_PATTERN.match(text) is not None


def _is_toc_line(style_name: str | None) -> bool:
    if not style_name:
        return False
    lower = style_name.lower()
    return "toc" in lower or "contents" in lower or "\u76ee\u5f55" in style_name


def _parse_toc_level_from_name(name: str | None) -> int | None:
    if not name:
        return None
    for pattern in _TOC_LEVEL_PATTERNS:
        match = pattern.search(name)
        if not match:
            continue
        try:
            return int(match.group(1))
        except ValueError:
            return None
    return None


def _parse_toc_level_from_text(text: str | None) -> int | None:
    if not text:
        return None
    clean = text.strip()
    if not clean:
        return None
    match = re.match(r"^\s*(\d+(?:\.\d+)*)\b", clean)
    if match:
        segments = match.group(1).split(".")
        level = len(segments)
        if 1 <= level <= 3:
            return level
    if _TEXT_CN_L1_PATTERN.match(clean):
        return 1
    if _TEXT_CN_L2_PATTERN.match(clean):
        return 2
    return None


def _looks_like_toc_entry(text: str | None) -> bool:
    if not text:
        return False
    clean = text.strip()
    if not clean:
        return False
    if _TEXT_TOC_PATTERN.match(clean):
        return False
    if re.search(r"\t\s*\d+\s*$", clean):
        return True
    if re.search(r"\.{2,}\s*\d+\s*$", clean):
        return True
    if _parse_toc_level_from_text(clean) is not None and re.search(r"\d+\s*$", clean):
        return True
    return False


def _extract_first_run_font_size(paragraph: Any) -> float | None:
    for run in getattr(paragraph, "runs", []):
        text = getattr(run, "text", "")
        if text and not text.isspace():
            return _extract_run_font_size(run)
    return None


def _is_document_title_candidate(
    paragraph: Any,
    text: str,
    align_enum: Any,
) -> bool:
    if not text:
        return False
    if _TEXT_ABSTRACT_PATTERN.match(text) or _TEXT_REFERENCE_PATTERN.match(text):
        return False
    if _TEXT_TOC_PATTERN.match(text):
        return False
    if _TEXT_FIGURE_PATTERN.match(text) or _TEXT_TABLE_PATTERN.match(text):
        return False
    if _is_keyword_line(text):
        return False
    alignment = _extract_paragraph_alignment(paragraph, align_enum)
    if alignment == "CENTER":
        return True
    font_size = _extract_first_run_font_size(paragraph)
    if font_size is not None and font_size >= 16.0:
        return True
    return False


def _is_document_title_en_candidate(
    text: str,
    paragraph: Any,
    align_enum: Any,
    exclude_abstract: bool = True,
) -> bool:
    if not text:
        return False
    if exclude_abstract and _TEXT_ABSTRACT_PATTERN.match(text):
        return False
    if _TEXT_REFERENCE_PATTERN.match(text):
        return False
    if _TEXT_KEYWORD_PATTERN.match(text):
        return False
    if not _TEXT_ENGLISH_TITLE_PATTERN.match(text.strip()):
        return False
    letters = sum(1 for ch in text if ch.isalpha())
    if letters < 10:
        return False
    alignment = _extract_paragraph_alignment(paragraph, align_enum)
    if alignment == "CENTER":
        return True
    font_size = _extract_first_run_font_size(paragraph)
    if font_size is not None and font_size >= 16.0:
        return True
    return False


def _is_blank_paragraph(paragraph: Any) -> bool:
    text = getattr(paragraph, "text", "")
    if text and text.strip():
        return False
    for run in getattr(paragraph, "runs", []):
        run_text = getattr(run, "text", "")
        if run_text and run_text.strip():
            return False
    return True


def _is_note_paragraph(text: str, style_name: str | None) -> bool:
    if text and _NOTE_TEXT_PATTERN.match(text):
        return True
    if style_name:
        lower = style_name.lower()
        return any(keyword in lower for keyword in _NOTE_STYLE_KEYWORDS)
    return False


def _is_table_paragraph(paragraph: Any) -> bool:
    element = getattr(paragraph, "_p", None)
    while element is not None:
        tag = getattr(element, "tag", "")
        if tag.endswith("}tc"):
            return True
        element = getattr(element, "getparent", lambda: None)()
    return False


def _paragraph_has_drawing(paragraph: Any) -> bool:
    element = getattr(paragraph, "_element", None)
    if element is None:
        return False
    for node in element.iter():
        tag = getattr(node, "tag", "")
        if not isinstance(tag, str):
            continue
        if tag.endswith("}drawing") or tag.endswith("}pict"):
            return True
    return False


def _choose_best_candidate(
    role: str,
    candidates: list[dict[str, Any]],
) -> tuple[StyleRule, dict[str, Any]] | None:
    if not candidates:
        return None
    sorted_candidates = sorted(
        candidates,
        key=lambda item: _candidate_sort_key(role, item),
    )
    best_candidate = sorted_candidates[0]
    best_rule = _build_style_rule(
        role,
        best_candidate["resolved"],
        best_candidate.get("stats"),
        best_candidate.get("prefer_samples", False),
    )
    return best_rule, best_candidate


def _candidate_sort_key(role: str, candidate: dict[str, Any]) -> tuple[int, int, int, int, int, str]:
    resolved = candidate["resolved"]
    stats = candidate.get("stats")
    rule = _build_style_rule(
        role,
        resolved,
        stats,
        candidate.get("prefer_samples", False),
    )
    count = stats.count if stats else 0
    order_index = candidate["order"]
    if stats and stats.first_index is not None:
        order_index = stats.first_index
    completeness = _required_field_score(rule)
    source_priority = _SOURCE_PRIORITY.get(candidate.get("source", "unknown"), 99)
    style_id = resolved.style_id or ""
    if role == "toc_body":
        toc_level = None
        if stats and stats.style_name:
            toc_level = _parse_toc_level_from_name(stats.style_name)
        if toc_level is None:
            toc_level = _parse_toc_level_from_name(resolved.name)
        toc_sort = toc_level if toc_level is not None else 99
        return (source_priority, toc_sort, order_index, -count, -completeness, style_id)
    return (source_priority, 0, -count, -completeness, order_index, style_id)


def _resolve_role_conflicts(
    role_candidates: dict[str, list[dict[str, Any]]],
    log_state: ParseLogState | None,
) -> dict[str, dict[str, Any]]:
    sorted_candidates: dict[str, list[dict[str, Any]]] = {}
    best_keys: dict[str, tuple[int, int, int, int, str]] = {}
    for role, candidates in role_candidates.items():
        if not candidates:
            continue
        ordered = sorted(candidates, key=lambda item: _candidate_sort_key(role, item))
        sorted_candidates[role] = ordered
        best_keys[role] = _candidate_sort_key(role, ordered[0])

    roles_by_group: dict[str | None, list[str]] = {}
    for role in sorted_candidates:
        group = _resolve_role_group(role)
        roles_by_group.setdefault(group, []).append(role)

    selected: dict[str, dict[str, Any]] = {}
    for group, roles in roles_by_group.items():
        if group is None:
            for role in roles:
                selected[role] = sorted_candidates[role][0]
            continue
        if group in {"cover", "toc"}:
            for role in roles:
                selected[role] = sorted_candidates[role][0]
            continue
        used_styles: dict[str, str] = {}
        roles.sort(key=lambda item: (best_keys[item], item))
        for role in roles:
            candidates = sorted_candidates[role]
            chosen: dict[str, Any] | None = None
            for candidate in candidates:
                style_id = candidate["resolved"].style_id or ""
                if not style_id or style_id not in used_styles:
                    chosen = candidate
                    break
                if log_state is not None:
                    _warn(
                        log_state,
                        rule="conflict_resolved",
                        reason=f"skip {style_id} used by {used_styles[style_id]}",
                        role=role,
                        style_id=style_id,
                    )
            if chosen is None:
                chosen = candidates[0]
                style_id = chosen["resolved"].style_id or ""
                if style_id and style_id in used_styles and log_state is not None:
                    other_role = used_styles[style_id]
                    _warn(
                        log_state,
                        rule="shared_style",
                        reason=f"shared with {other_role}",
                        role=role,
                        style_id=style_id,
                    )
            style_id = chosen["resolved"].style_id or ""
            if style_id:
                used_styles.setdefault(style_id, role)
            selected[role] = chosen

    return selected


def _required_field_score(rule: StyleRule) -> int:
    return sum(1 for field in REQUIRED_FIELDS if getattr(rule, field) is not None)


def _build_style_rule(
    role: str,
    resolved: Any,
    stats: SampleStats | None = None,
    prefer_samples: bool = False,
) -> StyleRule:
    line_rule, line_value, line_unit = _resolve_line_spacing(resolved)
    if stats:
        sample_line = _sample_mode(stats, "line_spacing")
        if sample_line and (prefer_samples or line_rule is None):
            line_rule, line_value, line_unit = sample_line
        elif sample_line and line_rule == sample_line[0]:
            if line_value is None:
                line_value = sample_line[1]
            if line_unit is None:
                line_unit = sample_line[2]
        sample_name = _sample_mode(stats, "font_name")
        prefer_sample_name = prefer_samples and sample_name is not None
        resolved_font_name = resolved.font_name
        if prefer_sample_name or resolved.font_name is None:
            if sample_name is not None:
                resolved_font_name = sample_name
    else:
        prefer_sample_name = False
        resolved_font_name = resolved.font_name
    if stats and (prefer_samples or resolved.fonts.ascii is None):
        ascii_name = _sample_mode(stats, "font_name_ascii") or resolved.fonts.ascii
    else:
        ascii_name = resolved.fonts.ascii
    if stats and (prefer_samples or resolved.fonts.hAnsi is None):
        h_ansi_name = _sample_mode(stats, "font_name_hAnsi") or resolved.fonts.hAnsi
    else:
        h_ansi_name = resolved.fonts.hAnsi
    if stats and (prefer_samples or resolved.fonts.eastAsia is None):
        east_asia_name = _sample_mode(stats, "font_name_eastAsia") or resolved.fonts.eastAsia
    else:
        east_asia_name = resolved.fonts.eastAsia
    preferred_name = east_asia_name or h_ansi_name or ascii_name
    if preferred_name is not None and not prefer_sample_name:
        resolved_font_name = preferred_name
    font_size_pt = resolved.font_size_pt
    if stats and (prefer_samples or font_size_pt is None):
        font_size_pt = _sample_mode(stats, "font_size_pt") or font_size_pt
    bold = resolved.bold
    if stats and (prefer_samples or bold is None):
        sample_bold = _sample_mode(stats, "bold")
        if sample_bold is not None:
            bold = sample_bold
    alignment = resolved.alignment
    if stats and (prefer_samples or alignment is None):
        alignment = _sample_mode(stats, "alignment") or alignment
    space_before_pt = resolved.space_before_pt
    if stats and (prefer_samples or space_before_pt is None):
        sample_before = _sample_mode(stats, "space_before_pt")
        if sample_before is not None:
            space_before_pt = sample_before
    space_after_pt = resolved.space_after_pt
    if stats and (prefer_samples or space_after_pt is None):
        sample_after = _sample_mode(stats, "space_after_pt")
        if sample_after is not None:
            space_after_pt = sample_after
    space_before_value: float | None = None
    space_before_unit: str | None = None
    space_after_value: float | None = None
    space_after_unit: str | None = None
    if stats:
        sample_before = _sample_mode(stats, "space_before")
        if sample_before is not None:
            space_before_value, space_before_unit = sample_before
        sample_after = _sample_mode(stats, "space_after")
        if sample_after is not None:
            space_after_value, space_after_unit = sample_after
    if space_before_unit is None and space_before_pt is not None:
        space_before_value = float(space_before_pt)
        space_before_unit = "PT"
    if space_after_unit is None and space_after_pt is not None:
        space_after_value = float(space_after_pt)
        space_after_unit = "PT"
    indent_left_pt = None
    indent_right_pt = None
    indent_first_line_pt = None
    indent_hanging_pt = None
    if stats and (prefer_samples or indent_left_pt is None):
        indent_left_pt = _sample_mode(stats, "indent_left_pt") or indent_left_pt
    if stats and (prefer_samples or indent_right_pt is None):
        indent_right_pt = _sample_mode(stats, "indent_right_pt") or indent_right_pt
    if stats and (prefer_samples or indent_first_line_pt is None):
        indent_first_line_pt = _sample_mode(stats, "indent_first_line_pt") or indent_first_line_pt
    if stats and (prefer_samples or indent_hanging_pt is None):
        indent_hanging_pt = _sample_mode(stats, "indent_hanging_pt") or indent_hanging_pt
    if alignment is None:
        alignment = "LEFT"
    if line_rule is None:
        line_rule, line_value, line_unit = "SINGLE", 1.0, "MULTIPLE"
    if space_before_pt is None:
        space_before_pt = 0.0
    if space_after_pt is None:
        space_after_pt = 0.0
    if indent_left_pt is None:
        indent_left_pt = 0.0
    if indent_right_pt is None:
        indent_right_pt = 0.0
    if indent_first_line_pt is None:
        indent_first_line_pt = 0.0
    if indent_hanging_pt is None:
        indent_hanging_pt = 0.0
    if bold is None:
        bold = False
    return StyleRule(
        role=role,
        font_name=resolved_font_name,
        font_name_ascii=ascii_name,
        font_name_eastAsia=east_asia_name,
        font_name_hAnsi=h_ansi_name,
        font_size_pt=font_size_pt,
        bold=bold,
        alignment=alignment,
        line_spacing_rule=line_rule,
        line_spacing_value=line_value,
        line_spacing_unit=line_unit,
        space_before_pt=space_before_pt,
        space_after_pt=space_after_pt,
        space_before_value=space_before_value,
        space_before_unit=space_before_unit,
        space_after_value=space_after_value,
        space_after_unit=space_after_unit,
        indent_left_pt=indent_left_pt,
        indent_right_pt=indent_right_pt,
        indent_first_line_pt=indent_first_line_pt,
        indent_hanging_pt=indent_hanging_pt,
    )


def _resolve_line_spacing(resolved: Any) -> tuple[str | None, float | None, str | None]:
    line_rule = resolved.line_rule
    line_twips = resolved.line_twips
    if line_rule in {"exact", "atLeast"} and line_twips is not None:
        rule = "EXACTLY" if line_rule == "exact" else "AT_LEAST"
        return rule, line_twips / 20.0, "PT"
    if line_rule in {None, "auto"} and line_twips is not None:
        multiple = line_twips / 240.0
        return "MULTIPLE", multiple, "MULTIPLE"
    return None, None, None


def _collect_paragraph_samples(
    template_path: Path,
    log_state: ParseLogState | None,
) -> dict[str, SampleStats]:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    except ImportError as exc:
        raise ImportError("python-docx is required to parse paragraph samples") from exc

    document = Document(str(template_path))
    samples: dict[str, SampleStats] = {}
    samples_by_name: dict[str, SampleStats] = {}
    paragraph_index = 0
    sample_index = 0
    for paragraph in _iter_paragraphs(document):
        paragraph_index += 1
        if _is_blank_paragraph(paragraph):
            continue
        style = paragraph.style
        if style is None:
            continue
        style_id = getattr(style, "style_id", None)
        if not style_id:
            continue
        style_name = getattr(style, "name", None)
        stats = samples.get(style_id)
        if stats is None:
            name_key = None
            if style_name and style_name.strip():
                name_key = style_name.strip().lower()
            if name_key:
                stats = samples_by_name.get(name_key)
                if stats is None:
                    stats = SampleStats()
                    samples_by_name[name_key] = stats
            else:
                stats = SampleStats()
            samples[style_id] = stats
        if stats.style_name is None and style_name:
            stats.style_name = style_name
        paragraph_text = getattr(paragraph, "text", None)
        if paragraph_text:
            trimmed = paragraph_text.strip()
            if trimmed and trimmed not in stats.text_samples:
                if len(stats.text_samples) < 3:
                    stats.text_samples.append(trimmed)
        stats.count += 1
        if stats.first_index is None:
            stats.first_index = paragraph_index
        outline_level = _extract_paragraph_outline_level(paragraph)
        if outline_level is not None:
            if stats.outline_min is None or outline_level < stats.outline_min:
                stats.outline_min = outline_level

        alignment = _extract_paragraph_alignment(paragraph, WD_ALIGN_PARAGRAPH)
        sample_index = _add_sample(stats, "alignment", alignment, 1.0, sample_index)

        line_spacing = _extract_paragraph_line_spacing(
            paragraph,
            WD_LINE_SPACING,
            log_state,
            style_id,
            paragraph_index,
        )
        sample_index = _add_sample(stats, "line_spacing", line_spacing, 1.0, sample_index)

        (
            space_before_pt,
            space_after_pt,
            space_before_value,
            space_before_unit,
            space_after_value,
            space_after_unit,
        ) = _extract_paragraph_spacing_detail(paragraph)
        sample_index = _add_sample(
            stats,
            "space_before_pt",
            space_before_pt,
            1.0,
            sample_index,
        )
        sample_index = _add_sample(
            stats,
            "space_after_pt",
            space_after_pt,
            1.0,
            sample_index,
        )
        if space_before_value is not None and space_before_unit is not None:
            sample_index = _add_sample(
                stats,
                "space_before",
                (space_before_value, space_before_unit),
                1.0,
                sample_index,
            )
        if space_after_value is not None and space_after_unit is not None:
            sample_index = _add_sample(
                stats,
                "space_after",
                (space_after_value, space_after_unit),
                1.0,
                sample_index,
            )
        indent_left, indent_right, indent_first, indent_hanging = _extract_paragraph_indentation(
            paragraph,
        )
        sample_index = _add_sample(
            stats,
            "indent_left_pt",
            indent_left,
            1.0,
            sample_index,
        )
        sample_index = _add_sample(
            stats,
            "indent_right_pt",
            indent_right,
            1.0,
            sample_index,
        )
        sample_index = _add_sample(
            stats,
            "indent_first_line_pt",
            indent_first,
            1.0,
            sample_index,
        )
        sample_index = _add_sample(
            stats,
            "indent_hanging_pt",
            indent_hanging,
            1.0,
            sample_index,
        )

        for run in paragraph.runs:
            text = run.text
            if not text or text.isspace():
                continue
            weight = float(len(text.strip()))
            if weight <= 0:
                continue
            font_name, font_ascii, font_hansi, font_east_asia = _extract_run_fonts(run)
            sample_index = _add_sample(
                stats,
                "font_name",
                font_name,
                weight,
                sample_index,
            )
            sample_index = _add_sample(
                stats,
                "font_name_ascii",
                font_ascii,
                weight,
                sample_index,
            )
            sample_index = _add_sample(
                stats,
                "font_name_hAnsi",
                font_hansi,
                weight,
                sample_index,
            )
            sample_index = _add_sample(
                stats,
                "font_name_eastAsia",
                font_east_asia,
                weight,
                sample_index,
            )
            font_size = _extract_run_font_size(run)
            sample_index = _add_sample(
                stats,
                "font_size_pt",
                font_size,
                weight,
                sample_index,
            )
            bold = _extract_run_bold(run)
            sample_index = _add_sample(stats, "bold", bold, weight, sample_index)
    return samples


def _collect_outline_levels(template_path: Path) -> tuple[set[int], int | None, int | None]:
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError("python-docx is required to parse outline levels") from exc

    document = Document(str(template_path))
    levels: set[int] = set()
    for paragraph in _iter_paragraphs(document):
        outline_level = _extract_paragraph_outline_level(paragraph)
        if outline_level is not None:
            levels.add(outline_level)
    if not levels:
        return set(), None, None
    return levels, min(levels), max(levels)


def _collect_body_candidates_by_stack(
    template_path: Path,
    name_map: dict[str, str],
    id_map: dict[str, str],
    max_heading_level: int | None,
    outline_level_max: int | None,
    outline_level_min: int | None,
    section_rules: Iterable[SectionRule] | None,
    enable_cover_detection: bool = True,
    toc_levels: dict[int, set[str]] | None = None,
    title_spacing: dict[str, dict[str, int]] | None = None,
) -> dict[str, dict[str, dict[str, object]]]:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    except ImportError as exc:
        raise ImportError("python-docx is required to parse body roles") from exc

    document = Document(str(template_path))
    paragraphs = list(_iter_paragraphs(document))
    section_rules = tuple(section_rules or ())

    def _paragraph_has_page_break(paragraph: Any) -> bool:
        element = getattr(paragraph, "_element", None)
        if element is None:
            return False
        for node in element.iter():
            tag = getattr(node, "tag", "")
            if not isinstance(tag, str):
                continue
            if "}" in tag:
                ns, local = tag[1:].split("}", 1)
            else:
                ns, local = "", tag
            if ns != _W_NS:
                continue
            if local == "lastRenderedPageBreak":
                return True
            if local == "br":
                br_type = node.get(_attr_name("type")) or ""
                if br_type == "page":
                    return True
        return False

    content_page_numbers: list[int] = []
    page_number = 1
    for paragraph in paragraphs:
        if not _is_blank_paragraph(paragraph) and not _is_table_paragraph(paragraph):
            content_page_numbers.append(page_number)
        if _paragraph_has_page_break(paragraph):
            page_number += 1

    spacing_targets = {"abstract_title", "abstract_en_title", "toc_title"}

    block_items = list(_iter_block_items(document))
    spacing_index_by_element: dict[int, int] = {}
    for idx, (kind, item) in enumerate(block_items):
        if kind != "paragraph":
            continue
        element = getattr(item, "_element", None)
        if element is not None:
            spacing_index_by_element[id(element)] = idx

    def _is_spacing_blank(item: tuple[str, Any]) -> bool:
        kind, obj = item
        if kind != "paragraph":
            return False
        if _is_table_paragraph(obj):
            return False
        return _is_blank_paragraph(obj)

    def _count_spacing_around(index: int) -> tuple[int, int]:
        idx = max(index - 1, 0)
        before = 0
        i = idx - 1
        while i >= 0 and _is_spacing_blank(block_items[i]):
            before += 1
            i -= 1
        after = 0
        i = idx + 1
        total = len(block_items)
        while i < total and _is_spacing_blank(block_items[i]):
            after += 1
            i += 1
        return before, after

    def _record_title_spacing(role: str, paragraph: Any) -> None:
        if title_spacing is None:
            return
        if role not in spacing_targets:
            return
        if role in title_spacing:
            return
        element = getattr(paragraph, "_element", None)
        if element is None:
            return
        spacing_index = spacing_index_by_element.get(id(element))
        if spacing_index is None:
            return
        before, after = _count_spacing_around(spacing_index)
        title_spacing[role] = {"before": before, "after": after}
    def _section_front_limit(total: int) -> int:
        if total <= 0:
            return 0
        base = max(5, total // 3)
        return min(30, base)

    def _resolve_caption_roles(
        items: list[tuple[str, Any]],
    ) -> tuple[dict[int, str], dict[str, bool]]:
        object_indices: dict[str, list[int]] = {"figure": [], "table": []}
        candidates: list[tuple[int, str, int]] = []
        for index, (kind, item) in enumerate(items):
            if kind == "table":
                object_indices["table"].append(index)
                continue
            paragraph = item
            if _paragraph_has_drawing(paragraph):
                object_indices["figure"].append(index)
            text = getattr(paragraph, "text", "") or ""
            text = text.strip()
            if not text:
                continue
            element = getattr(paragraph, "_element", None)
            if element is None:
                continue
            element_id = id(element)
            if _TEXT_FIGURE_PATTERN.match(text):
                candidates.append((index, "figure", element_id))
            elif _TEXT_TABLE_PATTERN.match(text):
                candidates.append((index, "table", element_id))
        presence = {key: bool(indices) for key, indices in object_indices.items()}
        confirmed: dict[int, str] = {}
        for index, kind, element_id in candidates:
            indices = object_indices[kind]
            if not indices:
                confirmed[element_id] = f"{kind}_caption"
                continue
            distance = min(abs(index - obj_index) for obj_index in indices)
            if distance <= 2:
                confirmed[element_id] = f"{kind}_caption"
        return confirmed, presence

    caption_roles_by_id, caption_object_presence = _resolve_caption_roles(block_items)

    content_total = sum(
        1
        for paragraph in paragraphs
        if not _is_blank_paragraph(paragraph) and not _is_table_paragraph(paragraph)
    )
    front_limit = _section_front_limit(content_total)
    back_start = max(0, content_total - front_limit)
    last_page = content_page_numbers[-1] if content_page_numbers else 1
    role_candidates: dict[str, dict[str, dict[str, object]]] = {}
    title_stack: list[int] = []
    abstract_mode = False
    abstract_en_mode = False
    reference_mode = False
    toc_mode = False
    cover_detected = False
    cover_title_recorded = False
    pending_cover_title: tuple[str | None, int, Any] | None = None
    document_title_found = False
    document_title_en_found = False
    document_title_index: int | None = None
    pending_captions: list[dict[str, object]] = []
    paragraph_index = 0
    content_index = -1
    active_section: dict[str, object] | None = None

    def _resolve_caption_roles(
        items: list[tuple[str, Any]],
    ) -> tuple[dict[int, str], dict[str, bool]]:
        object_indices: dict[str, list[int]] = {"figure": [], "table": []}
        candidates: list[tuple[int, str, int]] = []
        for index, (kind, item) in enumerate(items):
            if kind == "table":
                object_indices["table"].append(index)
                continue
            paragraph = item
            if _paragraph_has_drawing(paragraph):
                object_indices["figure"].append(index)
            text = getattr(paragraph, "text", "") or ""
            text = text.strip()
            if not text:
                continue
            element = getattr(paragraph, "_element", None)
            if element is None:
                continue
            element_id = id(element)
            if _TEXT_FIGURE_PATTERN.match(text):
                candidates.append((index, "figure", element_id))
            elif _TEXT_TABLE_PATTERN.match(text):
                candidates.append((index, "table", element_id))
        presence = {key: bool(indices) for key, indices in object_indices.items()}
        confirmed: dict[int, str] = {}
        for index, kind, element_id in candidates:
            indices = object_indices[kind]
            if not indices:
                confirmed[element_id] = f"{kind}_caption"
                continue
            distance = min(abs(index - obj_index) for obj_index in indices)
            if distance <= 2:
                confirmed[element_id] = f"{kind}_caption"
        return confirmed, presence

    def _record(
        role: str,
        style_id: str | None,
        index: int,
        source: str,
        paragraph: Any | None = None,
        prefer_first_run: bool = False,
        allow_missing_style: bool = False,
    ) -> None:
        if not style_id and not allow_missing_style:
            return
        role_map = role_candidates.setdefault(role, {})
        role_key = style_id or ""
        entry = role_map.get(role_key)
        if entry is None:
            entry = {
                "count": 0,
                "first_index": index,
                "source": source,
                "stats": SampleStats(),
                "sample_index": 0,
            }
            role_map[role_key] = entry
        entry["count"] = int(entry["count"]) + 1
        if entry.get("first_index") is None or index < int(entry["first_index"]):
            entry["first_index"] = index
        if entry.get("source") != "explicit" and source == "explicit":
            entry["source"] = source
        if paragraph is not None:
            _update_role_stats(
                entry,
                paragraph,
                style_id or "",
                index,
                WD_ALIGN_PARAGRAPH,
                WD_LINE_SPACING,
                prefer_first_run,
            )

    def _update_role_stats(
        entry: dict[str, object],
        paragraph: Any,
        style_id: str,
        index: int,
        align_enum: Any,
        line_enum: Any,
        prefer_first_run: bool,
    ) -> None:
        stats = entry["stats"]
        sample_index = int(entry.get("sample_index", 0))
        stats.count += 1
        if stats.first_index is None:
            stats.first_index = index
        paragraph_text = getattr(paragraph, "text", None)
        if paragraph_text:
            trimmed = paragraph_text.strip()
            if trimmed and trimmed not in stats.text_samples:
                if len(stats.text_samples) < 3:
                    stats.text_samples.append(trimmed)
        style = getattr(paragraph, "style", None)
        style_name = getattr(style, "name", None) if style is not None else None
        if stats.style_name is None and style_name:
            stats.style_name = style_name
        outline_level = _extract_paragraph_outline_level(paragraph)
        if outline_level is not None:
            if stats.outline_min is None or outline_level < stats.outline_min:
                stats.outline_min = outline_level
        if style_id:
            alignment = _extract_paragraph_alignment(paragraph, align_enum)
            sample_index = _add_sample(stats, "alignment", alignment, 1.0, sample_index)
            line_spacing = _extract_paragraph_line_spacing(
                paragraph,
                line_enum,
                None,
                style_id,
                index,
            )
            sample_index = _add_sample(stats, "line_spacing", line_spacing, 1.0, sample_index)
            (
                space_before_pt,
                space_after_pt,
                space_before_value,
                space_before_unit,
                space_after_value,
                space_after_unit,
            ) = _extract_paragraph_spacing_detail(paragraph)
            sample_index = _add_sample(
                stats,
                "space_before_pt",
                space_before_pt,
                1.0,
                sample_index,
            )
            sample_index = _add_sample(
                stats,
                "space_after_pt",
                space_after_pt,
                1.0,
                sample_index,
            )
            if space_before_value is not None and space_before_unit is not None:
                sample_index = _add_sample(
                    stats,
                    "space_before",
                    (space_before_value, space_before_unit),
                    1.0,
                    sample_index,
                )
            if space_after_value is not None and space_after_unit is not None:
                sample_index = _add_sample(
                    stats,
                    "space_after",
                    (space_after_value, space_after_unit),
                    1.0,
                    sample_index,
                )
            indent_left, indent_right, indent_first, indent_hanging = _extract_paragraph_indentation(
                paragraph,
            )
            sample_index = _add_sample(
                stats,
                "indent_left_pt",
                indent_left,
                1.0,
                sample_index,
            )
            sample_index = _add_sample(
                stats,
                "indent_right_pt",
                indent_right,
                1.0,
                sample_index,
            )
            sample_index = _add_sample(
                stats,
                "indent_first_line_pt",
                indent_first,
                1.0,
                sample_index,
            )
            sample_index = _add_sample(
                stats,
                "indent_hanging_pt",
                indent_hanging,
                1.0,
                sample_index,
            )
        if prefer_first_run:
            first_run = None
            for run in paragraph.runs:
                text = run.text
                if text and not text.isspace():
                    first_run = run
                    break
            if first_run is not None:
                weight = 1000.0
                font_name, font_ascii, font_hansi, font_east_asia = _extract_run_fonts(first_run)
                sample_index = _add_sample(stats, "font_name", font_name, weight, sample_index)
                sample_index = _add_sample(
                    stats,
                    "font_name_ascii",
                    font_ascii,
                    weight,
                    sample_index,
                )
                sample_index = _add_sample(
                    stats,
                    "font_name_hAnsi",
                    font_hansi,
                    weight,
                    sample_index,
                )
                sample_index = _add_sample(
                    stats,
                    "font_name_eastAsia",
                    font_east_asia,
                    weight,
                    sample_index,
                )
                font_size = _extract_run_font_size(first_run)
                sample_index = _add_sample(stats, "font_size_pt", font_size, weight, sample_index)
                bold = _extract_run_bold(first_run)
                sample_index = _add_sample(stats, "bold", bold, weight, sample_index)
        for run in paragraph.runs:
            text = run.text
            if not text or text.isspace():
                continue
            weight = float(len(text.strip()))
            if weight <= 0:
                continue
            font_name, font_ascii, font_hansi, font_east_asia = _extract_run_fonts(run)
            sample_index = _add_sample(
                stats,
                "font_name",
                font_name,
                weight,
                sample_index,
            )
            sample_index = _add_sample(
                stats,
                "font_name_ascii",
                font_ascii,
                weight,
                sample_index,
            )
            sample_index = _add_sample(
                stats,
                "font_name_hAnsi",
                font_hansi,
                weight,
                sample_index,
            )
            sample_index = _add_sample(
                stats,
                "font_name_eastAsia",
                font_east_asia,
                weight,
                sample_index,
            )
            font_size = _extract_run_font_size(run)
            sample_index = _add_sample(
                stats,
                "font_size_pt",
                font_size,
                weight,
                sample_index,
            )
            bold = _extract_run_bold(run)
            sample_index = _add_sample(stats, "bold", bold, weight, sample_index)
        entry["sample_index"] = sample_index

    def _record_toc_level(style_name: str | None, text: str | None) -> int | None:
        if toc_levels is None:
            return None
        level = _parse_toc_level_from_name(style_name)
        if level is None:
            level = _parse_toc_level_from_text(text)
        if level is None or level < 1 or level > 3:
            return None
        toc_levels.setdefault(level, set())
        if style_name and style_name.strip():
            toc_levels[level].add(style_name.strip())
        return level

    def _prune_pending(current_index: int) -> None:
        pending_captions[:] = [
            item for item in pending_captions if current_index - int(item["index"]) <= 2
        ]

    def _commit_cover_title() -> None:
        nonlocal pending_cover_title, cover_title_recorded
        if cover_title_recorded or pending_cover_title is None:
            return
        style_id, index, paragraph = pending_cover_title
        _record("cover_title", style_id, index, "stack", paragraph)
        cover_title_recorded = True
        pending_cover_title = None

    def _commit_document_title_from_pending() -> None:
        nonlocal pending_cover_title, document_title_found, document_title_index
        if document_title_found or pending_cover_title is None:
            return
        style_id, index, paragraph = pending_cover_title
        _record("document_title", style_id, index, "stack", paragraph)
        document_title_found = True
        document_title_index = index
        pending_cover_title = None

    def _section_title_role(rule: SectionRule) -> str:
        if rule.key == _COVER_SECTION_KEY:
            return "cover_title"
        return f"section_{rule.key}_title"

    def _section_body_role(rule: SectionRule) -> str:
        if rule.key == _COVER_SECTION_KEY:
            return "cover_info"
        return f"section_{rule.key}_body"

    def _content_page_for_index(index: int) -> int | None:
        if index < 0 or index >= len(content_page_numbers):
            return None
        return content_page_numbers[index]

    def _section_position_matches(rule: SectionRule, index: int) -> bool:
        if content_total <= 0 or index < 0:
            return False
        if rule.position == SectionPosition.FIRST_PAGE:
            return _content_page_for_index(index) == 1
        if rule.position == SectionPosition.FRONT:
            return index < front_limit
        if rule.position == SectionPosition.BACK:
            return index >= back_start
        if rule.position == SectionPosition.LAST_PAGE:
            return _content_page_for_index(index) == last_page
        return True

    def _match_section_title(text: str, style_name: str | None, index: int) -> SectionRule | None:
        if not section_rules or not text or index < 0:
            return None
        style_lower = style_name.lower() if style_name else ""
        for rule in section_rules:
            if not _section_position_matches(rule, index):
                continue
            if any(_keyword_matches(text, keyword) for keyword in rule.title_keywords if keyword):
                return rule
            if rule.title_style_names and style_lower:
                if any(style_lower == name.lower() for name in rule.title_style_names if name):
                    return rule
        return None

    def _match_section_content(text: str, index: int) -> SectionRule | None:
        if not section_rules or not text or index < 0:
            return None
        for rule in section_rules:
            if not rule.content_keywords:
                continue
            if not _section_position_matches(rule, index):
                continue
            if any(_keyword_matches(text, keyword) for keyword in rule.content_keywords if keyword):
                return rule
        return None

    def _next_active_section(rule: SectionRule) -> dict[str, object] | None:
        remaining = rule.body_paragraph_limit or 0
        if rule.body_range == BodyRangeRule.FIXED_PARAGRAPHS:
            remaining -= 1
            if remaining <= 0:
                return None
        return {
            "rule": rule,
            "body_role": _section_body_role(rule),
            "remaining": remaining,
        }

    def _keyword_matches(text: str, keyword: str) -> bool:
        clean = text.strip()
        kw = keyword.strip()
        if not clean or not kw:
            return False
        clean_lower = clean.lower()
        kw_lower = kw.lower()
        if clean_lower == kw_lower or kw_lower in clean_lower:
            return True
        clean_compact = re.sub(r"\s+", "", clean_lower)
        kw_compact = re.sub(r"\s+", "", kw_lower)
        if not clean_compact or not kw_compact:
            return False
        if clean_compact == kw_compact:
            return True
        return kw_compact in clean_compact

    def _paragraph_has_math(paragraph: Any) -> bool:
        element = getattr(paragraph, "_element", None)
        if element is None:
            return False
        for node in element.iter():
            tag = getattr(node, "tag", "")
            if not isinstance(tag, str):
                continue
            if "}" in tag:
                ns, local = tag[1:].split("}", 1)
            else:
                ns, local = "", tag
            if ns == _M_NS and local in {"oMath", "oMathPara"}:
                return True
            if ns == _O_NS and local == "OLEObject":
                prog_id = node.get("ProgID") or node.get("progId") or ""
                if "Equation" in prog_id or "MathType" in prog_id:
                    return True
        return False

    def _strip_formula_numbers(text: str) -> str:
        return _FORMULA_NUMBER_PATTERN.sub("", text)

    def _extract_formula_flags(text: str, paragraph: Any) -> tuple[bool, bool, bool]:
        has_math = _paragraph_has_math(paragraph)
        if not has_math:
            return False, False, False
        has_number = _FORMULA_NUMBER_PATTERN.search(text) is not None
        stripped = _strip_formula_numbers(text)
        has_non_number_text = bool(stripped.strip())
        return True, has_number, has_non_number_text

    for paragraph in paragraphs:
        paragraph_index += 1
        text = getattr(paragraph, "text", "") or ""
        style = getattr(paragraph, "style", None)
        style_id = getattr(style, "style_id", None) if style is not None else None
        style_name = getattr(style, "name", None) if style is not None else None
        is_blank = _is_blank_paragraph(paragraph)
        outline_level = _extract_paragraph_outline_level(paragraph)
        in_table = _is_table_paragraph(paragraph)
        has_drawing = _paragraph_has_drawing(paragraph)
        if not is_blank and not in_table:
            content_index += 1

        explicit_role = None
        if style_id:
            explicit_role = id_map.get(style_id.lower())
        if explicit_role is None and style_name:
            explicit_role = name_map.get(style_name.strip().lower())

        if (
            active_section
            and active_section["rule"].body_range == BodyRangeRule.UNTIL_BLANK
            and is_blank
            and not has_drawing
        ):
            active_section = None
            _prune_pending(paragraph_index)
            continue

        if abstract_en_mode and is_blank and not has_drawing:
            abstract_en_mode = False
            _prune_pending(paragraph_index)
            continue
        if abstract_mode and is_blank and not has_drawing:
            abstract_mode = False
            _prune_pending(paragraph_index)
            continue
        if reference_mode and is_blank and not has_drawing:
            _prune_pending(paragraph_index)
            continue
        has_math, has_formula_number, has_non_number_text = _extract_formula_flags(
            text,
            paragraph,
        )
        if has_math:
            if has_non_number_text:
                _record("formula_inline", style_id, paragraph_index, "stack", paragraph)
            else:
                _record("formula_block", style_id, paragraph_index, "stack", paragraph)
            if has_formula_number:
                _record(
                    "formula_number",
                    style_id,
                    paragraph_index,
                    "stack",
                    paragraph,
                    True,
                    allow_missing_style=True,
                )
            for run in paragraph.runs:
                script = _extract_run_script(run)
                if script == "superscript":
                    _register_run_style(
                        role_candidates,
                        "superscript",
                        run,
                        paragraph_index,
                        style_id,
                    )
                elif script == "subscript":
                    _register_run_style(
                        role_candidates,
                        "subscript",
                        run,
                        paragraph_index,
                        style_id,
                    )
            if not has_non_number_text:
                _prune_pending(paragraph_index)
                continue
        if in_table:
            if not is_blank:
                _record("table_body", style_id, paragraph_index, "stack", paragraph)
            _prune_pending(paragraph_index)
            continue

        if has_drawing:
            _record("figure_body", style_id, paragraph_index, "stack", paragraph)
            if is_blank:
                _prune_pending(paragraph_index)
                continue

        if toc_mode and not is_blank:
            if _is_toc_line(style_name) or _looks_like_toc_entry(text):
                toc_level = _record_toc_level(style_name, text)
                if toc_level is None:
                    role = "toc_body"
                else:
                    role = f"toc_body_L{toc_level}"
                _record(role, style_id, paragraph_index, "stack", paragraph)
                _prune_pending(paragraph_index)
                continue
            toc_mode = False

        matched_section = _match_section_title(text.strip(), style_name, content_index)
        matched_content = (
            _match_section_content(text.strip(), content_index)
            if not is_blank and matched_section is None
            else None
        )
        if matched_content is not None:
            if matched_content.key == _COVER_SECTION_KEY:
                if enable_cover_detection and (
                    active_section is None
                    or matched_content.key != active_section["rule"].key
                ):
                    cover_detected = True
                    _commit_cover_title()
                    _record("cover_info", style_id, paragraph_index, "stack", paragraph, True)
                    active_section = _next_active_section(matched_content)
                    _prune_pending(paragraph_index)
                    continue
            elif active_section is None or matched_content.key != active_section["rule"].key:
                _record(
                    _section_body_role(matched_content),
                    style_id,
                    paragraph_index,
                    "stack",
                    paragraph,
                )
                active_section = _next_active_section(matched_content)
                _prune_pending(paragraph_index)
                continue

        if (
            document_title_found
            and not document_title_en_found
            and not is_blank
            and document_title_index is not None
            and paragraph_index - document_title_index <= 6
        and _is_document_title_en_candidate(text, paragraph, WD_ALIGN_PARAGRAPH, exclude_abstract=True)
        ):
            _record("document_title_en", style_id, paragraph_index, "stack", paragraph)
            document_title_en_found = True
            _prune_pending(paragraph_index)
            continue

        if (
            not document_title_found
            and not is_blank
            and _is_document_title_candidate(paragraph, text, WD_ALIGN_PARAGRAPH)
        ):
            if cover_detected and enable_cover_detection:
                if not cover_title_recorded:
                    _record("cover_title", style_id, paragraph_index, "stack", paragraph)
                    cover_title_recorded = True
                _prune_pending(paragraph_index)
                continue
            if cover_detected:
                _prune_pending(paragraph_index)
                continue
            else:
                if pending_cover_title is None:
                    pending_cover_title = (style_id, paragraph_index, paragraph)
            _prune_pending(paragraph_index)
            continue

        inline_abstract_role = _match_inline_abstract_role(text)
        if inline_abstract_role and not is_blank:
            _record(inline_abstract_role, style_id, paragraph_index, "stack", paragraph)
            body_role = (
                "abstract_body"
                if inline_abstract_role == "abstract_title"
                else "abstract_en_body"
            )
            _record(body_role, style_id, paragraph_index, "stack", paragraph)
            if inline_abstract_role == "abstract_title":
                abstract_mode = True
                abstract_en_mode = False
                reference_mode = False
                title_stack = []
                toc_mode = False
            else:
                abstract_en_mode = True
                abstract_mode = False
                reference_mode = False
                title_stack = []
                toc_mode = False
            _prune_pending(paragraph_index)
            continue

        if (
            active_section is not None
            and matched_section is not None
            and matched_section.key == active_section["rule"].key
        ):
            matched_section = None
        if matched_section is not None and not is_blank:
            if (
                matched_section.key == _COVER_SECTION_KEY
                and enable_cover_detection
            ):
                cover_detected = True
                _commit_cover_title()
                _prune_pending(paragraph_index)
                continue
            if active_section is not None:
                active_section = None
            _record(
                _section_title_role(matched_section),
                style_id,
                paragraph_index,
                "stack",
                paragraph,
            )
            active_section = {
                "rule": matched_section,
                "body_role": _section_body_role(matched_section),
                "remaining": matched_section.body_paragraph_limit or 0,
            }
            _prune_pending(paragraph_index)
            continue

        title_role: str | None = None
        explicit_group = _resolve_role_group(explicit_role) if explicit_role else None
        element = getattr(paragraph, "_element", None)
        caption_role = (
            caption_roles_by_id.get(id(element)) if element is not None else None
        )
        if explicit_role and explicit_group in {"title", "special_title", "caption"}:
            title_role = explicit_role
        else:
            if caption_role is not None:
                title_role = caption_role
            else:
                text_role = _match_special_role_by_text_value(text)
                if text_role in {"figure_caption", "table_caption"}:
                    if text_role == "figure_caption" and caption_object_presence.get("figure"):
                        text_role = None
                    if text_role == "table_caption" and caption_object_presence.get("table"):
                        text_role = None
                if text_role:
                    title_role = text_role
            if title_role is None and outline_level_max is not None and outline_level_min is not None:
                if outline_level is not None and outline_level <= outline_level_max:
                    normalized_level = outline_level - outline_level_min + 1
                    if normalized_level > 0 and (
                        max_heading_level is None or normalized_level <= max_heading_level
                    ):
                        title_role = f"title_L{normalized_level}"
            if title_role is None and style_name:
                special = _match_special_role_by_style_name(style_name)
                if special:
                    title_role = special
                else:
                    level = _parse_heading_level_from_name(style_name)
                    if level is not None and (
                        max_heading_level is None or level <= max_heading_level
                    ):
                        title_role = f"title_L{level}"
            if title_role is None:
                heading_role = _match_title_role_by_text_value(text)
                if heading_role:
                    level = _extract_title_level_from_role(heading_role)
                    if level is None or (
                        max_heading_level is None or level <= max_heading_level
                    ):
                        title_role = heading_role

        if active_section is not None and not is_blank:
            rule = active_section["rule"]
            if rule.body_range == BodyRangeRule.FIXED_PARAGRAPHS:
                _record(active_section["body_role"], style_id, paragraph_index, "stack", paragraph)
                remaining = int(active_section.get("remaining", 0)) - 1
                if remaining <= 0:
                    active_section = None
                else:
                    active_section["remaining"] = remaining
                _prune_pending(paragraph_index)
                continue
            if rule.body_range == BodyRangeRule.UNTIL_NEXT_TITLE:
                if title_role is None:
                    _record(active_section["body_role"], style_id, paragraph_index, "stack", paragraph)
                    _prune_pending(paragraph_index)
                    continue
                active_section = None
            elif rule.body_range == BodyRangeRule.UNTIL_BLANK:
                _record(active_section["body_role"], style_id, paragraph_index, "stack", paragraph)
                _prune_pending(paragraph_index)
                continue

        if explicit_role and explicit_group in {"body", "special_body", "note"}:
            if not is_blank:
                _record(explicit_role, style_id, paragraph_index, "explicit", paragraph)
            _prune_pending(paragraph_index)
            continue

        if title_role and pending_cover_title is not None and not cover_detected:
            _commit_document_title_from_pending()

        if title_role:
            if not is_blank:
                _record_title_spacing(title_role, paragraph)
                source = "explicit" if explicit_role else "stack"
                _record(title_role, style_id, paragraph_index, source, paragraph)
            if title_role == "abstract_title":
                abstract_mode = True
                abstract_en_mode = False
                reference_mode = False
                title_stack = []
                toc_mode = False
            elif title_role == "abstract_en_title":
                abstract_en_mode = True
                abstract_mode = False
                reference_mode = False
                title_stack = []
                toc_mode = False
            elif title_role == "reference_title":
                reference_mode = True
                abstract_mode = False
                abstract_en_mode = False
                title_stack = []
                toc_mode = False
            elif title_role == "toc_title":
                toc_mode = True
                abstract_mode = False
                abstract_en_mode = False
                reference_mode = False
                title_stack = []
            elif title_role.startswith("title_L"):
                abstract_mode = False
                abstract_en_mode = False
                reference_mode = False
                toc_mode = False
                level = _extract_title_level_from_role(title_role) or 1
                while title_stack and title_stack[-1] >= level:
                    title_stack.pop()
                title_stack.append(level)
            elif title_role in {"figure_caption", "table_caption"}:
                abstract_mode = False
                reference_mode = False
                toc_mode = False
                pending_captions.append(
                    {"type": "figure" if title_role == "figure_caption" else "table", "index": paragraph_index}
                )
            _prune_pending(paragraph_index)
            continue

        if not is_blank and _is_keyword_line(text):
            _record("keyword_line", style_id, paragraph_index, "stack", paragraph, True)
            _prune_pending(paragraph_index)
            continue

        if abstract_en_mode:
            if not is_blank:
                _record("abstract_en_body", style_id, paragraph_index, "stack", paragraph)
            _prune_pending(paragraph_index)
            continue
        if abstract_mode:
            if not is_blank:
                _record("abstract_body", style_id, paragraph_index, "stack", paragraph)
            _prune_pending(paragraph_index)
            continue
        if reference_mode:
            if not is_blank:
                _record("reference_body", style_id, paragraph_index, "stack", paragraph)
            _prune_pending(paragraph_index)
            continue
        if pending_captions and _is_note_paragraph(text, style_name):
            candidates = [
                item for item in pending_captions if 1 <= paragraph_index - int(item["index"]) <= 2
            ]
            if candidates:
                chosen = min(
                    candidates,
                    key=lambda item: (paragraph_index - int(item["index"]), -int(item["index"])),
                )
                note_role = "figure_note" if chosen["type"] == "figure" else "table_note"
                _record(note_role, style_id, paragraph_index, "stack", paragraph)
                pending_captions.remove(chosen)
                _prune_pending(paragraph_index)
                continue

        if not is_blank:
            if title_stack:
                level = title_stack[-1]
                _record(f"body_L{level}", style_id, paragraph_index, "stack", paragraph)
            else:
                _record(_GLOBAL_BODY_CANDIDATE_ROLE, style_id, paragraph_index, "global", paragraph)

        for run in paragraph.runs:
            script = _extract_run_script(run)
            if script == "superscript":
                _register_run_style(
                    role_candidates,
                    "superscript",
                    run,
                    paragraph_index,
                    style_id,
                )
            elif script == "subscript":
                _register_run_style(
                    role_candidates,
                    "subscript",
                    run,
                    paragraph_index,
                    style_id,
                )

        _prune_pending(paragraph_index)

    if pending_cover_title is not None:
        if cover_detected and not cover_title_recorded:
            _commit_cover_title()
        elif not document_title_found:
            _commit_document_title_from_pending()

    return role_candidates


def _iter_paragraphs(document: Any) -> Iterable[Any]:
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table: Any) -> Iterable[Any]:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _iter_block_items(document: Any) -> Iterable[tuple[str, Any]]:
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    parent = document
    for child in parent.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield ("paragraph", Paragraph(child, parent))
        elif isinstance(child, CT_Tbl):
            yield ("table", Table(child, parent))


def _add_sample(
    stats: SampleStats,
    field: str,
    value: object | None,
    weight: float,
    index: int,
) -> int:
    if value is None:
        return index
    weight_map = stats.weights.setdefault(field, {})
    if value in weight_map:
        current_weight, current_index = weight_map[value]
        weight_map[value] = (current_weight + weight, current_index)
    else:
        index += 1
        weight_map[value] = (weight, index)
    return index


def _derive_role_stats(
    stats: SampleStats | None,
    count: int,
    first_index: int | None,
) -> SampleStats | None:
    if stats is None:
        return None
    return SampleStats(
        count=count if count > 0 else stats.count,
        first_index=first_index if first_index is not None else stats.first_index,
        outline_min=stats.outline_min,
        style_name=stats.style_name,
        weights=stats.weights,
        text_samples=list(stats.text_samples),
    )


def _sample_mode(stats: SampleStats, field: str) -> Any | None:
    weight_map = stats.weights.get(field)
    if not weight_map:
        return None
    best_value = None
    best_weight = -1.0
    best_index = 1_000_000_000
    for value, (weight, first_index) in weight_map.items():
        if weight > best_weight or (weight == best_weight and first_index < best_index):
            best_value = value
            best_weight = weight
            best_index = first_index
    return best_value


def _extract_paragraph_alignment(paragraph: Any, align_enum: Any) -> str | None:
    alignment = paragraph.alignment
    if alignment is None:
        style = getattr(paragraph, "style", None)
        paragraph_format = getattr(style, "paragraph_format", None) if style else None
        alignment = getattr(paragraph_format, "alignment", None)
    if alignment is None:
        return None
    if alignment == align_enum.LEFT:
        return "LEFT"
    if alignment == align_enum.CENTER:
        return "CENTER"
    if alignment == align_enum.RIGHT:
        return "RIGHT"
    if alignment in {align_enum.JUSTIFY, align_enum.DISTRIBUTE}:
        return "JUSTIFY"
    return "JUSTIFY"


def _extract_paragraph_line_spacing(
    paragraph: Any,
    line_enum: Any,
    log_state: ParseLogState | None,
    style_id: str,
    paragraph_index: int,
) -> tuple[str, float, str] | None:
    paragraph_format = paragraph.paragraph_format
    rule = paragraph_format.line_spacing_rule
    value = paragraph_format.line_spacing
    if rule is not None:
        if rule == line_enum.SINGLE:
            return "SINGLE", 1.0, "MULTIPLE"
        if rule == line_enum.ONE_POINT_FIVE:
            return "ONE_POINT_FIVE", 1.5, "MULTIPLE"
        if rule == line_enum.DOUBLE:
            return "DOUBLE", 2.0, "MULTIPLE"
        if rule == line_enum.MULTIPLE and value is not None:
            return "MULTIPLE", float(value), "MULTIPLE"
        if rule == line_enum.EXACTLY and value is not None:
            return "EXACTLY", _length_to_pt(value), "PT"
        if rule == line_enum.AT_LEAST and value is not None:
            return "AT_LEAST", _length_to_pt(value), "PT"
        return None
    if value is None:
        return None
    if isinstance(value, float):
        return "MULTIPLE", float(value), "MULTIPLE"
    line_rule, line_twips = _read_paragraph_line_rule(paragraph)
    if line_rule in {"exact", "atLeast"} and line_twips is not None:
        resolved = "EXACTLY" if line_rule == "exact" else "AT_LEAST"
        return resolved, line_twips / 20.0, "PT"
    if line_rule in {None, "auto"} and line_twips is not None:
        return "MULTIPLE", line_twips / 240.0, "MULTIPLE"
    inferred = _length_to_pt(value)
    if log_state is not None:
        _warn(
            log_state,
            rule="line_spacing",
            reason="inferred EXACTLY due to missing lineRule",
            style_id=style_id,
            paragraph_index=paragraph_index,
        )
    return "EXACTLY", inferred, "PT"


def _extract_paragraph_spacing_detail(
    paragraph: Any,
) -> tuple[
    float | None,
    float | None,
    float | None,
    str | None,
    float | None,
    str | None,
]:
    paragraph_format = paragraph.paragraph_format
    space_before = paragraph_format.space_before
    space_after = paragraph_format.space_after
    space_before_pt = _length_to_pt(space_before)
    space_after_pt = _length_to_pt(space_after)
    before_value = space_before_pt
    before_unit = "PT" if space_before_pt is not None else None
    after_value = space_after_pt
    after_unit = "PT" if space_after_pt is not None else None

    p_pr = getattr(paragraph._element, "pPr", None)
    if p_pr is not None:
        spacing = p_pr.find("w:spacing", namespaces=_PPR_NS)
        if spacing is not None:
            before_lines = _parse_spacing_lines(spacing.get(_attr_name("beforeLines")))
            after_lines = _parse_spacing_lines(spacing.get(_attr_name("afterLines")))
            if before_lines is not None:
                before_value = before_lines
                before_unit = "LINE"
            if after_lines is not None:
                after_value = after_lines
                after_unit = "LINE"
    return (
        space_before_pt,
        space_after_pt,
        before_value,
        before_unit,
        after_value,
        after_unit,
    )


def _extract_paragraph_spacing(paragraph: Any) -> tuple[float | None, float | None]:
    space_before_pt, space_after_pt, *_ = _extract_paragraph_spacing_detail(paragraph)
    return space_before_pt, space_after_pt


def _extract_paragraph_indentation(
    paragraph: Any,
) -> tuple[float | None, float | None, float | None, float | None]:
    paragraph_format = paragraph.paragraph_format
    left_indent = _length_to_pt(paragraph_format.left_indent)
    right_indent = _length_to_pt(paragraph_format.right_indent)
    first_line = _length_to_pt(paragraph_format.first_line_indent)
    hanging = None
    if first_line is not None and first_line < 0:
        hanging = abs(first_line)
        first_line = 0.0
    return left_indent, right_indent, first_line, hanging


def _length_to_pt(value: Any | None) -> float | None:
    if value is None:
        return None
    if hasattr(value, "pt"):
        return float(value.pt)
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _parse_spacing_lines(value: str | None) -> float | None:
    if value is None:
        return None
    try:
        return float(value) / 100.0
    except ValueError:
        return None


def _attr_name(name: str) -> str:
    return f"{{{_W_NS}}}{name}"


def _parse_int(value: str | None) -> int | None:
    if value is None:
        return None
    try:
        return int(value)
    except ValueError:
        return None


def _parse_int_from_object(value: object) -> int | None:
    if value is None:
        return None
    if isinstance(value, int):
        return value
    if isinstance(value, float):
        return int(value)
    if isinstance(value, str):
        return _parse_int(value)
    return None


def _read_paragraph_line_rule(paragraph: Any) -> tuple[str | None, int | None]:
    p_pr = getattr(paragraph._element, "pPr", None)
    if p_pr is None:
        return None, None
    spacing = p_pr.find("w:spacing", namespaces=_PPR_NS)
    if spacing is None:
        return None, None
    line_rule = spacing.get(_attr_name("lineRule"))
    line_val = spacing.get(_attr_name("line"))
    return line_rule, _parse_int(line_val)


def _extract_run_fonts(run: Any) -> tuple[str | None, str | None, str | None, str | None]:
    font_name = None
    ascii_name = None
    hansi_name = None
    east_asia_name = None
    r_pr = getattr(run._element, "rPr", None)
    if r_pr is not None:
        r_fonts = r_pr.find("w:rFonts", namespaces=_PPR_NS)
        if r_fonts is not None:
            ascii_name = r_fonts.get(_attr_name("ascii"))
            hansi_name = r_fonts.get(_attr_name("hAnsi"))
            east_asia_name = r_fonts.get(_attr_name("eastAsia"))
            font_name = east_asia_name or hansi_name or ascii_name
    if font_name is None:
        font_name = run.font.name
    return font_name, ascii_name, hansi_name, east_asia_name


def _extract_run_font_size(run: Any) -> float | None:
    size = run.font.size
    if size is None:
        return None
    return _length_to_pt(size)


def _extract_run_bold(run: Any) -> bool | None:
    if run.bold is not None:
        return run.bold
    return run.font.bold


def _extract_run_script(run: Any) -> str | None:
    r_pr = getattr(run._element, "rPr", None)
    if r_pr is not None:
        vert = r_pr.find("w:vertAlign", namespaces=_PPR_NS)
        if vert is not None:
            val = vert.get(_attr_name("val"))
            if val in {"superscript", "subscript"}:
                return val
    font = getattr(run, "font", None)
    if font is not None:
        if getattr(font, "superscript", None):
            return "superscript"
        if getattr(font, "subscript", None):
            return "subscript"
    return None


def _register_run_style(
    role_candidates: dict[str, dict[str, dict[str, object]]],
    role: str,
    run: Any,
    paragraph_index: int,
    style_id: str | None,
) -> None:
    role_key = style_id or ""
    entry = role_candidates.setdefault(role, {}).get(role_key)
    if entry is None:
        entry = {
            "count": 0,
            "first_index": paragraph_index,
            "source": "run",
            "stats": SampleStats(),
            "sample_index": 0,
        }
        role_candidates.setdefault(role, {})[role_key] = entry
    entry["count"] = int(entry["count"]) + 1
    if entry.get("first_index") is None or paragraph_index < int(entry["first_index"]):
        entry["first_index"] = paragraph_index
    stats = entry["stats"]
    sample_index = int(entry.get("sample_index", 0))
    font_name, font_ascii, font_hansi, font_east_asia = _extract_run_fonts(run)
    sample_index = _add_sample(stats, "font_name", font_name, 1.0, sample_index)
    sample_index = _add_sample(stats, "font_name_ascii", font_ascii, 1.0, sample_index)
    sample_index = _add_sample(stats, "font_name_hAnsi", font_hansi, 1.0, sample_index)
    sample_index = _add_sample(
        stats,
        "font_name_eastAsia",
        font_east_asia,
        1.0,
        sample_index,
    )
    font_size = _extract_run_font_size(run)
    sample_index = _add_sample(stats, "font_size_pt", font_size, 1.0, sample_index)
    bold = _extract_run_bold(run)
    sample_index = _add_sample(stats, "bold", bold, 1.0, sample_index)
    stats.count += 1
    entry["sample_index"] = sample_index


def _extract_paragraph_outline_level(paragraph: Any) -> int | None:
    p_pr = getattr(paragraph._element, "pPr", None)
    if p_pr is None:
        return None
    outline_elem = p_pr.find("w:outlineLvl", namespaces=_PPR_NS)
    if outline_elem is None:
        return None
    level = _parse_int(outline_elem.get(_attr_name("val")))
    if level is None:
        return None
    if level >= 9:
        return None
    return level


def _default_page_margins() -> dict[str, object]:
    return {
        "sections": [],
        "summary": {},
    }


def _default_table_borders() -> dict[str, object]:
    return {
        "summary": {},
        "tables": [],
    }


def _default_toc_levels() -> dict[str, object]:
    return {
        "levels": [],
        "styles": {},
    }


def _default_page_margin_summary() -> dict[str, object]:
    return {}


def _twips_to_pt(value: int | None) -> float | None:
    if value is None:
        return None
    return value / 20.0


def _parse_pg_mar(sect_pr: etree._Element) -> dict[str, float]:
    pg_mar = sect_pr.find("w:pgMar", namespaces=_PPR_NS)
    if pg_mar is None:
        return {}
    margins: dict[str, float] = {}
    for key in _PAGE_MARGIN_KEYS:
        raw = _parse_int(pg_mar.get(_attr_name(key)))
        pt = _twips_to_pt(raw)
        if pt is not None:
            margins[key] = pt
    return margins


def _extract_paragraph_text(paragraph: etree._Element) -> str:
    texts: list[str] = []
    for node in paragraph.iterfind(".//w:t", namespaces=_PPR_NS):
        if node.text:
            texts.append(node.text)
    return "".join(texts).strip()


def _contains_any_keyword(text: str, keywords: Iterable[str]) -> bool:
    if not text:
        return False
    clean = text.strip().lower()
    if not clean:
        return False
    for keyword in keywords:
        if not keyword:
            continue
        kw = keyword.strip().lower()
        if not kw:
            continue
        if clean == kw or kw in clean:
            return True
        clean_compact = re.sub(r"\s+", "", clean)
        kw_compact = re.sub(r"\s+", "", kw)
        if clean_compact and kw_compact:
            if clean_compact == kw_compact or kw_compact in clean_compact:
                return True
    return False


def _build_logical_part_keywords(
    section_rules: Iterable[SectionRule] | None,
) -> tuple[set[str], set[str], set[str]]:
    cover_keywords: set[str] = set()
    statement_keywords: set[str] = set()
    back_keywords: set[str] = set()
    if section_rules:
        for rule in section_rules:
            display = rule.display_name or ""
            key = rule.key or ""
            if rule.key == _COVER_SECTION_KEY:
                cover_keywords.update(rule.title_keywords)
                cover_keywords.update(rule.content_keywords)
            if rule.position in {SectionPosition.FRONT, SectionPosition.FIRST_PAGE}:
                if "statement" in key.lower() or "\u58f0\u660e" in display:
                    statement_keywords.update(rule.title_keywords)
            if rule.position in {SectionPosition.BACK, SectionPosition.LAST_PAGE}:
                back_keywords.update(rule.title_keywords)
    if not statement_keywords:
        statement_keywords.update(_DEFAULT_STATEMENT_KEYWORDS)
    if not back_keywords:
        back_keywords.update(_DEFAULT_BACK_KEYWORDS)
    return cover_keywords, statement_keywords, back_keywords


def _classify_section_part(
    section: dict[str, object],
    cover_markers: set[int],
    statement_markers: set[int],
    main_markers: set[int],
    back_markers: set[int],
    first_main_index: int | None,
    first_back_index: int | None,
) -> str:
    start = int(section.get("start", 0))
    end = int(section.get("end", 0))
    if _section_has_marker(start, end, statement_markers):
        return "statement"
    if _section_has_marker(start, end, back_markers):
        return "back"
    if _section_has_marker(start, end, cover_markers):
        return "cover"
    if _section_has_marker(start, end, main_markers):
        return "main"
    if first_back_index is not None and start >= first_back_index:
        return "back"
    if first_main_index is not None:
        if end < first_main_index:
            return "cover"
        if start >= first_main_index:
            return "main"
    return "unknown"


def _section_has_marker(start: int, end: int, markers: set[int]) -> bool:
    if start <= 0 or end <= 0 or not markers:
        return False
    return any(start <= index <= end for index in markers)


def _build_page_margin_summary(sections: list[dict[str, object]]) -> dict[str, object]:
    summary: dict[str, object] = {"sections": len(sections)}
    if not sections:
        return summary
    signatures: set[tuple[float | None, ...]] = set()
    for section in sections:
        margins = section.get("margins")
        if not isinstance(margins, dict):
            margins = {}
        signature = tuple(margins.get(key) for key in _PAGE_MARGIN_KEYS)
        signatures.add(signature)
        logical_part = section.get("logical_part")
        if isinstance(logical_part, str) and logical_part:
            summary[logical_part] = int(summary.get(logical_part, 0)) + 1
    summary["distinct_margins"] = len(signatures)
    return summary


def _parse_page_margins(
    template_path: Path,
    section_rules: Iterable[SectionRule] | None,
    log_state: ParseLogState | None,
    enable_cover_detection: bool = True,
) -> dict[str, object]:
    try:
        with ZipFile(template_path) as archive:
            document_bytes = archive.read("word/document.xml")
    except Exception as exc:
        if log_state is not None:
            _warn(log_state, rule="page_margins", reason=f"missing document.xml ({exc})")
        return _default_page_margins()
    try:
        root = etree.fromstring(document_bytes)
    except Exception as exc:
        if log_state is not None:
            _warn(log_state, rule="page_margins", reason=f"parse document.xml failed ({exc})")
        return _default_page_margins()
    body = root.find("w:body", namespaces=_PPR_NS)
    if body is None:
        return _default_page_margins()

    cover_keywords, statement_keywords, back_keywords = _build_logical_part_keywords(section_rules)
    paragraphs = list(body.iterfind(".//w:p", namespaces=_PPR_NS))
    paragraph_texts: list[tuple[int, str]] = []
    sections: list[dict[str, object]] = []
    section_start = 1
    paragraph_index = 0
    for paragraph in paragraphs:
        paragraph_index += 1
        text = _extract_paragraph_text(paragraph)
        if text:
            paragraph_texts.append((paragraph_index, text))
        sect_pr = paragraph.find("w:pPr/w:sectPr", namespaces=_PPR_NS)
        if sect_pr is not None:
            margins = _parse_pg_mar(sect_pr)
            sections.append(
                {
                    "index": len(sections) + 1,
                    "start": section_start,
                    "end": paragraph_index,
                    "margins": margins,
                }
            )
            section_start = paragraph_index + 1

    body_sect_pr = body.find("w:sectPr", namespaces=_PPR_NS)
    if body_sect_pr is not None and section_start <= paragraph_index:
        margins = _parse_pg_mar(body_sect_pr)
        sections.append(
            {
                "index": len(sections) + 1,
                "start": section_start,
                "end": paragraph_index,
                "margins": margins,
            }
        )

    if not sections:
        return _default_page_margins()

    cover_markers: set[int] = set()
    statement_markers: set[int] = set()
    back_markers: set[int] = set()
    main_markers: set[int] = set()
    for index, text in paragraph_texts:
        if enable_cover_detection and _contains_any_keyword(text, cover_keywords):
            cover_markers.add(index)
        if _TEXT_REFERENCE_PATTERN.match(text) or _contains_any_keyword(text, back_keywords):
            back_markers.add(index)
        if _contains_any_keyword(text, statement_keywords):
            statement_markers.add(index)
        if _match_title_role_by_text_value(text) is not None:
            main_markers.add(index)

    first_main_index = min(main_markers) if main_markers else None
    first_back_index = min(back_markers) if back_markers else None

    for section in sections:
        logical_part = _classify_section_part(
            section,
            cover_markers,
            statement_markers,
            main_markers,
            back_markers,
            first_main_index,
            first_back_index,
        )
        section["logical_part"] = logical_part

    summary = _build_page_margin_summary(sections)
    return {
        "sections": sections,
        "summary": summary or _default_page_margin_summary(),
    }


def _read_xml_from_docx(
    template_path: Path,
    part_name: str,
    log_state: ParseLogState | None,
    rule: str,
) -> bytes | None:
    try:
        with ZipFile(template_path) as archive:
            return archive.read(part_name)
    except Exception as exc:
        if log_state is not None:
            _warn(log_state, rule=rule, reason=f"missing {part_name} ({exc})")
        return None


def _parse_table_borders(
    template_path: Path,
    log_state: ParseLogState | None,
) -> dict[str, object]:
    document_bytes = _read_xml_from_docx(
        template_path,
        "word/document.xml",
        log_state,
        rule="table_borders",
    )
    if not document_bytes:
        return _default_table_borders()
    styles_bytes = _read_xml_from_docx(
        template_path,
        "word/styles.xml",
        log_state,
        rule="table_borders",
    )
    style_borders, style_names = _parse_table_style_borders(styles_bytes, log_state)
    try:
        root = etree.fromstring(document_bytes)
    except Exception as exc:
        if log_state is not None:
            _warn(log_state, rule="table_borders", reason=f"parse document.xml failed ({exc})")
        return _default_table_borders()
    tables = root.findall(".//w:tbl", namespaces=_PPR_NS)
    if not tables:
        return _default_table_borders()

    summary: dict[str, int] = {}
    table_items: list[dict[str, object]] = []
    for index, table in enumerate(tables, start=1):
        borders = _read_tbl_borders(table)
        style_id = _read_tbl_style_id(table)
        style_name = style_names.get(style_id) if style_id else None
        if not _has_any_border(borders) and style_id and style_id in style_borders:
            borders = style_borders[style_id]
        pattern = _classify_table_border_pattern(borders)
        summary[pattern] = summary.get(pattern, 0) + 1
        table_items.append(
            {
                "index": index,
                "style_id": style_id,
                "style_name": style_name,
                "pattern": pattern,
                "borders": borders,
            }
        )
    return {
        "summary": summary,
        "tables": table_items,
    }


def _parse_footnote_numbering(
    template_path: Path,
    log_state: ParseLogState | None,
) -> dict[str, object]:
    settings_bytes = _read_xml_from_docx(
        template_path,
        "word/settings.xml",
        log_state,
        rule="footnote_numbering",
    )
    if not settings_bytes:
        return {}
    try:
        root = etree.fromstring(settings_bytes)
    except Exception as exc:
        if log_state is not None:
            _warn(log_state, rule="footnote_numbering", reason=f"parse settings.xml failed ({exc})")
        return {}
    footnote_pr = root.find(".//w:footnotePr", namespaces=_PPR_NS)
    if footnote_pr is None:
        return {}
    num_fmt = footnote_pr.find("w:numFmt", namespaces=_PPR_NS)
    num_start = footnote_pr.find("w:numStart", namespaces=_PPR_NS)
    num_restart = footnote_pr.find("w:numRestart", namespaces=_PPR_NS)
    return {
        "format": num_fmt.get(_attr_name("val")) if num_fmt is not None else None,
        "start": _parse_int_from_object(num_start.get(_attr_name("val"))) if num_start is not None else None,
        "restart": num_restart.get(_attr_name("val")) if num_restart is not None else None,
    }


def _collect_footnote_samples(
    template_path: Path,
    log_state: ParseLogState | None,
) -> tuple[SampleStats | None, SampleStats | None, set[str], set[str]]:
    footnotes_bytes = _read_xml_from_docx(
        template_path,
        "word/footnotes.xml",
        log_state,
        rule="footnotes",
    )
    if not footnotes_bytes:
        return None, None, set(), set()
    try:
        root = etree.fromstring(footnotes_bytes)
    except Exception as exc:
        if log_state is not None:
            _warn(log_state, rule="footnotes", reason=f"parse footnotes.xml failed ({exc})")
        return None, None, set(), set()

    text_stats = SampleStats()
    ref_stats = SampleStats()
    text_style_ids: set[str] = set()
    ref_style_ids: set[str] = set()
    text_index = 0
    ref_index = 0

    for footnote in root.findall("w:footnote", namespaces=_PPR_NS):
        footnote_type = footnote.get(_attr_name("type"))
        if footnote_type in {"separator", "continuationSeparator"}:
            continue
        for paragraph in footnote.findall("w:p", namespaces=_PPR_NS):
            p_style = paragraph.find("w:pPr/w:pStyle", namespaces=_PPR_NS)
            if p_style is not None:
                style_val = p_style.get(_attr_name("val"))
                if style_val:
                    text_style_ids.add(style_val)
            for run in paragraph.findall("w:r", namespaces=_PPR_NS):
                r_style = run.find("w:rPr/w:rStyle", namespaces=_PPR_NS)
                style_val = r_style.get(_attr_name("val")) if r_style is not None else None
                if style_val:
                    ref_style_ids.add(style_val)
                if _run_has_footnote_ref(run) or _is_reference_style_id(style_val):
                    ref_index = _add_run_samples(ref_stats, run, ref_index)
                    continue
                text_index = _add_run_samples(text_stats, run, text_index)
    if text_stats.count <= 0:
        text_stats = None
    if ref_stats.count <= 0:
        ref_stats = None
    return text_stats, ref_stats, text_style_ids, ref_style_ids


def _parse_header_footer(
    template_path: Path,
    styles: dict[str, Any],
    defaults: StyleDefaults,
    theme_map: dict[str, str],
    page_margins: dict[str, object] | None,
    log_state: ParseLogState | None,
) -> dict[str, object]:
    try:
        from docx import Document
    except ImportError:
        return {}
    try:
        document = Document(str(template_path))
    except Exception as exc:
        if log_state is not None:
            _warn(log_state, rule="header_footer", reason=f"load document failed ({exc})")
        return {}

    margin_sections = []
    if isinstance(page_margins, dict):
        sections = page_margins.get("sections")
        if isinstance(sections, list):
            margin_sections = sections

    sections_meta: list[dict[str, object]] = []
    header_sections = 0
    footer_sections = 0
    headers_total = 0
    footers_total = 0

    for idx, section in enumerate(getattr(document, "sections", []), start=1):
        logical_part = None
        if idx - 1 < len(margin_sections):
            logical_part = margin_sections[idx - 1].get("logical_part")
        entry: dict[str, object] = {
            "index": idx,
            "logical_part": logical_part,
            "headers": [],
            "footers": [],
        }
        header_items = _collect_header_footer_entries(
            section,
            kind="header",
            styles=styles,
            defaults=defaults,
            theme_map=theme_map,
            log_state=log_state,
        )
        footer_items = _collect_header_footer_entries(
            section,
            kind="footer",
            styles=styles,
            defaults=defaults,
            theme_map=theme_map,
            log_state=log_state,
        )
        if header_items:
            entry["headers"] = header_items
            header_sections += 1
            headers_total += len(header_items)
        if footer_items:
            entry["footers"] = footer_items
            footer_sections += 1
            footers_total += len(footer_items)
        if header_items or footer_items:
            sections_meta.append(entry)

    if not sections_meta:
        return {}
    return {
        "summary": {
            "sections": len(getattr(document, "sections", [])),
            "header_sections": header_sections,
            "footer_sections": footer_sections,
            "headers": headers_total,
            "footers": footers_total,
        },
        "sections": sections_meta,
    }


def _collect_header_footer_entries(
    section: Any,
    kind: str,
    styles: dict[str, Any],
    defaults: StyleDefaults,
    theme_map: dict[str, str],
    log_state: ParseLogState | None,
) -> list[dict[str, object]]:
    entries: list[dict[str, object]] = []
    variants = []
    if kind == "header":
        variants = [
            ("default", getattr(section, "header", None)),
            ("first", getattr(section, "first_page_header", None)),
            ("even", getattr(section, "even_page_header", None)),
        ]
    else:
        variants = [
            ("default", getattr(section, "footer", None)),
            ("first", getattr(section, "first_page_footer", None)),
            ("even", getattr(section, "even_page_footer", None)),
        ]
    seen = set()
    for variant, header_footer in variants:
        if header_footer is None:
            continue
        paragraphs = list(getattr(header_footer, "paragraphs", []))
        if not any(not _is_blank_paragraph(p) for p in paragraphs):
            continue
        key = id(header_footer)
        if key in seen:
            continue
        seen.add(key)
        rule = _build_header_footer_rule(
            role=f"{kind}_{variant}",
            paragraphs=paragraphs,
            styles=styles,
            defaults=defaults,
            theme_map=theme_map,
            log_state=log_state,
        )
        if rule is None:
            continue
        entries.append(
            {
                "type": variant,
                "style": rule.to_dict(include_missing_fields=False),
            }
        )
    return entries


def _build_header_footer_rule(
    role: str,
    paragraphs: list[Any],
    styles: dict[str, Any],
    defaults: StyleDefaults,
    theme_map: dict[str, str],
    log_state: ParseLogState | None,
) -> StyleRule | None:
    stats, style_counts = _collect_paragraph_samples_from_paragraphs(
        paragraphs,
        log_state=log_state,
    )
    if stats.count <= 0:
        return None
    style_id = None
    if style_counts:
        style_id = max(style_counts, key=style_counts.get)
    if style_id and style_id in styles:
        resolved = resolve_style(style_id, styles, defaults, theme_map=theme_map)
    else:
        resolved = _build_default_resolved_style(defaults, theme_map, style_id or "")
    return _build_style_rule(role, resolved, stats, prefer_samples=True)


def _collect_paragraph_samples_from_paragraphs(
    paragraphs: list[Any],
    log_state: ParseLogState | None,
) -> tuple[SampleStats, dict[str, int]]:
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    except ImportError as exc:
        raise ImportError("python-docx is required to parse header/footer samples") from exc

    stats = SampleStats()
    style_counts: dict[str, int] = {}
    paragraph_index = 0
    sample_index = 0
    for paragraph in paragraphs:
        if _is_blank_paragraph(paragraph):
            continue
        paragraph_index += 1
        stats.count += 1
        if stats.first_index is None:
            stats.first_index = paragraph_index
        paragraph_text = getattr(paragraph, "text", None)
        if paragraph_text:
            trimmed = paragraph_text.strip()
            if trimmed and trimmed not in stats.text_samples:
                if len(stats.text_samples) < 3:
                    stats.text_samples.append(trimmed)
        style = getattr(paragraph, "style", None)
        style_id = getattr(style, "style_id", None) if style is not None else None
        style_name = getattr(style, "name", None) if style is not None else None
        if style_id:
            style_counts[style_id] = style_counts.get(style_id, 0) + 1
        if stats.style_name is None and style_name:
            stats.style_name = style_name
        outline_level = _extract_paragraph_outline_level(paragraph)
        if outline_level is not None:
            if stats.outline_min is None or outline_level < stats.outline_min:
                stats.outline_min = outline_level

        alignment = _extract_paragraph_alignment(paragraph, WD_ALIGN_PARAGRAPH)
        sample_index = _add_sample(stats, "alignment", alignment, 1.0, sample_index)
        line_spacing = _extract_paragraph_line_spacing(
            paragraph,
            WD_LINE_SPACING,
            log_state,
            style_id or "",
            paragraph_index,
        )
        sample_index = _add_sample(stats, "line_spacing", line_spacing, 1.0, sample_index)
        (
            space_before_pt,
            space_after_pt,
            space_before_value,
            space_before_unit,
            space_after_value,
            space_after_unit,
        ) = _extract_paragraph_spacing_detail(paragraph)
        sample_index = _add_sample(
            stats,
            "space_before_pt",
            space_before_pt,
            1.0,
            sample_index,
        )
        sample_index = _add_sample(
            stats,
            "space_after_pt",
            space_after_pt,
            1.0,
            sample_index,
        )
        if space_before_value is not None and space_before_unit is not None:
            sample_index = _add_sample(
                stats,
                "space_before",
                (space_before_value, space_before_unit),
                1.0,
                sample_index,
            )
        if space_after_value is not None and space_after_unit is not None:
            sample_index = _add_sample(
                stats,
                "space_after",
                (space_after_value, space_after_unit),
                1.0,
                sample_index,
            )
        indent_left, indent_right, indent_first, indent_hanging = _extract_paragraph_indentation(
            paragraph,
        )
        sample_index = _add_sample(
            stats,
            "indent_left_pt",
            indent_left,
            1.0,
            sample_index,
        )
        sample_index = _add_sample(
            stats,
            "indent_right_pt",
            indent_right,
            1.0,
            sample_index,
        )
        sample_index = _add_sample(
            stats,
            "indent_first_line_pt",
            indent_first,
            1.0,
            sample_index,
        )
        sample_index = _add_sample(
            stats,
            "indent_hanging_pt",
            indent_hanging,
            1.0,
            sample_index,
        )

        for run in paragraph.runs:
            text = run.text
            if not text or text.isspace():
                continue
            weight = float(len(text.strip()))
            if weight <= 0:
                continue
            font_name, font_ascii, font_hansi, font_east_asia = _extract_run_fonts(run)
            sample_index = _add_sample(
                stats,
                "font_name",
                font_name,
                weight,
                sample_index,
            )
            sample_index = _add_sample(
                stats,
                "font_name_ascii",
                font_ascii,
                weight,
                sample_index,
            )
            sample_index = _add_sample(
                stats,
                "font_name_hAnsi",
                font_hansi,
                weight,
                sample_index,
            )
            sample_index = _add_sample(
                stats,
                "font_name_eastAsia",
                font_east_asia,
                weight,
                sample_index,
            )
            font_size = _extract_run_font_size(run)
            sample_index = _add_sample(
                stats,
                "font_size_pt",
                font_size,
                weight,
                sample_index,
            )
            bold = _extract_run_bold(run)
            sample_index = _add_sample(stats, "bold", bold, weight, sample_index)
    return stats, style_counts


def _is_reference_style_id(style_id: str | None) -> bool:
    if not style_id:
        return False
    return style_id.strip().lower() in _FOOTNOTE_REFERENCE_STYLE_IDS


def _run_has_footnote_ref(run: etree._Element) -> bool:
    return (
        run.find("w:footnoteRef", namespaces=_PPR_NS) is not None
        or run.find("w:footnoteReference", namespaces=_PPR_NS) is not None
    )


def _add_run_samples(stats: SampleStats, run: etree._Element, index: int) -> int:
    text_value = _extract_run_text(run)
    weight = float(max(len(text_value.strip()), 1))
    r_pr = run.find("w:rPr", namespaces=_PPR_NS)
    font_name, font_ascii, font_hansi, font_east_asia = _parse_run_fonts_from_xml(r_pr)
    font_size = _parse_run_size_from_xml(r_pr)
    bold = _parse_run_bold_from_xml(r_pr)
    index = _add_sample(stats, "font_name", font_name, weight, index)
    index = _add_sample(stats, "font_name_ascii", font_ascii, weight, index)
    index = _add_sample(stats, "font_name_hAnsi", font_hansi, weight, index)
    index = _add_sample(stats, "font_name_eastAsia", font_east_asia, weight, index)
    index = _add_sample(stats, "font_size_pt", font_size, weight, index)
    index = _add_sample(stats, "bold", bold, weight, index)
    stats.count += 1
    return index


def _extract_run_text(run: etree._Element) -> str:
    texts = []
    for node in run.iterfind(".//w:t", namespaces=_PPR_NS):
        if node.text:
            texts.append(node.text)
    return "".join(texts)


def _parse_run_fonts_from_xml(
    r_pr: etree._Element | None,
) -> tuple[str | None, str | None, str | None, str | None]:
    if r_pr is None:
        return None, None, None, None
    r_fonts = r_pr.find("w:rFonts", namespaces=_PPR_NS)
    if r_fonts is None:
        return None, None, None, None
    ascii_name = r_fonts.get(_attr_name("ascii"))
    hansi_name = r_fonts.get(_attr_name("hAnsi"))
    east_asia_name = r_fonts.get(_attr_name("eastAsia"))
    preferred = east_asia_name or hansi_name or ascii_name
    return preferred, ascii_name, hansi_name, east_asia_name


def _parse_run_size_from_xml(r_pr: etree._Element | None) -> float | None:
    if r_pr is None:
        return None
    sz = r_pr.find("w:sz", namespaces=_PPR_NS)
    val = sz.get(_attr_name("val")) if sz is not None else None
    size = _parse_int(val)
    if size is None:
        return None
    return size / 2.0


def _parse_run_bold_from_xml(r_pr: etree._Element | None) -> bool | None:
    if r_pr is None:
        return None
    bold_elem = r_pr.find("w:b", namespaces=_PPR_NS)
    if bold_elem is None:
        return None
    val = bold_elem.get(_attr_name("val"))
    if val is None:
        return True
    if val.lower() in {"0", "false", "off"}:
        return False
    return True


def _find_footnote_style_id(
    styles: dict[str, Any],
    role: str,
    fallback_ids: set[str],
) -> str | None:
    if role == "footnote_text":
        id_candidates = _FOOTNOTE_TEXT_STYLE_IDS
        name_candidates = _FOOTNOTE_TEXT_STYLE_NAMES
    else:
        id_candidates = _FOOTNOTE_REFERENCE_STYLE_IDS
        name_candidates = _FOOTNOTE_REFERENCE_STYLE_NAMES
    for style_id, style in styles.items():
        if style_id.strip().lower() in id_candidates:
            return style_id
        name = (style.name or "").strip().lower()
        if name and name in name_candidates:
            return style_id
    for style_id in sorted(fallback_ids):
        if style_id.strip().lower() in id_candidates:
            return style_id
    return None


def _build_default_resolved_style(
    defaults: StyleDefaults,
    theme_map: dict[str, str],
    style_id: str = "",
) -> ResolvedStyle:
    fonts = defaults.fonts.apply_theme(theme_map)
    return ResolvedStyle(
        style_id=style_id,
        name=None,
        fonts=fonts,
        font_name=fonts.preferred_name(),
        font_size_pt=defaults.font_size_pt,
        bold=defaults.bold,
        alignment=defaults.alignment,
        space_before_pt=defaults.space_before_pt,
        space_after_pt=defaults.space_after_pt,
        line_rule=defaults.line_rule,
        line_twips=defaults.line_twips,
        outline_level=defaults.outline_level,
    )


def _read_tbl_borders(table: etree._Element) -> dict[str, bool]:
    tbl_pr = table.find("w:tblPr", namespaces=_PPR_NS)
    if tbl_pr is None:
        return {}
    return _read_tbl_borders_from_pr(tbl_pr)


def _read_tbl_borders_from_pr(tbl_pr: etree._Element) -> dict[str, bool]:
    tbl_borders = tbl_pr.find("w:tblBorders", namespaces=_PPR_NS)
    if tbl_borders is None:
        return {}
    borders: dict[str, bool] = {}
    for key in ("top", "bottom", "left", "right", "insideH", "insideV"):
        elem = tbl_borders.find(f"w:{key}", namespaces=_PPR_NS)
        borders[key] = _border_present(elem)
    return borders


def _read_tbl_style_id(table: etree._Element) -> str | None:
    tbl_pr = table.find("w:tblPr", namespaces=_PPR_NS)
    if tbl_pr is None:
        return None
    tbl_style = tbl_pr.find("w:tblStyle", namespaces=_PPR_NS)
    if tbl_style is None:
        return None
    style_id = tbl_style.get(_attr_name("val"))
    if not style_id:
        return None
    return style_id


def _has_any_border(borders: dict[str, bool]) -> bool:
    return any(borders.values()) if borders else False


def _parse_table_style_borders(
    styles_bytes: bytes | None,
    log_state: ParseLogState | None,
) -> tuple[dict[str, dict[str, bool]], dict[str, str]]:
    if not styles_bytes:
        return {}, {}
    try:
        root = etree.fromstring(styles_bytes)
    except Exception as exc:
        if log_state is not None:
            _warn(log_state, rule="table_borders", reason=f"parse styles.xml failed ({exc})")
        return {}, {}
    borders_map: dict[str, dict[str, bool]] = {}
    names: dict[str, str] = {}
    for style in root.findall("w:style", namespaces=_PPR_NS):
        style_type = style.get(_attr_name("type"))
        if style_type != "table":
            continue
        style_id = style.get(_attr_name("styleId"))
        if not style_id:
            continue
        name_elem = style.find("w:name", namespaces=_PPR_NS)
        name = name_elem.get(_attr_name("val")) if name_elem is not None else None
        if name:
            names[style_id] = name
        tbl_pr = style.find("w:tblPr", namespaces=_PPR_NS)
        if tbl_pr is None:
            continue
        borders = _read_tbl_borders_from_pr(tbl_pr)
        if _has_any_border(borders):
            borders_map[style_id] = borders
    return borders_map, names


def _border_present(elem: etree._Element | None) -> bool:
    if elem is None:
        return False
    val = elem.get(_attr_name("val"))
    if val is None:
        return True
    return val not in {"nil", "none"}


def _classify_table_border_pattern(borders: dict[str, bool]) -> str:
    outer_keys = ("top", "bottom", "left", "right")
    inner_keys = ("insideH", "insideV")
    outer_any = any(borders.get(key, False) for key in outer_keys)
    inner_any = any(borders.get(key, False) for key in inner_keys)
    if not outer_any and not inner_any:
        return "none"
    outer_all = all(borders.get(key, False) for key in outer_keys)
    inner_all = all(borders.get(key, False) for key in inner_keys)
    if outer_all and inner_all:
        return "grid"
    if outer_all and not inner_any:
        return "outer_only"
    if inner_all and not outer_any:
        return "inner_only"
    return "mixed"


def _build_footnote_rules(
    styles: dict[str, Any],
    defaults: StyleDefaults,
    theme_map: dict[str, str],
    text_stats: SampleStats | None,
    ref_stats: SampleStats | None,
    text_style_ids: set[str],
    ref_style_ids: set[str],
) -> dict[str, StyleRule]:
    rules: dict[str, StyleRule] = {}
    for role, stats, style_ids in (
        ("footnote_text", text_stats, text_style_ids),
        ("footnote_reference", ref_stats, ref_style_ids),
    ):
        style_id = _find_footnote_style_id(styles, role, style_ids)
        if style_id and style_id in styles:
            resolved = resolve_style(style_id, styles, defaults, theme_map=theme_map)
        elif stats is not None and stats.count > 0:
            resolved = _build_default_resolved_style(defaults, theme_map, style_id or "")
        else:
            continue
        prefer_samples = stats is not None and stats.count > 0
        rule = _build_style_rule(role, resolved, stats, prefer_samples=prefer_samples)
        rules[role] = rule
    return rules


def _serialize_toc_levels(levels: dict[int, set[str]]) -> dict[str, object]:
    if not levels:
        return _default_toc_levels()
    sorted_levels = sorted(levels.keys())
    styles: dict[str, list[str]] = {}
    for level in sorted_levels:
        names = sorted(name for name in levels[level] if name)
        if names:
            styles[str(level)] = names
    return {
        "levels": sorted_levels,
        "styles": styles,
    }


def _apply_fallbacks(
    rules: dict[str, StyleRule],
    allow_fallback: bool,
    strict: bool,
    log_state: ParseLogState | None,
    required_roles: Iterable[str] | None = None,
    required_on_presence_map: dict[str, str] | None = None,
    global_body_rule: StyleRule | None = None,
) -> dict[str, StyleRule]:
    if strict:
        return rules
    required = _normalize_required_roles(required_roles)
    conditional_targets = set(
        _conditional_required_roles(rules, required_on_presence_map)
    )
    for role in conditional_targets:
        if role not in required:
            required.append(role)

    def _fallback_chain(role: str) -> str:
        normalized = _normalize_role(role)
        if role in conditional_targets:
            return "body_L1 -> default_body"
        if normalized.startswith("title_L") and normalized != _BASE_TITLE_ROLE:
            return "title_L1 -> default_title"
        if normalized.startswith("body_L"):
            if global_body_rule is not None:
                return "global_body -> default_body"
            return "default_body"
        if normalized == _BASE_TITLE_ROLE:
            return "default_title"
        if normalized == _BASE_BODY_ROLE:
            if global_body_rule is not None:
                return "global_body -> default_body"
            return "default_body"
        return f"default({normalized})"

    def _fill_from_fallback(role: str, target: StyleRule) -> None:
        normalized = _normalize_role(role)
        if role in conditional_targets:
            base = rules.get(_BASE_BODY_ROLE)
            if base is not None and base is not target:
                _fill_missing_fields(target, base)
            _fill_missing_fields(target, _default_rule(_BASE_BODY_ROLE))
            return
        if normalized.startswith("title_L") and normalized != _BASE_TITLE_ROLE:
            base = rules.get(_BASE_TITLE_ROLE)
            if base is not None and base is not target:
                _fill_missing_fields(target, base)
            _fill_missing_fields(target, _default_rule(_BASE_TITLE_ROLE))
            return
        if normalized.startswith("body_L"):
            if global_body_rule is not None and global_body_rule is not target:
                _fill_missing_fields(target, global_body_rule)
            _fill_missing_fields(target, _default_rule(_BASE_BODY_ROLE))
            return
        if normalized == _BASE_TITLE_ROLE:
            _fill_missing_fields(target, _default_rule(_BASE_TITLE_ROLE))
            return
        if normalized == _BASE_BODY_ROLE:
            if global_body_rule is not None and global_body_rule is not target:
                _fill_missing_fields(target, global_body_rule)
            _fill_missing_fields(target, _default_rule(_BASE_BODY_ROLE))
            return
        _fill_missing_fields(target, _default_rule(normalized))

    for role in required:
        if role not in rules:
            rules[role] = StyleRule(role=role)
        missing = _missing_required_fields(rules[role])
        if missing and log_state is not None:
            action = "fallback" if allow_fallback else "missing"
            missing_text = ", ".join(missing)
            reason = f"{missing_text} ({action})"
            if allow_fallback:
                reason = f"{missing_text} ({action} via {_fallback_chain(role)})"
            _warn(
                log_state,
                rule="missing_fields",
                reason=reason,
                role=role,
            )
        if allow_fallback:
            _fill_from_fallback(role, rules[role])

    if allow_fallback:
        for role, rule in rules.items():
            _fill_from_fallback(role, rule)
    return rules


def _default_rule(role: str) -> StyleRule:
    normalized = _normalize_role(role)
    if normalized == _BASE_TITLE_ROLE:
        return StyleRule(
            role=role,
            font_name="宋体",
            font_name_eastAsia="宋体",
            font_size_pt=16.0,
            bold=True,
            alignment="CENTER",
            line_spacing_rule="ONE_POINT_FIVE",
            line_spacing_value=1.5,
            line_spacing_unit="MULTIPLE",
            space_before_pt=12.0,
            space_after_pt=12.0,
            space_before_value=12.0,
            space_before_unit="PT",
            space_after_value=12.0,
            space_after_unit="PT",
            indent_left_pt=0.0,
            indent_right_pt=0.0,
            indent_first_line_pt=0.0,
            indent_hanging_pt=0.0,
        )
    if normalized == _BASE_BODY_ROLE:
        return StyleRule(
            role=role,
            font_name="宋体",
            font_name_eastAsia="宋体",
            font_size_pt=12.0,
            bold=False,
            alignment="JUSTIFY",
            line_spacing_rule="ONE_POINT_FIVE",
            line_spacing_value=1.5,
            line_spacing_unit="MULTIPLE",
            space_before_pt=0.0,
            space_after_pt=0.0,
            space_before_value=0.0,
            space_before_unit="PT",
            space_after_value=0.0,
            space_after_unit="PT",
            indent_left_pt=0.0,
            indent_right_pt=0.0,
            indent_first_line_pt=0.0,
            indent_hanging_pt=0.0,
        )
    return StyleRule(
        role=role,
        indent_left_pt=0.0,
        indent_right_pt=0.0,
        indent_first_line_pt=0.0,
        indent_hanging_pt=0.0,
    )


def _fill_missing_fields(target: StyleRule, defaults: StyleRule) -> None:
    for field in target.__dataclass_fields__:
        if field in {"role", "missing_fields"}:
            continue
        if getattr(target, field) is None:
            setattr(target, field, getattr(defaults, field))


def _missing_required_fields(rule: StyleRule) -> list[str]:
    missing: list[str] = []
    for field in REQUIRED_FIELDS:
        if getattr(rule, field) is None:
            missing.append(field)
    return missing


def _describe_candidate(candidate: dict[str, Any]) -> str:
    resolved = candidate["resolved"]
    source = candidate.get("source", "unknown")
    name = resolved.name or ""
    style_id = resolved.style_id or ""
    return f"source={source} style_name={name} style_id={style_id}"


def _tag_warnings_for_style(
    log_state: ParseLogState,
    role: str,
    style_id: str | None,
) -> None:
    if not style_id:
        return
    for warning in log_state.warnings:
        if warning.role is None and warning.style_id == style_id:
            warning.role = role


def _warn(
    log_state: ParseLogState,
    rule: str,
    reason: str,
    role: str | None = None,
    style_id: str | None = None,
    paragraph_index: int | None = None,
) -> None:
    log_state.warnings.append(
        WarningEntry(
            rule=rule,
            reason=reason,
            role=role,
            style_id=style_id,
            paragraph_index=paragraph_index,
        )
    )


def _write_log(log_state: ParseLogState) -> None:
    config.ensure_base_dirs()
    log_path = config.build_log_path(log_state.start_time)
    lines = [
        f"template_path: {log_state.template_path}",
        f"elapsed_sec: {log_state.elapsed_sec:.3f}"
        if log_state.elapsed_sec is not None
        else "elapsed_sec: unknown",
        f"styles_count: {log_state.style_count}",
    ]
    if log_state.detected_heading_levels:
        lines.append(f"detected_heading_levels: {log_state.detected_heading_levels}")
    if log_state.detected_heading_levels_overflow:
        lines.append(
            "detected_heading_levels_overflow: "
            f"{log_state.detected_heading_levels_overflow}"
        )
    for role, source in log_state.role_sources.items():
        lines.append(f"role_source[{role}]: {source}")
    for role, candidates in log_state.role_candidates.items():
        for candidate in candidates:
            lines.append(f"role_candidate[{role}]: {candidate}")
    if log_state.error:
        lines.append(f"error: {log_state.error}")
    lines.append(f"warnings_count: {len(log_state.warnings)}")
    for warning in log_state.warnings:
        role = warning.role or "unknown"
        parts = [f"role={role}", f"rule={warning.rule}", f"reason={warning.reason}"]
        if warning.style_id:
            parts.append(f"style_id={warning.style_id}")
        if warning.paragraph_index is not None:
            parts.append(f"paragraph_index={warning.paragraph_index}")
        lines.append("warning: " + " ".join(parts))
    log_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
