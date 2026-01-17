from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED

from lxml import etree
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
FIXTURES_DIR = ROOT / "tests" / "fixtures"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


def _save(doc: Document, name: str) -> Path:
    path = FIXTURES_DIR / name
    doc.save(path)
    return path


def _patch_zip(path: Path, updates: dict[str, bytes]) -> None:
    temp_path = path.with_suffix(".tmp")
    with ZipFile(path, "r") as src, ZipFile(temp_path, "w", ZIP_DEFLATED) as dst:
        for info in src.infolist():
            content = src.read(info.filename)
            if info.filename in updates:
                content = updates[info.filename]
            dst.writestr(info, content)
        for name, content in updates.items():
            if name not in src.namelist():
                dst.writestr(name, content)
    temp_path.replace(path)


def _theme_xml() -> bytes:
    return b"""<?xml version="1.0" encoding="UTF-8"?>
<a:theme xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
  <a:themeElements>
    <a:fontScheme name="Office">
      <a:majorFont>
        <a:latin typeface="ThemeMajorLatin"/>
        <a:ea typeface="ThemeMajorEA"/>
      </a:majorFont>
      <a:minorFont>
        <a:latin typeface="ThemeMinorLatin"/>
        <a:ea typeface="ThemeMinorEA"/>
      </a:minorFont>
    </a:fontScheme>
  </a:themeElements>
</a:theme>
"""


def _remove_spacing_from_style(styles_xml: bytes, style_id: str) -> bytes:
    root = etree.fromstring(styles_xml)
    for style in root.findall("w:style", namespaces=NS):
        sid = style.get(f"{{{W_NS}}}styleId")
        if sid != style_id:
            continue
        p_pr = style.find("w:pPr", namespaces=NS)
        if p_pr is None:
            continue
        for tag in ("w:spacing", "w:before", "w:after"):
            elem = p_pr.find(tag, namespaces=NS)
            if elem is not None:
                p_pr.remove(elem)
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8")


def _set_docdefaults(styles_xml: bytes) -> bytes:
    root = etree.fromstring(styles_xml)
    doc_defaults = root.find("w:docDefaults", namespaces=NS)
    if doc_defaults is None:
        doc_defaults = etree.SubElement(root, f"{{{W_NS}}}docDefaults")
    rpr_default = doc_defaults.find("w:rPrDefault", namespaces=NS)
    if rpr_default is None:
        rpr_default = etree.SubElement(doc_defaults, f"{{{W_NS}}}rPrDefault")
    r_pr = rpr_default.find("w:rPr", namespaces=NS)
    if r_pr is None:
        r_pr = etree.SubElement(rpr_default, f"{{{W_NS}}}rPr")
    for tag in ("w:rFonts", "w:sz", "w:b"):
        elem = r_pr.find(tag, namespaces=NS)
        if elem is not None:
            r_pr.remove(elem)
    r_fonts = etree.SubElement(r_pr, f"{{{W_NS}}}rFonts")
    r_fonts.set(f"{{{W_NS}}}ascii", "Times New Roman")
    r_fonts.set(f"{{{W_NS}}}hAnsi", "Times New Roman")
    r_fonts.set(f"{{{W_NS}}}eastAsia", "宋体")
    sz = etree.SubElement(r_pr, f"{{{W_NS}}}sz")
    sz.set(f"{{{W_NS}}}val", "24")
    bold = etree.SubElement(r_pr, f"{{{W_NS}}}b")
    bold.set(f"{{{W_NS}}}val", "1")

    ppr_default = doc_defaults.find("w:pPrDefault", namespaces=NS)
    if ppr_default is None:
        ppr_default = etree.SubElement(doc_defaults, f"{{{W_NS}}}pPrDefault")
    p_pr = ppr_default.find("w:pPr", namespaces=NS)
    if p_pr is None:
        p_pr = etree.SubElement(ppr_default, f"{{{W_NS}}}pPr")
    for tag in ("w:spacing", "w:jc"):
        elem = p_pr.find(tag, namespaces=NS)
        if elem is not None:
            p_pr.remove(elem)
    spacing = etree.SubElement(p_pr, f"{{{W_NS}}}spacing")
    spacing.set(f"{{{W_NS}}}before", "120")
    spacing.set(f"{{{W_NS}}}after", "240")
    spacing.set(f"{{{W_NS}}}line", "360")
    spacing.set(f"{{{W_NS}}}lineRule", "auto")
    jc = etree.SubElement(p_pr, f"{{{W_NS}}}jc")
    jc.set(f"{{{W_NS}}}val", "center")
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8")


def _strip_style_properties(styles_xml: bytes) -> bytes:
    root = etree.fromstring(styles_xml)
    for style in root.findall("w:style", namespaces=NS):
        style_type = style.get(f"{{{W_NS}}}type")
        if style_type != "paragraph":
            continue
        for tag in ("w:rPr", "w:pPr"):
            elem = style.find(tag, namespaces=NS)
            if elem is not None:
                style.remove(elem)
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8")


def _strip_document_paragraphs(document_xml: bytes) -> bytes:
    root = etree.fromstring(document_xml)
    body = root.find("w:body", namespaces=NS)
    if body is None:
        return document_xml
    body_sect_pr = body.find("w:sectPr", namespaces=NS)
    if body_sect_pr is None:
        sect_pr = root.find(".//w:pPr/w:sectPr", namespaces=NS)
        if sect_pr is not None:
            body_sect_pr = deepcopy(sect_pr)
    for paragraph in list(body.findall("w:p", namespaces=NS)):
        body.remove(paragraph)
    if body_sect_pr is not None and body_sect_pr.getparent() is None:
        body.append(body_sect_pr)
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8")


def _make_basic() -> Path:
    doc = Document()
    doc.add_paragraph("第一章", style="Heading 1")
    doc.add_paragraph("正文内容", style="Normal")
    return _save(doc, "TPL_BASIC.docx")


def _make_theme() -> Path:
    doc = Document()
    doc.add_paragraph("主题标题", style="Heading 1")
    doc.add_paragraph("主题正文", style="Normal")
    path = _save(doc, "TPL_THEME.docx")
    _patch_zip(path, {"word/theme/theme1.xml": _theme_xml()})
    return path


def _make_docdefaults() -> Path:
    doc = Document()
    doc.add_paragraph("默认样式", style="Normal")
    path = _save(doc, "TPL_DOCDEFAULTS.docx")
    with ZipFile(path, "r") as archive:
        styles_xml = archive.read("word/styles.xml")
    patched = _set_docdefaults(styles_xml)
    patched = _strip_style_properties(patched)
    _patch_zip(path, {"word/styles.xml": patched})
    return path


def _make_no_sample() -> Path:
    doc = Document()
    doc.add_paragraph("占位段落，将被移除")
    path = _save(doc, "TPL_NO_SAMPLE.docx")
    with ZipFile(path, "r") as archive:
        document_xml = archive.read("word/document.xml")
    patched = _strip_document_paragraphs(document_xml)
    _patch_zip(path, {"word/document.xml": patched})
    return path


def _make_table() -> Path:
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.paragraphs[0].text = "表格标题"
    cell.paragraphs[0].style = "Heading 1"
    cell.add_paragraph("表格正文", style="Normal")
    return _save(doc, "TPL_TABLE.docx")


def _make_multi_style() -> Path:
    doc = Document()
    styles = doc.styles
    style_a = styles.add_style("TitleA", WD_STYLE_TYPE.PARAGRAPH)
    style_a.base_style = styles["Heading 1"]
    style_b = styles.add_style("TitleB", WD_STYLE_TYPE.PARAGRAPH)
    style_b.base_style = styles["Heading 1"]
    doc.add_paragraph("标题A-1", style="TitleA")
    doc.add_paragraph("标题A-2", style="TitleA")
    doc.add_paragraph("标题B-1", style="TitleB")
    doc.add_paragraph("正文", style="Normal")
    return _save(doc, "TPL_MULTI_STYLE.docx")


def _make_missing_fields() -> Path:
    doc = Document()
    doc.add_paragraph("缺失字段标题", style="Heading 1")
    doc.add_paragraph("缺失字段正文", style="Normal")
    path = _save(doc, "TPL_MISSING_FIELDS.docx")
    with ZipFile(path, "r") as archive:
        styles_xml = archive.read("word/styles.xml")
    patched = _remove_spacing_from_style(styles_xml, "Normal")
    _patch_zip(path, {"word/styles.xml": patched})
    return path


def _make_line_rule() -> Path:
    doc = Document()
    styles = doc.styles
    style_single = styles.add_style("正文_LineSingle", WD_STYLE_TYPE.PARAGRAPH)
    style_single.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    doc.add_paragraph("单倍行距", style="正文_LineSingle")

    style_one_five = styles.add_style("正文_LineOnePointFive", WD_STYLE_TYPE.PARAGRAPH)
    style_one_five.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    doc.add_paragraph("1.5 倍行距", style="正文_LineOnePointFive")

    style_double = styles.add_style("正文_LineDouble", WD_STYLE_TYPE.PARAGRAPH)
    style_double.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    doc.add_paragraph("2 倍行距", style="正文_LineDouble")

    style_multiple = styles.add_style("正文_LineMultiple", WD_STYLE_TYPE.PARAGRAPH)
    style_multiple.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style_multiple.paragraph_format.line_spacing = 1.2
    doc.add_paragraph("多倍行距 1.2", style="正文_LineMultiple")

    style_exactly = styles.add_style("正文_LineExactly", WD_STYLE_TYPE.PARAGRAPH)
    style_exactly.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style_exactly.paragraph_format.line_spacing = Pt(18)
    doc.add_paragraph("固定行距 18pt", style="正文_LineExactly")

    style_at_least = styles.add_style("正文_LineAtLeast", WD_STYLE_TYPE.PARAGRAPH)
    style_at_least.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    style_at_least.paragraph_format.line_spacing = Pt(20)
    doc.add_paragraph("最小行距 20pt", style="正文_LineAtLeast")

    style_inferred = styles.add_style("正文_LineInferred", WD_STYLE_TYPE.PARAGRAPH)
    p = doc.add_paragraph("缺失 lineRule", style="正文_LineInferred")
    p.paragraph_format.line_spacing = Pt(16)
    return _save(doc, "TPL_LINE_RULE.docx")


def _make_align() -> Path:
    doc = Document()
    styles = doc.styles
    left_style = styles.add_style("正文_AlignLeft", WD_STYLE_TYPE.PARAGRAPH)
    left_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph("左对齐", style="正文_AlignLeft")

    center_style = styles.add_style("正文_AlignCenter", WD_STYLE_TYPE.PARAGRAPH)
    center_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("居中", style="正文_AlignCenter")

    right_style = styles.add_style("正文_AlignRight", WD_STYLE_TYPE.PARAGRAPH)
    right_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph("右对齐", style="正文_AlignRight")

    justify_style = styles.add_style("正文_AlignJustify", WD_STYLE_TYPE.PARAGRAPH)
    justify_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph("两端对齐", style="正文_AlignJustify")

    dist_style = styles.add_style("正文_AlignDistribute", WD_STYLE_TYPE.PARAGRAPH)
    dist_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
    doc.add_paragraph("分散对齐", style="正文_AlignDistribute")
    return _save(doc, "TPL_ALIGN.docx")


def _make_invalid() -> Path:
    path = FIXTURES_DIR / "TPL_INVALID.docx"
    path.write_text("invalid docx content", encoding="utf-8")
    return path


def _make_role_map() -> Path:
    data = {
        "Normal": "body",
        "Heading 1": "chapter_title",
        "id:Heading1": "chapter_title",
        "ID:TITLE": "chapter_title",
        "Heading1": "body",
    }
    path = FIXTURES_DIR / "ROLE_MAP.json"
    path.write_text(
        json_dumps(data),
        encoding="utf-8",
    )
    return path


def json_dumps(data: dict) -> str:
    import json

    return json.dumps(data, ensure_ascii=False, indent=2)


def main() -> None:
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)
    _make_basic()
    _make_theme()
    _make_docdefaults()
    _make_no_sample()
    _make_table()
    _make_multi_style()
    _make_missing_fields()
    _make_line_rule()
    _make_align()
    _make_invalid()
    _make_role_map()
    print(f"fixtures generated in {FIXTURES_DIR}")


if __name__ == "__main__":
    main()
